import contextvars, difflib, hashlib, hmac, html, inspect, json, math, os, re, secrets, shutil, smtplib, ssl, uuid, zipfile, requests, gradio as gr
from datetime import date, datetime, timezone
from email.headerregistry import Address
from email.message import EmailMessage
from email.utils import parseaddr
from http.cookies import SimpleCookie
from pathlib import Path
from typing import Iterable
from urllib.parse import parse_qs, quote

from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, Response

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

CORE=os.getenv("CORE_URL","http://core:8080")
RENDERER=os.getenv("RENDERER_URL","http://renderer:8090")
WORKSPACE=Path(os.getenv("GRANT_WORKSPACE","/workspace"))
ORGANIZATION_NAME=os.getenv("ORGANIZATION_NAME","Organization")
CONFIG_ROOT=Path(os.getenv("UI_CONFIG_ROOT",str(Path(__file__).resolve().parents[1]/"config")))
COMPETITIVE_UI_POLL_SECONDS=max(60,min(86400,int(os.getenv("COMPETITIVE_UI_POLL_SECONDS","14400"))))
COLLABORATION_UI_POLL_SECONDS=max(12,min(60,int(os.getenv("COLLABORATION_UI_POLL_SECONDS","12"))))
COMPILATION_UI_POLL_SECONDS=max(65,min(300,int(os.getenv("COMPILATION_UI_POLL_SECONDS","65"))))
COMPETITIVE_UPDATE_LABEL=os.getenv("COMPETITIVE_UPDATE_LABEL","Competitive Edge Auto-Update").strip() or "Competitive Edge Auto-Update"
GRANT_BUILD_VERSION=os.getenv("GRANT_BUILD_VERSION","0.8.0")
REQUEST_IDENTITY_HEADERS=contextvars.ContextVar("request_identity_headers",default=None)

def load_trusted_gateway_secret():
    if os.getenv("AUTH_MODE","local_single_user")!="trusted_headers":return None
    path=os.getenv("TRUSTED_GATEWAY_SECRET_FILE","").strip()
    if not path:raise RuntimeError("TRUSTED_GATEWAY_SECRET_FILE is required in trusted_headers mode")
    value=Path(path).read_text()
    if len(value)!=64 or any(character not in "0123456789abcdefABCDEF" for character in value):
        raise RuntimeError("trusted gateway secret must contain exactly 64 hexadecimal characters")
    return value

TRUSTED_GATEWAY_SECRET=load_trusted_gateway_secret()


def load_default_sections():
    raw=os.getenv("GRANT_SECTIONS","").strip()
    if raw:return [x.strip() for x in raw.split(",") if x.strip()]
    try:return [str(x).strip() for x in json.loads((CONFIG_ROOT/"default_sections.json").read_text()) if str(x).strip()]
    except Exception:return []
DEFAULT_SECTIONS=load_default_sections()

def load_editor_sections():
    key_for=lambda title:re.sub(r"[^a-z0-9]+","_",str(title).lower()).strip("_")
    try:
        items=json.loads((CONFIG_ROOT/"editor_sections.json").read_text())
        return [{"key":key_for(item["title"]),"title":str(item["title"]).strip(),"description":str(item.get("description") or "").strip()} for item in items if str(item.get("title") or "").strip()]
    except Exception:
        return [{"key":key_for(title),"title":title,"description":""} for title in DEFAULT_SECTIONS]

EDITOR_SECTIONS=load_editor_sections()

# Project roles are an API contract, not free-form user data. All project-role
# controls share these constrained values.
PROJECT_ROLE_CHOICES=[
    ("Project owner","owner"),
    ("Principal investigator","pi"),
    ("Contributor / scientific writer","contributor"),
    ("Reviewer","reviewer"),
    ("Approver","approver"),
    ("Research administrator","research_administrator"),
    ("Viewer","viewer"),
]
PROJECT_ROLE_LABELS={value:label for label,value in PROJECT_ROLE_CHOICES}

CSS="""
:root{--gs-bg:#0d0b14;--gs-panel:#15121d;--gs-panel-2:#1b1724;--gs-border:#30293b;--gs-muted:#9a92a5;--gs-copy:#f7f3fb;--gs-purple:#a855f7;--gs-magenta:#d946ef;--gs-cyan:#55d9f2;--gs-green:#42d59c}
.gradio-container{max-width:100%!important;background:var(--gs-bg)!important;color:var(--gs-copy)!important}
.wizard-lightbox{position:fixed!important;inset:0!important;z-index:9999!important;background:radial-gradient(circle at 63% 15%,rgba(88,28,135,.14),transparent 34%),var(--gs-bg)!important;overflow:auto!important}
.wizard-shell{min-height:100vh!important;display:grid;grid-template-columns:minmax(230px,16vw) 1fr!important;gap:0!important}
.wizard-rail{height:100vh;position:sticky;top:0;border-right:1px solid var(--gs-border);padding:38px 24px;background:#100d17;overflow-y:auto}
.wizard-rail-label{font-size:11px;letter-spacing:.22em;text-transform:uppercase;color:#746c7e;margin-bottom:28px}
.wizard-rail-step{display:grid;grid-template-columns:34px 1fr;gap:12px;align-items:center;padding:10px 4px;color:#716a7a}
.wizard-rail-step .number{width:30px;height:30px;border:1px solid #393243;border-radius:50%;display:grid;place-items:center;font-size:12px}
.wizard-rail-step.active{color:#fff}.wizard-rail-step.active .number{background:#7132a7;border-color:#a855f7;box-shadow:0 0 24px rgba(168,85,247,.28)}
.wizard-rail-step.done .number{background:#17362d;border-color:#24463b;color:var(--gs-green)}
.wizard-rail-step b{display:block;font-size:13px}.wizard-rail-step small{display:block;font-size:10px;color:#6d6575;margin-top:3px}
.wizard-current{margin-top:24px;padding:14px;border:1px solid #302a38;border-radius:12px;background:#15111b}.wizard-current b,.wizard-current span{display:block}.wizard-current b{font-size:11px;color:#a855f7;letter-spacing:.08em}.wizard-current span{margin-top:5px;font-size:12px;color:#d9d2df}
.wizard-main{max-width:1480px!important;width:100%!important;margin:auto!important;padding:7vh 6vw 110px!important}
.wizard-lightbox .gr-group,.wizard-lightbox .gr-group>.styler{background:transparent!important;border:0!important}
.wizard-lightbox .block:not(.wizard-panel){background:transparent!important}
.wizard-lightbox input,.wizard-lightbox textarea,.wizard-lightbox [role=listbox],.wizard-lightbox .wrap{background:#18141f!important;color:#f7f3fb!important;border-color:#3a3245!important}
.wizard-lightbox label,.wizard-lightbox .label-wrap,.wizard-lightbox .prose{color:#f7f3fb!important}
.wizard-lightbox .prose h1,.wizard-lightbox .prose h2,.wizard-lightbox .prose h3,.wizard-lightbox .prose p,.wizard-lightbox .prose strong{color:#f7f3fb!important}
.wizard-lightbox button{background:#19151f!important;border-color:#3a3245!important;color:#f7f3fb!important}.wizard-lightbox button.primary{background:linear-gradient(100deg,#7c3aed,#d946ef)!important;border:0!important}
.wizard-kicker{color:#bc6cf6;font-size:11px;letter-spacing:.2em;text-transform:uppercase;font-weight:700}
.wizard-title h1,.wizard-title h2{font-family:Georgia,'Times New Roman',serif!important;font-weight:400!important;line-height:1.06!important;letter-spacing:-.035em!important;color:#fbf8ff!important}.wizard-title h1{font-size:clamp(42px,4.2vw,72px)!important}.wizard-title h2{font-size:clamp(36px,3.4vw,58px)!important}
.wizard-title p{max-width:900px;color:var(--gs-muted);font-size:16px;line-height:1.7}
.wizard-title,.wizard-title *{color:var(--gs-copy)!important}.wizard-title .accent{color:#c06cf4!important}
.wizard-panel{background:rgba(22,18,29,.86)!important;border:1px solid var(--gs-border)!important;border-radius:18px!important;padding:22px!important}
.wizard-hero{border-color:#573468!important;background:linear-gradient(145deg,rgba(168,85,247,.11),rgba(21,18,29,.96))!important;min-height:220px}
.wizard-option-grid .wrap{gap:14px!important}
.wizard-option-grid label{border:1px solid var(--gs-border)!important;border-radius:14px!important;background:var(--gs-panel)!important;padding:18px!important;min-height:72px;transition:.18s ease}.wizard-option-grid label:has(input:checked){border-color:#8e4db8!important;background:#23162d!important;box-shadow:inset 0 0 0 1px rgba(192,108,244,.25)}
.wizard-option-grid input[type=radio],.wizard-choice input[type=radio]{accent-color:#c06cf4!important;width:18px!important;height:18px!important;opacity:1!important}.wizard-option-grid label:has(input:checked) input[type=radio],.wizard-choice label:has(input:checked) input[type=radio]{outline:2px solid #f0c8ff!important;outline-offset:2px!important}
.wizard-summary-table{display:flex!important;flex-direction:column!important;gap:0!important;background:transparent!important}.wizard-summary-row{display:grid!important;grid-template-columns:minmax(0,1fr) minmax(330px,430px)!important;gap:24px!important;align-items:center!important;padding:16px 4px!important;border:0!important;border-bottom:1px solid rgba(75,63,87,.42)!important;background:transparent!important}.wizard-summary-row:last-child{border-bottom:0!important}.wizard-summary-copy h3{margin:0 0 4px!important;font-size:16px!important}.wizard-summary-copy p{margin:0!important;color:#a69dac!important;font-size:12px!important;line-height:1.5!important}.wizard-summary-copy small{display:block!important;margin-top:6px!important;color:#766d7e!important}.wizard-choice{min-width:0!important}.wizard-choice .wrap{padding:5px!important}.wizard-choice label{min-height:42px!important;padding:8px 10px!important;border-radius:8px!important}.core-included{justify-self:end;padding:8px 12px;border-radius:999px;background:#18362d;color:#62deb0;font-size:12px;font-weight:800}.skip-all-panel{display:flex;align-items:center;justify-content:space-between;gap:18px;padding:13px 16px;margin-bottom:6px;border-radius:10px;background:#18141f}.skip-all-panel p{margin:0;color:#bcb3c4;font-size:12px}.skip-all-panel button{width:auto!important;min-width:230px!important}
.core-flow{display:grid;grid-template-columns:1fr;gap:0;margin:24px 0;border:1px solid var(--gs-border);border-radius:18px;overflow:hidden;background:var(--gs-panel)}
.core-step{display:grid;grid-template-columns:54px 1fr minmax(150px,240px);gap:16px;align-items:center;padding:20px 24px;border-bottom:1px solid var(--gs-border);min-height:98px}.core-step:last-child{border-bottom:0}.core-step .step-no{height:44px;width:44px;border:1px solid #56326c;border-radius:12px;display:grid;place-items:center;color:#c06cf4}.core-step b{display:block;color:#f5eff9;margin-bottom:5px}.core-step span,.core-step small{font-size:13px;color:var(--gs-muted)}.core-step .output{text-align:right;font-size:12px;color:#b2a8bc}
.module-catalog{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:14px;margin:16px 0 24px}.module-card{border:1px solid var(--gs-border);border-radius:15px;padding:18px;background:var(--gs-panel);min-height:155px}.module-card.gated{border-left:3px solid #8f45bc}.module-card b{color:#f4eef8}.module-card .placement{color:#b25ee8;font-size:11px;margin:8px 0}.module-card p{color:var(--gs-muted);font-size:12px;line-height:1.55}.module-card small{color:#746c7e}
.workflow-preview{border:1px solid var(--gs-border);border-radius:18px;background:var(--gs-panel);padding:20px}.workflow-preview-row{display:grid;grid-template-columns:38px 1fr auto;gap:14px;align-items:center;padding:14px 4px;border-bottom:1px solid #292332}.workflow-preview-row:last-child{border-bottom:0}.workflow-preview-row .n{height:34px;width:34px;border:1px solid #6a3a83;border-radius:9px;display:grid;place-items:center;color:#c06cf4}.workflow-preview-row small{color:var(--gs-muted)}
.privacy-note{border:1px solid #24444d;border-radius:14px;padding:16px;background:rgba(25,92,105,.12);color:#bcecf3}
#wizard-create-progress{display:none;margin:18px 0;padding:14px 16px;border:1px solid #3d3150;border-radius:12px;background:#15121d}#wizard-create-progress.active,#wizard-create-progress.complete,#wizard-create-progress.failed{display:block}#wizard-create-progress .track{height:10px;overflow:hidden;border-radius:999px;background:#282231;margin-top:10px}#wizard-create-progress .bar{height:100%;width:34%;border-radius:999px;background:linear-gradient(90deg,#7c3aed,#d946ef,#55d9f2);animation:wizard-progress 1.4s ease-in-out infinite alternate}#wizard-create-progress.complete .bar{width:100%;animation:none;background:#42d59c}#wizard-create-progress.failed .bar{width:100%;animation:none;background:#ef476f}@keyframes wizard-progress{from{transform:translateX(-90%)}to{transform:translateX(280%)}}
#refresh-shared-updates{position:fixed!important;top:14px;right:18px;z-index:10005;width:auto!important;min-width:210px!important;background:#19151f!important;border:1px solid #65417a!important;color:#f7f3fb!important;box-shadow:0 10px 30px #0007!important}
.wizard-footer{position:fixed!important;z-index:10001;bottom:0;left:0;right:0;border-top:1px solid var(--gs-border);background:rgba(13,11,20,.96);backdrop-filter:blur(14px);padding:14px 28px!important}
.wizard-footer button.primary{background:linear-gradient(100deg,#7c3aed,#d946ef)!important;border:0!important;color:white!important}
.workspace-shell{background:var(--gs-bg)!important;min-height:100vh}.workspace-shell h1,.workspace-shell h2,.workspace-shell h3{color:#f7f3fb}
.team-chat{max-height:520px;overflow:auto;border:1px solid var(--gs-border);border-radius:14px;padding:14px;background:var(--gs-panel)}.chat-message{background:var(--gs-panel-2);border:1px solid var(--gs-border);border-radius:12px;padding:10px 12px;margin:8px 0}.chat-meta{font-size:11px;color:var(--gs-muted);margin-bottom:4px}
.page-frame{background:#e7e7e7;padding:24px;border-radius:14px;overflow:auto}
.page-frame iframe{width:100%;min-height:11.2in;border:0;background:#e7e7e7}
.version-diff{overflow:auto;border:1px solid var(--gs-border);border-radius:12px;background:#fff;color:#17131c;padding:12px}.version-diff table{border-collapse:collapse;width:100%;font-family:ui-monospace,SFMono-Regular,Menlo,monospace;font-size:12px}.version-diff th,.version-diff td{padding:3px 6px;vertical-align:top;white-space:pre-wrap}.version-diff .diff_header{background:#ece8f0}.version-diff .diff_add{background:#d9f8e9}.version-diff .diff_sub{background:#ffe1e1}.version-diff .diff_chg{background:#fff2b8}
.status{font-size:12px;padding:8px 12px;border-radius:8px;background:var(--gs-panel)}
.question-card{padding:18px;border:1px solid var(--gs-border);border-radius:10px;background:var(--gs-panel)}
.global-nav{width:100%;box-sizing:border-box;display:flex;align-items:center;gap:18px;padding:9px 20px;margin:0 0 12px;border-bottom:1px solid #33283d;background:#100d17}.global-nav .brand{margin-right:auto;color:#f5f0fa;font-size:13px;font-weight:800;letter-spacing:.04em}.global-nav button{appearance:none;border:0!important;background:transparent!important;padding:2px 0!important;min-height:22px!important;min-width:0!important;font-size:12px!important;font-weight:750!important;cursor:pointer;background-image:linear-gradient(90deg,#d7b5ff,#8b5cf6)!important;background-clip:text!important;-webkit-background-clip:text!important;color:transparent!important;-webkit-text-fill-color:transparent!important;text-shadow:1px 1px 1px #fff8!important}.global-nav button:hover{filter:brightness(1.25)}.global-nav a{font-size:12px;font-weight:700;color:#a998b8;text-decoration:none}.global-nav a:hover{color:#fff}
.grant-editor-shell{width:calc(100vw - 40px)!important;max-width:none!important;margin-left:calc(50% - 50vw + 20px)!important;margin-right:20px!important;padding:4px 0 110px!important}.grant-editor-header{display:flex;align-items:center;justify-content:space-between;padding:6px 0 12px;border-bottom:1px solid var(--gs-border);margin-bottom:10px}.grant-editor-header h2{font-size:20px!important;margin:0 0 2px!important}
.grant-editor-layout{display:grid!important;grid-template-columns:minmax(220px,300px) minmax(700px,1fr) minmax(220px,280px)!important;align-items:start!important;gap:12px!important}.grant-outline-panel{width:100%!important;min-width:0!important;max-width:300px!important;position:sticky!important;top:12px!important;background:#111019!important;border:1px solid var(--gs-border)!important;border-radius:10px!important;padding:9px!important}.grant-document-panel{width:100%!important;min-width:0!important;max-width:none!important;background:#e7e8eb!important;color:#19151f!important;border-radius:10px!important;padding:10px!important;box-shadow:0 14px 38px #0004}.grant-guidance-panel{width:100%!important;min-width:0!important;max-width:280px!important;position:sticky!important;top:12px!important;background:#15121d!important;border:1px solid var(--gs-border)!important;border-radius:10px!important;padding:10px!important}
.grant-editor-shell button{min-height:30px!important;padding:4px 10px!important;border-radius:6px!important;font-size:12px!important}.grant-editor-toolbar{align-items:center!important;gap:7px!important;margin-bottom:8px!important}.grant-editor-toolbar>div{min-width:0!important}.grant-editor-toolbar label{font-size:10px!important}.grant-editor-toolbar input{min-height:30px!important;font-size:12px!important}.grant-document-scroll{height:calc(100vh - 315px);min-height:620px;overflow:auto;scroll-behavior:smooth;padding:24px 22px 70px;background:#dfe1e5;border:1px solid #c5c7cb;border-radius:8px}.grant-doc-section{box-sizing:border-box;max-width:920px;min-height:420px;margin:0 auto 24px;padding:54px 68px 70px;background:#fff;color:#1d1b20;border:1px solid #c9c9c9;box-shadow:0 2px 9px #0002;scroll-margin-top:18px}.grant-doc-section.selected{box-shadow:0 0 0 2px #8b5cf6,0 3px 12px #0002}.grant-doc-version{font:600 10px/1.2 system-ui,-apple-system,sans-serif;color:#7c7680;text-transform:uppercase;letter-spacing:.08em;margin-bottom:22px}.unsaved-label{display:none;margin-left:10px;color:#8a6110}.grant-doc-section.unsaved .unsaved-label{display:inline}.grant-doc-section [contenteditable=true]{outline:none;border-radius:3px}.grant-doc-section [contenteditable=true]:focus{box-shadow:0 0 0 2px #9bc1ff;background:#f8fbff}.grant-doc-section h2{font:700 30px/1.25 Georgia,'Times New Roman',serif;margin:0 0 7px}.grant-doc-description{min-height:20px;margin-bottom:28px;color:#6b6570;font:italic 13px/1.45 system-ui,-apple-system,sans-serif}.grant-doc-body{min-height:250px;white-space:pre-wrap;font:17px/1.72 Georgia,'Times New Roman',serif}.grant-doc-body:empty:before{content:'Write here, or choose Rewrite selected to create an evidence-grounded draft.';color:#aaa4ad}.grant-doc-description:empty:before{content:'Optional section purpose or drafting direction';color:#aaa4ad}
.editor-outline{display:flex;flex-direction:column;gap:7px}.editor-outline-head{display:flex;justify-content:space-between;align-items:center;padding:4px 4px 10px;color:#c9c1d0}.editor-outline-item{display:grid;grid-template-columns:20px 1fr auto;gap:8px;align-items:center;padding:10px 8px;border:1px solid transparent;border-radius:10px;background:#191620;color:#ded7e4;cursor:pointer}.editor-outline-item:hover{border-color:#52465f}.editor-outline-item.active{background:#263a2b;border-color:#5e8f63;color:white}.editor-outline-item.dragging{opacity:.45}.editor-outline-item .drag{cursor:grab;color:#7e7488}.editor-outline-item .outline-actions{display:flex;gap:2px}.editor-outline-item button{border:0!important;background:transparent!important;color:#aba1b5!important;padding:2px 4px!important;min-width:20px!important}.editor-outline-item button:hover{color:white!important}.editor-outline-description{font-size:10px;color:#8d8496;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;max-width:185px}.guidance-card{border:1px solid #332b3e;background:#1a1622;border-radius:10px;padding:10px;margin:8px 0}.guidance-meta{font-size:10px;color:#a89fb1;margin-bottom:5px}.guidance-type{font-size:10px;letter-spacing:.08em;font-weight:800;color:#a6d490}.guidance-resolved{opacity:.55}.editor-empty{padding:24px;border:1px dashed #aaa;border-radius:12px;background:#fff;color:#554d59}
#publish-grant-floating{position:fixed!important;left:22px!important;bottom:22px!important;z-index:9000!important;width:auto!important;min-width:190px!important;background:linear-gradient(to right,#b1ce8c 0%,#8ab66b 56%,#63bf4e 100%)!important;border:1px solid #7ca567!important;color:white!important;font-weight:800!important;text-shadow:2px 2px 2px #000!important;box-shadow:0 10px 28px #0007!important}
#editor-outline-command,#editor-document-payload,#global-navigation-command{display:none!important}.editor-secondary{font-size:12px;color:#9d94a6}
@media(max-width:900px){.wizard-shell{grid-template-columns:1fr!important}.wizard-rail{height:auto;position:static;padding:12px;display:flex;overflow:auto}.wizard-rail-label,.wizard-rail-step b,.wizard-rail-step small{display:none}.wizard-rail-step{display:block}.wizard-main{padding:35px 20px 100px!important}.module-catalog{grid-template-columns:1fr}.core-step{grid-template-columns:48px 1fr}.core-step .output{display:none}}
@media(max-width:1250px){.grant-editor-layout{grid-template-columns:minmax(190px,240px) minmax(560px,1fr) minmax(190px,230px)!important}.grant-doc-section{padding:42px 46px 58px}.grant-document-scroll{padding:18px 14px 60px}}
@media(max-width:980px){.grant-editor-shell{width:calc(100vw - 24px)!important;margin-left:calc(50% - 50vw + 12px)!important}.grant-editor-layout{display:block!important}.grant-outline-panel,.grant-guidance-panel{position:static!important;max-width:none!important;margin-bottom:10px}.grant-document-scroll{height:65vh;min-height:500px}.grant-doc-section{padding:36px 30px 50px}}
"""

def core_request_headers(method):
    headers={}
    if method.upper() not in {"GET","HEAD","OPTIONS"}:
        headers["Idempotency-Key"]=str(uuid.uuid4())
    auth_mode=os.getenv("AUTH_MODE","local_single_user")
    if auth_mode=="trusted_headers":
        identity=REQUEST_IDENTITY_HEADERS.get() or {}
        headers.update(identity)
        headers["X-Grantspace-Gateway-Secret"]=TRUSTED_GATEWAY_SECRET
        if "X-Grantspace-User-Id" not in headers or "X-Grantspace-Organization-Id" not in headers:
            raise gr.Error("The authenticated gateway did not provide a stable user and organization identity for this browser session.")
    elif auth_mode=="internal_accounts":
        identity=REQUEST_IDENTITY_HEADERS.get() or {}
        authorization=identity.get("Authorization")
        if not authorization:raise gr.Error("Your login session is missing or expired. Sign in again.")
        headers["Authorization"]=authorization
    return headers

def gateway_callback(function):
    """Bind one Gradio callback to the identity headers on its own browser request."""
    signature=inspect.signature(function)
    has_varargs=any(parameter.kind==inspect.Parameter.VAR_POSITIONAL for parameter in signature.parameters.values())
    request_kind=inspect.Parameter.KEYWORD_ONLY if has_varargs else inspect.Parameter.POSITIONAL_OR_KEYWORD
    request_parameter=inspect.Parameter("request",request_kind,default=None,annotation=gr.Request)
    def wrapped(*args,**kwargs):
        request=kwargs.pop("request",None)
        positional_request=request is None and bool(args) and isinstance(args[-1],gr.Request)
        if positional_request:request=args[-1]
        call_args=args[:-1] if positional_request else args
        identity={}
        if request is not None:
            incoming={str(key).lower():str(value).strip() for key,value in dict(request.headers).items()}
            if os.getenv("AUTH_MODE","local_single_user")=="trusted_headers":
                supplied=incoming.get("x-grantspace-gateway-secret","")
                if not TRUSTED_GATEWAY_SECRET or not hmac.compare_digest(supplied,TRUSTED_GATEWAY_SECRET):
                    raise gr.Error("The request did not originate from the authenticated gateway.")
            for header in ("x-grantspace-user-id","x-grantspace-organization-id","x-grantspace-user-email","x-grantspace-user-name"):
                if incoming.get(header):identity["-".join(part.capitalize() for part in header.split("-"))]=incoming[header]
            if os.getenv("AUTH_MODE","local_single_user")=="internal_accounts":
                session_token=incoming.get("x-grantspace-session")
                if not session_token:session_token=dict(getattr(request,"cookies",{}) or {}).get("grantspace_session")
                if not session_token and incoming.get("cookie"):
                    parsed_cookies=SimpleCookie()
                    try:parsed_cookies.load(incoming["cookie"])
                    except Exception:parsed_cookies=SimpleCookie()
                    morsel=parsed_cookies.get("grantspace_session")
                    if morsel:session_token=morsel.value
                if session_token:identity["Authorization"]=f"Bearer {session_token}"
        token=REQUEST_IDENTITY_HEADERS.set(identity)
        try:return function(*call_args,**kwargs)
        finally:REQUEST_IDENTITY_HEADERS.reset(token)
    wrapped.__name__=f"authenticated_{getattr(function,'__name__','callback')}"
    wrapped.__signature__=signature.replace(parameters=[*signature.parameters.values(),request_parameter])
    wrapped.__annotations__={**getattr(function,"__annotations__",{}),"request":gr.Request}
    return wrapped

def api(method,path,**kwargs):
    headers=core_request_headers(method)
    headers.update(kwargs.pop("headers",{}) or {})
    r=requests.request(method,f"{CORE}{path}",timeout=kwargs.pop("timeout",300),headers=headers,**kwargs)
    if not r.ok:
        try:
            payload=r.json()
            detail=payload.get("error") or payload.get("detail") or r.text
        except Exception:detail=r.text
        detail=str(detail or "").strip()
        if not detail:detail=f"{method.upper()} {path} failed with HTTP {r.status_code} {r.reason or 'error'}"
        raise gr.Error(detail)
    return r.json()

def renderer_api(path,payload,timeout=180):
    r=requests.post(f"{RENDERER}{path}",json=payload,timeout=timeout)
    if not r.ok:
        try:detail=r.json().get("detail",r.text)
        except Exception:detail=r.text
        raise gr.Error(detail)
    return r.json()

def file_path(value):
    if value is None:return None
    return str(getattr(value,"name",value))

def iter_docx_blocks(doc):
    for child in doc.element.body.iterchildren():
        if isinstance(child,CT_P):yield Paragraph(child,doc)
        elif isinstance(child,CT_Tbl):yield Table(child,doc)

def extract_docx(path):
    doc=DocxDocument(path);parts=[]
    for block in iter_docx_blocks(doc):
        if isinstance(block,Paragraph):
            if block.text.strip():parts.append(block.text.strip())
        else:
            for row in block.rows:
                cells=[c.text.strip() for c in row.cells]
                if any(cells):parts.append(" | ".join(cells))
    for label,collection in (("HEADER",[s.header for s in doc.sections]),("FOOTER",[s.footer for s in doc.sections])):
        seen=set()
        for item in collection:
            container=[]
            container.extend(p.text.strip() for p in item.paragraphs if p.text.strip())
            for table in item.tables:
                for row in table.rows:
                    cells=[c.text.strip() for c in row.cells]
                    if any(cells):container.append(" | ".join(cells))
            text="\n".join(container)
            if text and text not in seen:parts.append(f"[{label}]\n{text}");seen.add(text)
    return "\n\n".join(parts)

def extract_file(value):
    p=Path(file_path(value));ext=p.suffix.lower()
    if ext==".pdf":
        # Form-feed is an immutable page boundary in the stored source buffer.
        # Rust uses it to derive source_page while copying exact source spans.
        reader=PdfReader(str(p));text="\n\f\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
        if len(text)<40:raise gr.Error(f"{p.name} has no usable text layer. Please upload a searchable/OCR-processed PDF rather than allowing silent requirement loss.")
        return text
    if ext==".docx":
        text=extract_docx(str(p)).strip()
        if not text:raise gr.Error(f"{p.name} contains no readable paragraphs/tables/headers/footers.")
        return text
    if ext in {".html",".htm"}:return BeautifulSoup(p.read_text(errors="ignore"),"html.parser").get_text("\n",strip=True)
    if ext in {".txt",".md",".csv",".json",".xml"}:return p.read_text(errors="ignore")
    raise gr.Error(f"Unsupported text-ingestion file type: {ext or 'unknown'}. Images belong in Branding/Layout Inspiration, not grant evidence ingestion.")

def push_doc(project,name,kind,text):
    if not text.strip():return False
    return bool(api("POST",f"/api/projects/{project}/documents",json={"name":name,"kind":kind,"text":text}).get("added",False))

def copy_brand_assets(project,assets):
    dest=WORKSPACE/"projects"/project/"branding";dest.mkdir(parents=True,exist_ok=True);out=[]
    for item in assets or []:
        src=Path(file_path(item)); digest=hashlib.sha256(src.read_bytes()).hexdigest()[:12]; target=dest/src.name
        if target.exists() and hashlib.sha256(target.read_bytes()).hexdigest()[:12]!=digest:
            target=dest/f"{src.stem}_{digest}{src.suffix.lower()}"
        if not target.exists():shutil.copy2(src,target)
        out.append(str(target))
    return out

def build_design_profile(project,sponsor,assets):
    return renderer_api("/design-profile",{"project_id":project,"sponsor":sponsor or None,"organization_name":ORGANIZATION_NAME,"asset_paths":assets})

def competitive_update_reason_label(update):
    reasons=set(str(x) for x in ((update or {}).get("refresh_reason") or []))
    if "public_intelligence_refresh_due" in reasons:
        return "Fresh public competitor intelligence was discovered during the scheduled refresh"
    if "competitive_config_changed" in reasons:
        return "Competitive-scoring/search configuration changed and public intelligence was recomputed"
    if "project_inputs_changed" in reasons:
        return "The grant or clinical study changed, so competitor intelligence was rebuilt against the new design"
    if "manual_force" in reasons:
        return "A fresh public competitor scan found a material positioning change"
    return "Fresh public competitive intelligence changed the positioning context"

def section_preview(project,project_title,section_title,body,section_key=None,version=None,competitive_update=None):
    if not project:return '<div class="page-frame"><div style="background:white;padding:32px">Create or open a project to preview a section.</div></div>'
    if competitive_update and competitive_update.get("status")=="pending" and competitive_update.get("base_body") is not None:
        includes_human_edits=bool(version and competitive_update.get("proposed_version") and int(version)!=int(competitive_update.get("proposed_version")))
        d=renderer_api("/preview-diff",{"project_id":project,"title":project_title or "Grant","organization_name":ORGANIZATION_NAME,"section":{"section_key":section_key,"title":section_title or "Section","body":body or "","version":version},"baseline_body":competitive_update.get("base_body") or "","update_summary":competitive_update.get("summary") or "","update_reason":competitive_update_reason_label(competitive_update),"includes_human_edits":includes_human_edits})
    else:
        d=renderer_api("/preview",{"project_id":project,"title":project_title or "Grant","organization_name":ORGANIZATION_NAME,"sections":[{"section_key":section_key,"title":section_title or "Section","body":body or "","version":version}],"include_document_title":False})
    doc=html.escape(d["html"],quote=True)
    return f'<div class="page-frame"><iframe sandbox="" srcdoc="{doc}"></iframe></div>'

def competitive_update_banner(update,current_version=None):
    if not update or update.get("status")!="pending":return ""
    summary=update.get("summary") or "New public competitive information changed the positioning strategy."
    reason=competitive_update_reason_label(update)
    edited_after=bool(current_version and update.get("proposed_version") and int(current_version)!=int(update.get("proposed_version")))
    highlight=("Highlighted text shows all differences from the pre-update version, including edits you made after the automatic proposal."
               if edited_after else
               "Highlighted text shows the proposed changes caused by the refreshed competitive intelligence.")
    return (f"### ⚡ {COMPETITIVE_UPDATE_LABEL}\n"
            f"**{reason}.** {summary}\n\n"
            f"This section was refreshed automatically. **{highlight}** "
            "Review it, edit it if needed, then approve the exact version you want in the final grant. Your previously approved version remains protected until you approve again.")

def global_competitive_update_banner(data):
    if not data:return ""
    pending=int(data.get("pending") or 0)
    processing=int(data.get("processing_pending") or 0)
    events=data.get("events") or []
    latest=events[0] if events else {}
    if pending:
        pending_sections=data.get("pending_sections") or []
        names=[str(x.get("title") or x.get("section_key") or "Section") for x in pending_sections[:8]]
        section_line=("\n\n**Affected sections:** "+", ".join(names)+(f" +{len(pending_sections)-8} more" if len(pending_sections)>8 else "")) if names else ""
        return (f"## ⚡ {COMPETITIVE_UPDATE_LABEL}\n"
                f"Fresh public competitor data produced **{pending} highlighted grant section update(s) waiting for your review.** "
                "Human-approved text was not silently overwritten. Open an affected section to see exactly what changed, then edit or approve it."
                f"{section_line}\n\n{latest.get('summary','')}")
    if processing:
        return (f"## ⚡ {COMPETITIVE_UPDATE_LABEL}\n"
                "A fresh competitive-intelligence update is being reconciled with the grant. The agent will retry any incomplete section refresh automatically; final export remains protected until reconciliation finishes.\n\n"
                f"{latest.get('summary','')}")
    if latest.get("material"):
        return f"**Competitive intelligence is current.** Latest auto-update: {latest.get('summary','')}"
    return ""


def poll_competitive_updates(project):
    if not project:return ""
    try:
        data=api("GET",f"/api/projects/{project}/competitive/updates",timeout=2400)
        return global_competitive_update_banner(data)
    except Exception as exc:
        # Background public-intelligence refresh must not destroy an in-progress edit.
        # Surface the retry state in the banner; the backend remains fail-closed for export.
        return f"## ⚡ {COMPETITIVE_UPDATE_LABEL}\nAutomatic competitor refresh will retry. Current detail: `{html.escape(str(exc))}`"

def full_document_preview(project,payload):
    if not project or not payload.get("sections"):
        return '<div class="page-frame"><div style="background:white;padding:32px">No human-approved sections have been added to the final document yet.</div></div>'
    meta=payload.get("project") or {}
    d=renderer_api("/preview",{
        "project_id":project,
        "title":meta.get("title") or "Grant Application",
        "sponsor":meta.get("sponsor"),
        "organization_name":ORGANIZATION_NAME,
        "sections":[{"section_key":x.get("section_key"),"title":x.get("title") or "Section","body":x.get("body") or "","version":x.get("version")} for x in payload.get("sections",[])],
        "include_document_title":True,
        "design_profile":payload.get("design_profile"),
    })
    doc=html.escape(d["html"],quote=True)
    return f'<div class="page-frame"><iframe sandbox="" srcdoc="{doc}"></iframe></div>'

def requirement_rows(reqs):
    return [[r.get("requirement"),"Required" if r.get("mandatory") else "Optional",", ".join(r.get("evidence_needed") or []) or "No separate evidence specified",str(r.get("status") or "Needs review").replace("_"," ").title()] for r in reqs]

def model_response_preview(value,word_limit=20):
    words=str(value or "").split()
    return html.escape(" ".join(words[:word_limit]))+("…" if words else "")

def requirements_response_preview(requirements):
    text=" ".join(str(item.get("requirement") or "") for item in (requirements or []))
    return model_response_preview(text)

def generation_elapsed(started_at):
    try:
        started=datetime.strptime(str(started_at),"%Y-%m-%d %H:%M:%S").replace(tzinfo=timezone.utc)
        seconds=max(0,int((datetime.now(timezone.utc)-started).total_seconds()))
        minutes,seconds=divmod(seconds,60)
        return f"{minutes} min {seconds} sec"
    except (TypeError,ValueError):return "an unknown duration"
def evidence_rows(items):return [[e.get("id"),e.get("requirement_id"),e.get("source_type"),e.get("claim"),e.get("status"),round(float(e.get("confidence",0)),2),e.get("url") or ""] for e in items]
def slug(s):
    out=[];last=False
    for ch in (s or ""):
        if ch.isalnum():out.append(ch.lower());last=False
        elif out and not last:out.append("_");last=True
    return "".join(out).strip("_")

def load_workflow_registry():
    """Load the same versioned registry used by the Rust gate evaluator."""
    configured=Path(os.getenv("WORKFLOW_DEFINITIONS_PATH",str(CONFIG_ROOT/"workflow_definitions.json")))
    candidates=[configured,CONFIG_ROOT/"workflow_definitions.json",Path("/app/config/workflow_definitions.json")]
    for candidate in candidates:
        try:
            data=json.loads(candidate.read_text())
            if len(data.get("core_steps") or [])!=5:raise ValueError("registry must contain five core steps")
            return data
        except FileNotFoundError:continue
    raise RuntimeError("workflow_definitions.json is required; the UI will not invent a workflow")

WORKFLOW_REGISTRY=load_workflow_registry()
WORKFLOW_MODULES={item["key"]:item for item in WORKFLOW_REGISTRY["optional_modules"]}
WORKFLOW_PRESETS={item["key"]:item for item in WORKFLOW_REGISTRY["presets"]}
MODULE_CHOICES=[(item["title"],item["key"]) for item in WORKFLOW_REGISTRY["optional_modules"]]
GATE_CHOICES=[(item["title"],item["key"]) for item in WORKFLOW_REGISTRY["optional_modules"] if item.get("gate_configurable",False)]
PRESET_CHOICES=[(item["title"],item["key"]) for item in WORKFLOW_REGISTRY["presets"] if item["key"]!=WORKFLOW_REGISTRY["legacy_preset_key"]]
REVIEW_MODE_CHOICES=[(item["title"],item["key"]) for item in WORKFLOW_REGISTRY["review_modes"]]
REVIEWER_ARCHETYPE_ROWS=[[item["title"],item["description"]] for item in WORKFLOW_REGISTRY["reviewer_archetypes"]]

def core_flow_html():
    rows=[]
    for index,step in enumerate(WORKFLOW_REGISTRY["core_steps"],1):
        rows.append(
            f'<div class="core-step"><div class="step-no">{index:02d}</div><div><b>{html.escape(step["title"])}</b>'
            f'<span>{html.escape(step["description"])}</span></div><div class="output">OUTPUT<br><small>{html.escape(step["output"])}</small></div></div>'
        )
    return '<div class="core-flow">'+"".join(rows)+"</div>"

def module_catalog_html():
    cards=[]
    for item in WORKFLOW_REGISTRY["optional_modules"]:
        configurable=item.get("gate_configurable",False)
        gate="Optional gate available" if configurable else "Advisory tool only"
        cards.append(
            f'<article class="module-card"><b>{html.escape(item["title"])}</b>'
            f'<div class="placement">{html.escape(item["placement"].replace("_"," ").upper())} · {gate}</div>'
            f'<p>{html.escape(item["description"])}</p><small>Produces: {html.escape(item["output"])}<br>{html.escape(item["runtime_implication"])}</small></article>'
        )
    return '<div class="module-catalog">'+"".join(cards)+"</div>"

def workflow_item_screen_html(item,kind,position,total):
    label="Core outcome" if kind=="core" else "Optional tool"
    placement=(item.get("placement") or "ordered").replace("_"," ")
    runtime=item.get("runtime_implication")
    runtime_html=(f'<div class="privacy-note"><b>Runtime and data use</b><br>{html.escape(runtime)}</div>' if runtime else "")
    return (
        f'<div class="wizard-title"><div class="wizard-kicker">{html.escape(label)} · {position} of {total}</div>'
        f'<h2>{html.escape(item["title"])}</h2><p>{html.escape(item["description"])}</p></div>'
        f'<div class="wizard-panel"><p><b>Produces</b><br>{html.escape(item["output"])}</p>'
        f'<p><b>Where it appears</b><br>{html.escape(placement.title())}</p></div>{runtime_html}'
    )

def selected_modules_from_modes(*modes):
    if len(modes)!=len(WORKFLOW_REGISTRY["optional_modules"]):
        raise gr.Error("The optional-tool choices are incomplete. Reopen the setup wizard and choose Include or Skip on every tool screen.")
    return [
        module["key"]
        for module,mode in zip(WORKFLOW_REGISTRY["optional_modules"],modes)
        if mode=="include"
    ]

WIZARD_CORE_PAGE=4
WIZARD_MODULE_PAGE=5
WIZARD_REVIEW_PAGE=6
WIZARD_TEAM_PAGE=WIZARD_REVIEW_PAGE+1
WIZARD_ROUTING_PAGE=WIZARD_TEAM_PAGE+1
WIZARD_PREVIEW_PAGE=WIZARD_ROUTING_PAGE+1
WIZARD_PAGE_COUNT=WIZARD_PREVIEW_PAGE
WIZARD_PAGE_TITLES=["Start","Grant details","Grant source","Core workflow","Optional tools","Review setup","Team","Model routing","Workflow preview"]
WIZARD_RAIL_SECTIONS=[
    ("Start",1,1,"Create or open"),
    ("Grant ask",2,3,"Details and source"),
    ("Core outcomes",WIZARD_CORE_PAGE,WIZARD_CORE_PAGE,"Complete core workflow"),
    ("Optional tools",WIZARD_MODULE_PAGE,WIZARD_MODULE_PAGE,"Choose tools together"),
    ("Review setup",WIZARD_REVIEW_PAGE,WIZARD_REVIEW_PAGE,"Advisory configuration"),
    ("Team",WIZARD_TEAM_PAGE,WIZARD_TEAM_PAGE,"Invite collaborators"),
    ("Routing",WIZARD_ROUTING_PAGE,WIZARD_ROUTING_PAGE,"Choose model policy"),
    ("Preview",WIZARD_PREVIEW_PAGE,WIZARD_PREVIEW_PAGE,"Confirm and create"),
]

def wizard_rail_html(active):
    rows=[]
    for index,(title,first,last,subtitle) in enumerate(WIZARD_RAIL_SECTIONS,1):
        state="done" if active>last else ("active" if first<=active<=last else "")
        marker="✓" if active>last else str(index)
        rows.append(f'<div class="wizard-rail-step {state}"><span class="number">{marker}</span><span><b>{html.escape(title)}</b><small>{html.escape(subtitle)}</small></span></div>')
    current=WIZARD_PAGE_TITLES[active-1]
    return (f'<nav class="wizard-rail"><div class="wizard-rail-label">Compose workflow</div>'
            +"".join(rows)
            +f'<div class="wizard-current"><b>{active} of {WIZARD_PAGE_COUNT}</b><span>{html.escape(current)}</span></div></nav>')

def wizard_page_updates(active):
    return [gr.Column(visible=index==active) for index in range(1,WIZARD_PAGE_COUNT+1)]+[wizard_rail_html(active),f"{active} of {WIZARD_PAGE_COUNT}"]

def wizard_go(active):
    return lambda: wizard_page_updates(active)

def wizard_nav_js(active):
    return f"""() => {{
      for (let index = 1; index <= {WIZARD_PAGE_COUNT}; index += 1) {{
        const page = document.getElementById(`wizard-page-${{index}}`);
        if (page) page.style.setProperty('display', index === {int(active)} ? 'flex' : 'none', 'important');
      }}
      return [];
    }}"""

def wizard_nav_from_progress_js():
    return f"""() => {{
      const progress = document.getElementById('wizard-progress');
      const match = progress && progress.textContent.match(/(\\d+)\\s+of\\s+\\d+/);
      const active = match ? Number(match[1]) : 1;
      for (let index = 1; index <= {WIZARD_PAGE_COUNT}; index += 1) {{
        const page = document.getElementById(`wizard-page-${{index}}`);
        if (page) page.style.setProperty('display', index === active ? 'flex' : 'none', 'important');
      }}
      return [];
    }}"""

WIZARD_CREATE_CLICK_JS="""() => {
  if (window.__grantspaceCreateClickHandler) return [];
  window.__grantspaceCreateClickHandler = (event) => {
    const target = event.target;
    const control = target instanceof Element ? target.closest('#wizard-create-button') : null;
    const button = control && control.matches('button') ? control : control?.querySelector('button');
    if (!button) return;
    if (button.dataset.creationBusy === 'true' && !button.disabled) button.dataset.creationBusy = 'false';
    if (button.dataset.creationBusy === 'true') return;
    const editingExisting = button.textContent.includes('Save workflow');
    button.dataset.creationBusy = 'true';
    button.setAttribute('aria-busy', 'true');
    button.textContent = editingExisting ? 'Saving workflow changes…' : 'Creating shared grant…';
    const status = document.querySelector('#wizard-create-status .prose, #wizard-create-status');
    if (status) status.innerHTML = editingExisting ? '<h3>Saving workflow changes…</h3><p>Validating the composition and updating this grant without deleting historical artifacts.</p>' : '<h3>Creating and drafting your shared grant…</h3><p>The app will stay here while it saves the grant ask, derives the sponsor-specific outline, drafts every section in bounded model chunks, assembles the responses, and verifies the saved document.</p>';
    const progress = document.getElementById('wizard-create-progress');
    if (progress) {
      progress.className = 'active';
      const label = progress.querySelector('.label');
      if (label) label.textContent = editingExisting ? 'Validating the updated workflow configuration…' : 'Stage 1 of 4 · Saving the shared grant and authoritative source…';
    }
    setTimeout(() => { button.disabled = true; }, 0);
  };
  document.addEventListener('click', window.__grantspaceCreateClickHandler, false);
  window.__grantspaceCreateStatusObserver = new MutationObserver(() => {
    const status = document.getElementById('wizard-create-status');
    const progress = document.getElementById('wizard-create-progress');
    if (!status || !progress) return;
    const text = status.textContent || '';
    const label = progress.querySelector('.label');
    if (text.includes('Shared grant created')) {
      progress.className = 'complete';
      if (label) label.textContent = 'Stage 4 of 4 · Every model response is assembled and saved. Opening the editor…';
      setTimeout(() => {
        const overlay = document.getElementById('wizard-overlay');
        if (overlay) {
          overlay.style.setProperty('display', 'none', 'important');
          overlay.setAttribute('aria-hidden', 'true');
        }
        window.scrollTo({top: 0, behavior: 'smooth'});
      }, 250);
    } else if (text.includes('Grant creation failed')) {
      progress.className = 'failed';
      if (label) label.textContent = 'Creation failed · Read the error below, correct it, and retry.';
    }
  });
  window.__grantspaceCreateStatusObserver.observe(document.body, {subtree: true, childList: true, characterData: true});
  return [];
}"""

SESSION_STORAGE_AUTH_JS="""async () => {
  const storageKey = 'grantspace_session';
  const originalFetch = window.__grantspaceOriginalFetch || window.fetch.bind(window);
  window.__grantspaceOriginalFetch = originalFetch;
  try {
    const response = await originalFetch('/session-token', {credentials: 'same-origin', cache: 'no-store'});
    if (!response.ok) throw new Error('session bootstrap failed');
    const payload = await response.json();
    if (!payload.access_token) throw new Error('session token missing');
    window.sessionStorage.setItem(storageKey, payload.access_token);
  } catch (_) {
    window.sessionStorage.removeItem(storageKey);
    window.location.replace('/login');
    return [];
  }
  if (!window.__grantspaceSessionFetchInstalled) {
    window.fetch = (input, init = {}) => {
      const requestUrl = new URL(typeof input === 'string' ? input : input.url, window.location.href);
      if (requestUrl.origin === window.location.origin && requestUrl.pathname.startsWith('/app/')) {
        const headers = new Headers(init.headers || (typeof input !== 'string' ? input.headers : undefined));
        const token = window.sessionStorage.getItem(storageKey);
        if (token) headers.set('X-Grantspace-Session', token);
        init = {...init, headers, credentials: 'same-origin'};
      }
      return originalFetch(input, init);
    };
    window.__grantspaceSessionFetchInstalled = true;
  }
  return [];
}"""

WIZARD_HIDE_JS="""() => {
  const overlay = document.getElementById('wizard-overlay');
  if (overlay) {
    overlay.style.setProperty('display', 'none', 'important');
    overlay.setAttribute('aria-hidden', 'true');
  }
  window.scrollTo({top: 0, behavior: 'smooth'});
  return [];
}"""

WIZARD_HIDE_AFTER_CREATE_JS="""(status) => {
  if (!String(status || '').includes('Shared grant created and loaded')) return [];
  const overlay = document.getElementById('wizard-overlay');
  if (overlay) {
    overlay.style.setProperty('display', 'none', 'important');
    overlay.setAttribute('aria-hidden', 'true');
  }
  window.scrollTo({top: 0, behavior: 'smooth'});
  return [];
}"""

def apply_workflow_preset(preset_key):
    preset=WORKFLOW_PRESETS.get(preset_key)
    if not preset:raise gr.Error("Choose a valid workflow preset.")
    return list(preset["enabled_modules"]),list(preset["required_modules"])

def reconcile_module_gates(selected,required):
    selected=list(selected or [])
    allowed={key for _,key in GATE_CHOICES if key in selected}
    required=[key for key in (required or []) if key in allowed]
    choices=[choice for choice in GATE_CHOICES if choice[1] in allowed]
    return gr.update(choices=choices,value=required)

def validate_grant_details_and_continue(title,deadline):
    if not (title or "").strip():raise gr.Error("Working title is required.")
    if (deadline or "").strip():
        try:date.fromisoformat(deadline.strip())
        except ValueError:raise gr.Error("Sponsor deadline must use YYYY-MM-DD.")
    return wizard_page_updates(3)

def validate_grant_source_and_continue(source,source_url,source_text,edit_project=None):
    if not (edit_project or "").strip() and not source and not (source_url or "").strip() and not (source_text or "").strip():
        raise gr.Error("Upload, link, or paste the authoritative grant ask.")
    return wizard_page_updates(WIZARD_CORE_PAGE)

def optional_tools_continue(*module_modes):
    selected=selected_modules_from_modes(*module_modes)
    destination=WIZARD_REVIEW_PAGE if WORKFLOW_REGISTRY["review_module_key"] in selected else WIZARD_TEAM_PAGE
    return wizard_page_updates(destination)

def optional_tools_skip_all():
    return ["skip" for _ in WORKFLOW_REGISTRY["optional_modules"]]+wizard_page_updates(WIZARD_TEAM_PAGE)

def team_back_from_optional_tools(*module_modes):
    selected=selected_modules_from_modes(*module_modes)
    destination=WIZARD_REVIEW_PAGE if WORKFLOW_REGISTRY["review_module_key"] in selected else WIZARD_MODULE_PAGE
    return wizard_page_updates(destination)

def validate_routing_and_preview(title,sponsor,mechanism,deadline,review_mode,routing_mode,team_rows,*module_modes):
    selected=selected_modules_from_modes(*module_modes)
    required=[]
    review_required=False
    if WORKFLOW_REGISTRY["review_module_key"] not in selected:review_mode=None
    preview=workflow_preview_html(title,sponsor,mechanism,deadline,selected,required,review_mode,review_required,routing_mode,team_rows)
    return wizard_page_updates(WIZARD_PREVIEW_PAGE)+[preview,selected,required,review_required]

def validate_team_and_preview(title,sponsor,mechanism,deadline,selected,required,review_mode,review_required,routing_mode,team_rows):
    identity=api("GET","/api/me")
    if not identity.get("id"):raise gr.Error("Your authenticated account could not be resolved. Sign in again.")
    return wizard_to_preview(title,sponsor,mechanism,deadline,selected,required,review_mode,review_required,routing_mode,team_rows)

def authenticated_identity():
    identity=api("GET","/api/me")
    system=api("GET","/api/system/info")
    display=identity.get("display_name") or identity.get("username") or identity.get("email") or "signed-in user"
    account=identity.get("username") or identity.get("email") or "authenticated account"
    email_delivery=system.get("email_delivery") or {}
    if email_delivery.get("configured"):
        email_notice=f"Email delivery is configured through **{email_delivery.get('mode') or 'SMTP'}**."
    else:
        email_notice="**Email delivery is not configured.** Invitations will still be stored, but the app will show a one-time link instead of claiming an email was sent."
    return identity.get("id") or "",f"**Project owner:** {display} (`{account}`)  \nOwnership, edits, and approvals use your authenticated account automatically. No internal user ID is required.  \n{email_notice}"

def _normalized_team_rows(team_rows):
    rows=[]
    for row in _records(team_rows,["Email","Role"]):
        email=_cell(row.get("Email")).strip().lower()
        role=_cell(row.get("Role")).strip()
        if email:rows.append([email,role])
    return rows

def add_wizard_team_invitation(team_rows,email,role):
    email=(email or "").strip().lower()
    if not email or "@" not in email or email.startswith("@") or email.endswith("@"):
        raise gr.Error("Enter a valid teammate email address.")
    if role not in PROJECT_ROLE_LABELS:raise gr.Error("Choose a project role from the list.")
    rows=_normalized_team_rows(team_rows)
    if any(existing_email==email for existing_email,_ in rows):raise gr.Error("That email is already in the invitation list.")
    rows.append([email,role])
    remove_choices=[(f"{existing_email} · {PROJECT_ROLE_LABELS.get(existing_role,existing_role)}",existing_email) for existing_email,existing_role in rows]
    return rows,"",gr.update(choices=remove_choices,value=None),f"Added **{email}** as **{PROJECT_ROLE_LABELS[role]}**."

def remove_wizard_team_invitation(team_rows,email):
    target=(email or "").strip().lower()
    if not target:raise gr.Error("Choose a teammate to remove.")
    rows=[[existing_email,role] for existing_email,role in _normalized_team_rows(team_rows) if existing_email!=target]
    remove_choices=[(f"{existing_email} · {PROJECT_ROLE_LABELS.get(role,role)}",existing_email) for existing_email,role in rows]
    return rows,gr.update(choices=remove_choices,value=None),f"Removed **{target}** from the invitation list."

def account_rows():
    payload=api("GET","/api/admin/users")
    return [[item.get("id"),item.get("username"),item.get("email"),item.get("display_name"),item.get("system_role"),bool(item.get("must_change_password")),bool(item.get("active")),item.get("last_seen_at"),item.get("locked_until")] for item in payload.get("users",[])]

def create_account(username,email,display_name,temporary_password):
    result=api("POST","/api/admin/users",json={"username":username,"email":email,"display_name":display_name or None,"temporary_password":temporary_password})
    delivery="Temporary-password email sent." if result.get("email_sent") else f"Account created, but email delivery failed: {result.get('delivery_error') or 'unknown SMTP error'}"
    return account_rows(),f"**{result.get('user',{}).get('username')}** created. {delivery}",""

def set_account_status(user_id,active):
    if not (user_id or "").strip():raise gr.Error("Select or paste a user ID.")
    api("POST",f"/api/admin/users/{user_id.strip()}/{'enable' if active else 'disable'}",json={})
    return account_rows(),f"Account {'enabled' if active else 'disabled'}."

def send_account_reset(user_id):
    if not (user_id or "").strip():raise gr.Error("Select or paste a user ID.")
    result=api("POST",f"/api/admin/users/{user_id.strip()}/password-reset",json={})
    return f"Single-use reset link sent. It expires at {result.get('expires_at')}."

def project_workflow_ui(project):
    workflow=api("GET",f"/api/projects/{project}/workflow")
    status=api("GET",f"/api/projects/{project}/workflow/status")
    config=workflow.get("config") or workflow
    enabled=set(config.get("enabled_modules") or [])
    required=set(config.get("required_modules") or [])
    if config.get("review_required"):required.add(WORKFLOW_REGISTRY["review_module_key"])
    summary=(f"### Grant workflow\nDefinition **v{workflow.get('definition_version',config.get('definition_version'))}** · "
             f"configuration **v{workflow.get('config_version','—')}** · {len(enabled)} optional tools selected · routing **{config.get('model_routing_mode') or 'deployment default'}**")
    # Workflow configuration still controls model context and generated content,
    # but normal contributors work in one document-first surface. Specialist
    # records remain available through the API and portable audit export.
    hidden=gr.update(visible=False)
    return (summary,status,hidden,hidden,hidden,hidden,hidden,hidden,hidden)

def workflow_preview_html(title,sponsor,mechanism,deadline,selected,required,review_mode,review_required,routing_mode,team_rows):
    selected=set(selected or []);required=set(required or [])
    rows=[];position=0
    for step in WORKFLOW_REGISTRY["core_steps"]:
        before=[m for m in WORKFLOW_REGISTRY["optional_modules"] if m["key"] in selected and m["placement"]=="before_"+step["key"]]
        for module in before:
            position+=1;rows.append((position,module["title"],module["output"],module["key"] in required))
        position+=1;rows.append((position,step["title"],step["output"],True))
        placement={"solicitation":"after_solicitation","framework":"between_framework_and_aims","aims":"after_aims","literature":"after_literature","proposal":"after_draft"}.get(step["key"])
        for module in WORKFLOW_REGISTRY["optional_modules"]:
            if module["key"] in selected and module["placement"]==placement:
                position+=1;rows.append((position,module["title"],module["output"],module["key"] in required or (module["key"]==WORKFLOW_REGISTRY["review_module_key"] and review_required)))
    core_titles={step["title"] for step in WORKFLOW_REGISTRY["core_steps"]}
    body="".join(f'<div class="workflow-preview-row"><span class="n">{n}</span><span><b>{html.escape(name)}</b><br><small>{html.escape(output)}</small></span><small>{"Core outcome" if name in core_titles else "Optional tool"}</small></div>' for n,name,output,_gated in rows)
    cross=[WORKFLOW_MODULES[key]["title"] for key in selected if WORKFLOW_MODULES[key]["placement"] in {"cross_cutting","view"}]
    team_count=len(_records(team_rows,["Email","Role"]))
    meta=f'<p><b>{html.escape(title or "Untitled grant")}</b> · {html.escape(sponsor or "Sponsor not entered")} {html.escape(mechanism or "")} · {html.escape(deadline or "No deadline entered")}</p>'
    extras=f'<p><small>Cross-cutting capabilities: {html.escape(", ".join(sorted(cross)) or "None")} · Team invitations: {team_count} · Review: {html.escape(review_mode or "not selected")} · Routing: {html.escape(routing_mode or "deployment default")}</small></p>'
    return '<div class="workflow-preview">'+meta+body+extras+'</div>'

def build_workflow_config(preset_key,selected,required,grant_type,deadline,review_mode,review_required,routing_mode):
    selected=list(dict.fromkeys(selected or []));required=list(dict.fromkeys(required or []))
    review_key=WORKFLOW_REGISTRY["review_module_key"]
    return {
        "schema_version":WORKFLOW_REGISTRY["schema_version"],
        "definition_version":WORKFLOW_REGISTRY["definition_version"],
        "template":preset_key,
        "enabled_modules":selected,
        "required_modules":required,
        "review_mode":review_mode if review_key in selected else None,
        "review_required":bool(review_required and review_key in selected),
        "grant_type":grant_type or None,
        "target_deadline":deadline or None,
        "model_routing_mode":routing_mode or None,
        "local_model_provider":os.getenv("LOCAL_LLM_PROVIDER") or None,
        "local_model":os.getenv("LOCAL_LLM_API_MODEL") or os.getenv("LOCAL_LLM_MODEL") or None,
        "cloud_model":os.getenv("CLAUDE_MODEL") or None,
        "cloud_task_kinds":[x.strip() for x in os.getenv("CLAUDE_TASK_KINDS",os.getenv("CLOUD_TASK_KINDS","")).split(",") if x.strip()],
    }

def update_project_workflow_mode(project,core_only):
    project=_require_project(project)
    current=api("GET",f"/api/projects/{project}/workflow")
    proposed=json.loads(json.dumps(current.get("config") or current))
    previous_enabled=list(proposed.get("enabled_modules") or [])
    previous_required=list(proposed.get("required_modules") or [])
    review_was_required=bool(proposed.get("review_required"))
    proposed["required_modules"]=[]
    proposed["review_required"]=False
    if core_only:
        proposed["template"]="custom_configuration_v1"
        proposed["enabled_modules"]=[]
        proposed["review_mode"]=None
    changed=(previous_required or review_was_required or (core_only and previous_enabled))
    if not changed:
        action=("This grant already uses only the five core outcomes."
                if core_only else "All optional tools are already advisory and cannot block completion.")
        return action,*project_workflow_ui(project)
    impact=api("POST",f"/api/projects/{project}/workflow/impact",json={"workflow":proposed})
    if impact.get("destructive"):
        raise gr.Error("The server marked this workflow change destructive; no change was applied.")
    actor=(api("GET","/api/me").get("id") or "").strip()
    if not actor:raise gr.Error("Your authenticated account could not be resolved. Sign in again.")
    api("PATCH",f"/api/projects/{project}/workflow",json={"workflow":proposed,"expected_config_version":int(current.get("config_version")),"actor":actor})
    if core_only:
        action=(f"Workflow simplified to the five core outcomes. {len(previous_enabled)} optional tool(s) were hidden; "
                "their artifacts and history remain preserved and auditable.")
    else:
        action=(f"Removed {len(previous_required)+(1 if review_was_required else 0)} optional completion gate(s). "
                "Enabled tools remain available, but they can no longer block the proposal.")
    return action,*project_workflow_ui(project)

def make_optional_tools_advisory(project):return update_project_workflow_mode(project,False)
def use_five_core_outcomes(project):return update_project_workflow_mode(project,True)

def artifact_editor_load(project,artifact_type):
    if not project:raise gr.Error("Open a project first.")
    data=api("GET",f"/api/projects/{project}/workflow/artifacts/{artifact_type}")
    body=data.get("body")
    text=json.dumps(body,indent=2,ensure_ascii=False) if body is not None else ""
    version=data.get("version")
    status=(f"Version **{version}** · {'approved' if data.get('approved') else 'awaiting human approval'} · SHA-256 `{str(data.get('sha256') or '')[:16]}…`"
            if version else "No artifact exists yet.")
    return text,version,status,data

def parse_artifact_editor(text,artifact_type):
    try:body=json.loads(text or "")
    except json.JSONDecodeError as exc:raise gr.Error(f"{artifact_type} JSON is invalid at line {exc.lineno}, column {exc.colno}: {exc.msg}")
    if not isinstance(body,dict):raise gr.Error(f"{artifact_type} must be a JSON object.")
    return body

def artifact_editor_save(project,artifact_type,text,actor,current_version):
    if not (actor or "").strip():raise gr.Error("Authenticated actor ID is required.")
    body=parse_artifact_editor(text,artifact_type)
    data=api("POST",f"/api/projects/{project}/workflow/artifacts/{artifact_type}",json={"body":body,"source":"human_editor","author":actor.strip(),"expected_version":int(current_version) if current_version else None})
    return json.dumps(data["body"],indent=2,ensure_ascii=False),data["version"],f"Saved {artifact_type} v{data['version']}; approval is still required.",data

def mark_solicitation_facts_human_approved(body):
    for collection in ("eligibility","requirements","deadlines","budget_rules","attachments","review_criteria"):
        for item in body.get(collection) or []:
            if not item.get("sources"):
                label=item.get("label") or item.get("title") or "Unnamed item"
                raise gr.Error(f"The source text for “{label}” could not be located in the uploaded grant ask. Correct the wording or source before approval.")
            item["status"]="human_approved"
    return body

def artifact_editor_approve(project,artifact_type,text,actor,current_version):
    if not (actor or "").strip():raise gr.Error("Authenticated approver ID is required.")
    body=parse_artifact_editor(text,artifact_type)
    if artifact_type=="solicitation_profile":body=mark_solicitation_facts_human_approved(body)
    current=api("GET",f"/api/projects/{project}/workflow/artifacts/{artifact_type}")
    if body!=current.get("body"):
        current=api("POST",f"/api/projects/{project}/workflow/artifacts/{artifact_type}",json={"body":body,"source":"human_editor","author":actor.strip(),"expected_version":int(current_version) if current_version else None})
    version=current.get("version")
    if not version:raise gr.Error(f"No {artifact_type} version exists to approve.")
    approved=api("POST",f"/api/projects/{project}/workflow/artifacts/{artifact_type}/approve",json={"version":int(version),"approver":actor.strip()})
    workflow_status=api("GET",f"/api/projects/{project}/workflow/status")
    return json.dumps(approved["body"],indent=2,ensure_ascii=False),approved["version"],f"✓ Approved exact {artifact_type} v{approved['version']}.",approved,workflow_status

def artifact_editor_generate(project,artifact_type,actor):
    if not (actor or "").strip():raise gr.Error("Authenticated actor ID is required.")
    generated=api("POST",f"/api/projects/{project}/workflow/artifacts/{artifact_type}/generate",json={"actor":actor.strip(),"high_value":True},timeout=2400)
    data=generated["artifact"]
    return json.dumps(data["body"],indent=2,ensure_ascii=False),data["version"],f"Generated {artifact_type} v{data['version']} with `{generated['model']}` from approved {generated['upstream_artifact_type']} v{generated['upstream_version']}. Review and edit before approval.",data

SOLICITATION_FACT_HEADERS=["Rule or requirement"]
SOLICITATION_CRITERION_HEADERS=["Review criterion"]
FRAMEWORK_HEADERS=["Key","Title","Position","Requirement IDs","Criterion IDs","Narrative purpose","Key argument","Aim IDs","Evidence needs","Missing investigator inputs","Owner user ID","Approver user ID","Target words","Dependencies"]
CORE_AIM_HEADERS=["ID","Title","Statement","Rationale","Approach summary","Expected outcome","Impact","Innovation","Classification","Dependencies","Supporting evidence IDs"]
LITERATURE_QUERY_HEADERS=["ID","Query","Rationale","Aim IDs","Requirement IDs","Criterion IDs","Preferred domains"]
EVIDENCE_NEED_HEADERS=["Evidence need ID","Disposition","Evidence IDs","Rationale"]
REFERENCE_HEADERS=["Reference type","ID","Label","Details"]

def _editor_context(project):
    return api("GET",f"/api/projects/{project}/workflow/editor-context")

def _approved_artifact(context,artifact_type):
    return ((context or {}).get("approved_artifacts") or {}).get(artifact_type) or {}

def _reference_rows(context,kinds):
    rows=[]
    contract=(context or {}).get("contract") or {}
    solicitation=(_approved_artifact(context,"solicitation_profile").get("body") or {})
    framework=(_approved_artifact(context,"research_framework").get("body") or {})
    aims=(_approved_artifact(context,"aim_set").get("body") or {})
    if "requirements" in kinds:
        rows.extend(["requirement",item.get("id"),item.get("label"),"mandatory" if item.get("mandatory") else "optional"] for item in solicitation.get("requirements") or [])
    if "criteria" in kinds:
        rows.extend(["review criterion",item.get("id"),item.get("title"),"scored" if item.get("scored") else "narrative"] for item in solicitation.get("review_criteria") or [])
    if "members" in kinds:
        rows.extend(["project member",item.get("user_id"),item.get("name"),item.get("role")] for item in context.get("members") or [])
    if "framework_nodes" in kinds:
        rows.extend(["framework node",item.get("key"),item.get("title"),f"position {item.get('position')}"] for item in framework.get("nodes") or [])
    if "aims" in kinds:
        rows.extend(["aim",item.get("id"),item.get("title"),item.get("classification")] for item in aims.get("aims") or [])
    if "evidence" in kinds:
        rows.extend(["evidence",item.get("id"),item.get("claim"),item.get("status")] for item in context.get("evidence") or [])
    if "sources" in kinds:
        rows.extend(["research source",item.get("id"),item.get("title"),item.get("url")] for item in context.get("sources") or [])
    if "citations" in kinds:
        rows.extend(["citation",item.get("id"),item.get("title"),f"evidence {item.get('evidence_id')} · {'verified' if item.get('verified') else 'unverified'}"] for item in context.get("citations") or [])
    if "classifications" in kinds:
        rows.extend(["allowed classification",value,value,"server artifact contract"] for value in contract.get("assertion_classifications") or [])
    if "dispositions" in kinds:
        rows.extend(["allowed disposition",value,value,"server artifact contract"] for value in contract.get("evidence_need_dispositions") or [])
    return rows

def _split_values(value):
    if value is None:return []
    if isinstance(value,list):return [str(item).strip() for item in value if str(item).strip()]
    return [item.strip() for item in str(value).replace("\n",",").split(",") if item.strip()]

def _int_values(value):
    try:return [int(item) for item in _split_values(value)]
    except ValueError as exc:raise gr.Error(f"Expected comma-separated integer IDs: {exc}")

def _display_value(value):
    if isinstance(value,(dict,list)):return json.dumps(value,ensure_ascii=False)
    if value is None:return ""
    return str(value)

def _parse_value(value):
    text=str(value or "").strip()
    if not text:return ""
    try:return json.loads(text)
    except json.JSONDecodeError:return text

def _provenance_text(sources):
    return " | ".join(f"doc {source.get('document_id')} · {source.get('locator')} · {source.get('excerpt','')[:120]}" for source in (sources or []))

def _solicitation_fact_text(item):
    label=str(item.get("label") or "").strip()
    value=_display_value(item.get("value")).strip()
    return label if not value or value==label else f"{label}: {value}"

def _solicitation_criterion_text(item):
    title=str(item.get("title") or "").strip()
    description=str(item.get("description") or "").strip()
    return title if not description or description==title else f"{title} — {description}"

def _stable_human_id(prefix,text,used):
    base=f"{prefix}-{hashlib.sha256(text.encode('utf-8')).hexdigest()[:12].upper()}"
    candidate=base;counter=2
    while candidate in used:
        candidate=f"{base}-{counter}";counter+=1
    used.add(candidate)
    return candidate

def _artifact_status(data):
    version=data.get("version")
    if not version:return "No artifact exists yet."
    return f"Version **{version}** · {'approved' if data.get('approved') else 'awaiting human approval'} · SHA-256 `{str(data.get('sha256') or '')[:16]}…`"

def _approval_message(data,label):
    if data.get("approved"):return f"✓ Approved exact {label} v{data.get('version')}."
    progress=data.get("approval_progress") or {}
    return f"Approval recorded for {label} v{data.get('version')} ({progress.get('approvals',0)} of {progress.get('minimum_approvals','?')}); the artifact remains unapproved until the configured threshold is met."

def return_artifact_for_revision(project,artifact_type,version,rationale):
    if not version:raise gr.Error("Load the current artifact version before returning it for revision.")
    rationale=str(rationale or "").strip()
    if not rationale:raise gr.Error("Explain what must change before returning approved work for revision.")
    data=api("POST",f"/api/projects/{project}/workflow/artifacts/{artifact_type}/return-for-revision",json={"version":int(version),"rationale":rationale})
    status=api("GET",f"/api/projects/{project}/workflow/status")
    return f"↩ Returned exact {artifact_type} v{version} for revision. The prior version and approval history remain auditable.",data,status,""

def load_solicitation_form(project):
    data=api("GET",f"/api/projects/{project}/workflow/artifacts/solicitation_profile");context=_editor_context(project);data["editor_context"]=context
    body=data.get("body") or {}
    rows=[]
    for category in (context.get("contract") or {}).get("solicitation_fact_categories") or []:
        for item in body.get(category) or []:
            rows.append([_solicitation_fact_text(item)])
    criteria=[[_solicitation_criterion_text(item)] for item in body.get("review_criteria") or []]
    questions=[[item] for item in body.get("open_questions") or []]
    return body.get("working_title",""),body.get("sponsor",""),body.get("mechanism") or "",body.get("purpose",""),rows,criteria,questions,data.get("version"),_artifact_status(data),data

def solicitation_body(metadata,working_title,sponsor,mechanism,purpose,fact_rows,criterion_rows,question_rows):
    prior=(metadata or {}).get("body") or {}
    categories=((metadata or {}).get("editor_context") or {}).get("contract",{}).get("solicitation_fact_categories") or []
    if not categories:raise gr.Error("The server did not provide the solicitation editor contract. Reload the profile before editing.")
    body={"schema_version":1,"working_title":str(working_title or "").strip(),"sponsor":str(sponsor or "").strip(),"mechanism":str(mechanism or "").strip() or None,"purpose":str(purpose or "").strip(),"eligibility":[],"requirements":[],"review_criteria":[],"deadlines":[],"budget_rules":[],"attachments":[],"open_questions":[]}
    prior_fact_rows={}
    used_ids=set()
    for category in categories:
        for item in prior.get(category) or []:
            prior_fact_rows.setdefault(_solicitation_fact_text(item),[]).append((category,item))
            used_ids.add(str(item.get("id") or ""))
    for row in _records(fact_rows,SOLICITATION_FACT_HEADERS):
        text=str(row["Rule or requirement"] or "").strip()
        if not text:continue
        matches=prior_fact_rows.get(text) or []
        if matches:
            category,prior_item=matches.pop(0)
            body[category].append(prior_item)
        else:
            item_id=_stable_human_id("RULE",text,used_ids)
            body["requirements"].append({"id":item_id,"label":text,"value":text,"mandatory":False,"status":"human_corrected","sources":[]})
    prior_criteria={}
    for item in prior.get("review_criteria") or []:
        prior_criteria.setdefault(_solicitation_criterion_text(item),[]).append(item)
        used_ids.add(str(item.get("id") or ""))
    for row in _records(criterion_rows,SOLICITATION_CRITERION_HEADERS):
        text=str(row["Review criterion"] or "").strip()
        if not text:continue
        matches=prior_criteria.get(text) or []
        if matches:body["review_criteria"].append(matches.pop(0))
        else:
            item_id=_stable_human_id("CRITERION",text,used_ids)
            body["review_criteria"].append({"id":item_id,"title":text,"description":text,"scored":False,"scale":None,"status":"human_corrected","sources":[]})
    body["open_questions"]=[str(row["Question"]).strip() for row in _records(question_rows,["Question"]) if str(row["Question"] or "").strip()]
    return body

def save_solicitation_form(project,actor,version,metadata,*fields):
    body=solicitation_body(metadata,*fields)
    data=api("POST",f"/api/projects/{project}/workflow/artifacts/solicitation_profile",json={"body":body,"source":"human_structured_editor","author":actor,"expected_version":int(version) if version else None})
    return (*load_solicitation_form(project)[:-3],data["version"],f"Saved solicitation profile v{data['version']}; approval is still required.",data)

def approve_solicitation_form(project,actor,version,metadata,*fields):
    body=solicitation_body(metadata,*fields)
    current=api("GET",f"/api/projects/{project}/workflow/artifacts/solicitation_profile")
    collections=("eligibility","requirements","deadlines","budget_rules","attachments","review_criteria")
    needs_source_location=any(not item.get("sources") for collection in collections for item in body.get(collection) or [])
    if body!=current.get("body") or needs_source_location:
        current=api("POST",f"/api/projects/{project}/workflow/artifacts/solicitation_profile",json={"body":body,"source":"human_structured_editor","author":actor,"expected_version":int(version) if version else None})
    approved_body=json.loads(json.dumps(current.get("body") or {}))
    mark_solicitation_facts_human_approved(approved_body)
    if approved_body!=current.get("body"):
        current=api("POST",f"/api/projects/{project}/workflow/artifacts/solicitation_profile",json={"body":approved_body,"source":"human_approval_review","author":actor,"expected_version":int(current["version"])})
    approved=api("POST",f"/api/projects/{project}/workflow/artifacts/solicitation_profile/approve",json={"version":int(current["version"]),"approver":actor})
    loaded=load_solicitation_form(project)
    return (*loaded[:-2],_approval_message(approved,"solicitation profile"),approved,api("GET",f"/api/projects/{project}/workflow/status"))

def load_framework_form(project):
    data=api("GET",f"/api/projects/{project}/workflow/artifacts/research_framework");body=data.get("body") or {};context=_editor_context(project);data["editor_context"]=context
    rows=[[n.get("key"),n.get("title"),n.get("position"),", ".join(n.get("requirement_ids") or []),", ".join(n.get("review_criterion_ids") or []),n.get("narrative_purpose"),n.get("key_argument"),", ".join(n.get("linked_aim_ids") or []),", ".join(n.get("evidence_needs") or []),", ".join(n.get("missing_investigator_inputs") or []),n.get("owner_user_id"),n.get("approver_user_id"),n.get("target_words"),", ".join(n.get("dependencies") or [])] for n in body.get("nodes") or []]
    return body.get("overall_argument",""),rows,data.get("version"),_artifact_status(data),data,_reference_rows(context,{"requirements","criteria","members"})

def framework_body(metadata,argument,rows):
    prior=(metadata or {}).get("body") or {}
    approved=_approved_artifact((metadata or {}).get("editor_context"),"solicitation_profile")
    solicitation_version=prior.get("solicitation_profile_version") or approved.get("version")
    if not solicitation_version:raise gr.Error("Approve the solicitation profile before creating a research framework.")
    nodes=[]
    for row in _records(rows,FRAMEWORK_HEADERS):
        nodes.append({"key":str(row["Key"] or "").strip(),"title":str(row["Title"] or "").strip(),"position":int(row["Position"] or 0),"requirement_ids":_split_values(row["Requirement IDs"]),"review_criterion_ids":_split_values(row["Criterion IDs"]),"narrative_purpose":str(row["Narrative purpose"] or "").strip(),"key_argument":str(row["Key argument"] or "").strip(),"linked_aim_ids":_split_values(row["Aim IDs"]),"evidence_needs":_split_values(row["Evidence needs"]),"missing_investigator_inputs":_split_values(row["Missing investigator inputs"]),"owner_user_id":str(row["Owner user ID"] or "").strip(),"approver_user_id":str(row["Approver user ID"] or "").strip(),"target_words":int(row["Target words"] or 0),"dependencies":_split_values(row["Dependencies"])})
    return {"schema_version":1,"solicitation_profile_version":int(solicitation_version),"overall_argument":str(argument or "").strip(),"nodes":nodes}

def save_framework_form(project,actor,version,metadata,argument,rows,approve=False):
    body=framework_body(metadata,argument,rows);current=api("GET",f"/api/projects/{project}/workflow/artifacts/research_framework")
    if body!=current.get("body"):
        current=api("POST",f"/api/projects/{project}/workflow/artifacts/research_framework",json={"body":body,"source":"human_structured_editor","author":actor,"expected_version":int(version) if version else None})
    if approve:current=api("POST",f"/api/projects/{project}/workflow/artifacts/research_framework/approve",json={"version":int(current["version"]),"approver":actor})
    loaded=load_framework_form(project);status=(_approval_message(current,"research framework") if approve else f"Saved research framework v{current['version']}; approval is still required.")
    return (*loaded[:3],status,current,loaded[5])

def generate_framework_form(project,actor):
    artifact_editor_generate(project,"research_framework",actor)
    return load_framework_form(project)

def load_aims_form(project):
    data=api("GET",f"/api/projects/{project}/workflow/artifacts/aim_set");body=data.get("body") or {};context=_editor_context(project);data["editor_context"]=context
    rows=[[a.get("id"),a.get("title"),a.get("statement"),a.get("rationale"),a.get("approach_summary"),a.get("expected_outcome"),a.get("impact"),a.get("innovation"),a.get("classification"),", ".join(a.get("dependencies") or []),", ".join(str(x) for x in a.get("supporting_evidence_ids") or [])] for a in body.get("aims") or []]
    return body.get("overall_objective",""),body.get("central_hypothesis_or_thesis",""),rows,data.get("version"),_artifact_status(data),data,_reference_rows(context,{"framework_nodes","evidence","classifications"})

def aims_body(metadata,objective,hypothesis,rows):
    prior=(metadata or {}).get("body") or {};approved=_approved_artifact((metadata or {}).get("editor_context"),"research_framework");framework_version=prior.get("framework_version") or approved.get("version");aims=[]
    if not framework_version:raise gr.Error("Approve the research framework before creating aims.")
    for row in _records(rows,CORE_AIM_HEADERS):
        aims.append({"id":str(row["ID"] or "").strip(),"title":str(row["Title"] or "").strip(),"statement":str(row["Statement"] or "").strip(),"rationale":str(row["Rationale"] or "").strip(),"approach_summary":str(row["Approach summary"] or "").strip(),"expected_outcome":str(row["Expected outcome"] or "").strip(),"impact":str(row["Impact"] or "").strip(),"innovation":str(row["Innovation"] or "").strip(),"classification":str(row["Classification"] or "assumption").strip(),"dependencies":_split_values(row["Dependencies"]),"supporting_evidence_ids":_int_values(row["Supporting evidence IDs"])})
    return {"schema_version":1,"framework_version":int(framework_version),"overall_objective":str(objective or "").strip(),"central_hypothesis_or_thesis":str(hypothesis or "").strip(),"aims":aims}

def save_aims_form(project,actor,version,metadata,objective,hypothesis,rows,approve=False):
    body=aims_body(metadata,objective,hypothesis,rows);current=api("GET",f"/api/projects/{project}/workflow/artifacts/aim_set")
    if body!=current.get("body"):
        current=api("POST",f"/api/projects/{project}/workflow/artifacts/aim_set",json={"body":body,"source":"human_structured_editor","author":actor,"expected_version":int(version) if version else None})
    if approve:current=api("POST",f"/api/projects/{project}/workflow/artifacts/aim_set/approve",json={"version":int(current["version"]),"approver":actor})
    loaded=load_aims_form(project);status=(_approval_message(current,"aim set") if approve else f"Saved aim set v{current['version']}; approval is still required.")
    return (*loaded[:4],status,current,loaded[6])

def generate_aims_form(project,actor):
    artifact_editor_generate(project,"aim_set",actor)
    return load_aims_form(project)

def _literature_query_rows(body):
    return [[q.get("id"),q.get("query"),q.get("rationale"),", ".join(q.get("aim_ids") or []),", ".join(q.get("requirement_ids") or []),", ".join(q.get("criterion_ids") or []),", ".join(q.get("preferred_domains") or [])] for q in body.get("queries") or []]

def load_search_plan_form(project):
    data=api("GET",f"/api/projects/{project}/workflow/artifacts/literature_search_plan");body=data.get("body") or {};context=_editor_context(project);data["editor_context"]=context
    upstream=f"Solicitation v{body.get('solicitation_profile_version','—')} · Framework v{body.get('framework_version','—')} · Aims v{body.get('aim_set_version','—')}"
    return upstream,_literature_query_rows(body),data.get("version"),_artifact_status(data),data,_reference_rows(context,{"aims","requirements","criteria"})

def search_plan_body(metadata,query_rows):
    prior=(metadata or {}).get("body") or {};context=(metadata or {}).get("editor_context") or {}
    solicitation=prior.get("solicitation_profile_version") or _approved_artifact(context,"solicitation_profile").get("version")
    framework=prior.get("framework_version") or _approved_artifact(context,"research_framework").get("version")
    aims=prior.get("aim_set_version") or _approved_artifact(context,"aim_set").get("version")
    if not all((solicitation,framework,aims)):raise gr.Error("Approve the solicitation profile, research framework, and aims before saving a literature search plan.")
    queries=[{"id":str(r["ID"] or "").strip(),"query":str(r["Query"] or "").strip(),"rationale":str(r["Rationale"] or "").strip(),"aim_ids":_split_values(r["Aim IDs"]),"requirement_ids":_split_values(r["Requirement IDs"]),"criterion_ids":_split_values(r["Criterion IDs"]),"preferred_domains":_split_values(r["Preferred domains"])} for r in _records(query_rows,LITERATURE_QUERY_HEADERS)]
    return {"schema_version":1,"solicitation_profile_version":int(solicitation),"framework_version":int(framework),"aim_set_version":int(aims),"queries":queries}

def save_search_plan_form(project,actor,version,metadata,query_rows,approve=False):
    body=search_plan_body(metadata,query_rows);current=api("GET",f"/api/projects/{project}/workflow/artifacts/literature_search_plan")
    if body!=current.get("body"):
        current=api("POST",f"/api/projects/{project}/workflow/artifacts/literature_search_plan",json={"body":body,"source":"human_structured_editor","author":actor,"expected_version":int(version) if version else None})
    if approve:current=api("POST",f"/api/projects/{project}/workflow/artifacts/literature_search_plan/approve",json={"version":int(current["version"]),"approver":actor})
    loaded=load_search_plan_form(project);status=(_approval_message(current,"literature search plan") if approve else f"Saved literature search plan v{current['version']}; approval is still required.")
    return (*loaded[:3],status,current,loaded[5])

def generate_search_plan_form(project,max_queries):
    result=api("POST",f"/api/projects/{project}/research/plan",json={"max_queries":int(max_queries)},timeout=1200)
    loaded=list(load_search_plan_form(project));warnings=result.get("planner_warnings") or []
    loaded[3]=f"Generated search plan v{loaded[2]} with `{result.get('provider')}:{result.get('model')}`. Review and approve it before any external search runs."+(f" {len(warnings)} invalid or duplicate model suggestion(s) were excluded." if warnings else "")
    return tuple(loaded)

def load_literature_form(project):
    data=api("GET",f"/api/projects/{project}/workflow/artifacts/literature_manifest");body=data.get("body") or {};context=_editor_context(project);data["editor_context"]=context
    queries=_literature_query_rows(body)
    needs=[[n.get("evidence_need_id"),n.get("disposition"),", ".join(str(x) for x in n.get("evidence_ids") or []),n.get("rationale")] for n in body.get("evidence_needs") or []]
    contradictions=[[value] for value in body.get("contradictions") or []]
    manifest=f"Run `{body.get('run_id','—')}` · {body.get('search_provider','—')} · {body.get('started_at','—')} → {body.get('completed_at','—')} · source IDs {body.get('source_ids') or []} · citation IDs {body.get('citation_ids') or []}"
    return manifest,queries,needs,contradictions,data.get("version"),_artifact_status(data),data,_reference_rows(context,{"aims","requirements","criteria","evidence","sources","citations","dispositions"})

def literature_body(metadata,query_rows,need_rows,contradiction_rows):
    body=dict((metadata or {}).get("body") or {})
    body["queries"]=[{"id":str(r["ID"] or "").strip(),"query":str(r["Query"] or "").strip(),"rationale":str(r["Rationale"] or "").strip(),"aim_ids":_split_values(r["Aim IDs"]),"requirement_ids":_split_values(r["Requirement IDs"]),"criterion_ids":_split_values(r["Criterion IDs"]),"preferred_domains":_split_values(r["Preferred domains"])} for r in _records(query_rows,LITERATURE_QUERY_HEADERS)]
    body["evidence_needs"]=[{"evidence_need_id":str(r["Evidence need ID"] or "").strip(),"disposition":str(r["Disposition"] or "unresolved_risk").strip(),"evidence_ids":_int_values(r["Evidence IDs"]),"rationale":str(r["Rationale"] or "").strip()} for r in _records(need_rows,EVIDENCE_NEED_HEADERS)]
    body["contradictions"]=[str(r["Contradiction"] or "").strip() for r in _records(contradiction_rows,["Contradiction"]) if str(r["Contradiction"] or "").strip()]
    return body

def save_literature_form(project,actor,version,metadata,query_rows,need_rows,contradiction_rows,approve=False):
    body=literature_body(metadata,query_rows,need_rows,contradiction_rows);current=api("GET",f"/api/projects/{project}/workflow/artifacts/literature_manifest")
    if body!=current.get("body"):
        current=api("POST",f"/api/projects/{project}/workflow/artifacts/literature_manifest",json={"body":body,"source":"human_structured_editor","author":actor,"expected_version":int(version) if version else None})
    if approve:current=api("POST",f"/api/projects/{project}/workflow/artifacts/literature_manifest/approve",json={"version":int(current["version"]),"approver":actor})
    loaded=load_literature_form(project);status=(_approval_message(current,"literature manifest") if approve else f"Saved literature manifest v{current['version']}; approval is still required.")
    return (*loaded[:5],status,current,loaded[7])

def collaboration_markdown(data):
    messages=data.get("messages") or []
    if not messages:return '<div class="team-chat">No messages yet. Start the project conversation.</div>'
    blocks=[]
    for m in messages:
        blocks.append(f'<div class="chat-message"><div class="chat-meta"><b>{html.escape(str(m.get("author") or ""))}</b> · {html.escape(str(m.get("created_at") or ""))}</div>{html.escape(str(m.get("body") or "")).replace(chr(10),"<br>")}</div>')
    return '<div class="team-chat">'+"".join(blocks)+"</div>"

def _require_project(project):
    if not (project or "").strip():raise gr.Error("Open a project first.")
    return project.strip()

def _choice_update(items,label_key="title",id_key="id",current=None):
    choices=[(str(item.get(label_key) or item.get(id_key)),str(item.get(id_key))) for item in items if item.get(id_key) is not None]
    values={value for _,value in choices};selected=str(current) if current is not None and str(current) in values else (choices[0][1] if choices else None)
    return gr.update(choices=choices,value=selected)

def _member_choices(members):
    return [(f"{item.get('name') or item.get('email') or item.get('user_id')} · {item.get('role')}",str(item.get("user_id"))) for item in members if item.get("user_id")]

def team_workspace_rows(data):
    members=data.get("members") or [];invites=data.get("invites") or [];tasks=data.get("tasks") or [];notifications=data.get("notifications") or [];routing=(data.get("approval_routing") or {}).get("routes") or [];health=(data.get("health") or {}).get("issues") or []
    member_rows=[[x.get("user_id"),x.get("name"),x.get("email"),x.get("role"),bool(x.get("present")),x.get("joined_at"),x.get("last_seen_at")] for x in members]
    activity_rows=[[x.get("created_at"),x.get("actor"),x.get("kind"),x.get("detail")] for x in data.get("activity") or []]
    invite_rows=[]
    for x in invites:
        status="accepted" if x.get("accepted_at") else ("revoked" if x.get("revoked_at") else "pending")
        invite_rows.append([x.get("id"),x.get("email"),x.get("role"),status,x.get("expires_at"),x.get("created_at")])
    task_rows=[[x.get("id"),x.get("priority"),x.get("status"),x.get("title"),x.get("owner_user_id"),x.get("due_at"),x.get("source"),", ".join(x.get("dependencies") or []),x.get("updated_at")] for x in tasks]
    notification_rows=[[x.get("id"),not bool(x.get("read_at")),x.get("kind"),_display_value(x.get("payload")),x.get("created_at"),x.get("read_at")] for x in notifications]
    routing_rows=[[x.get("title"),x.get("current_version"),x.get("owner_user_id"),", ".join(x.get("approver_user_ids") or []),x.get("approvals"),x.get("minimum_approvals"),bool(x.get("threshold_met")),bool(x.get("approved"))] for x in routing]
    health_rows=[[x.get("severity"),x.get("kind"),x.get("title"),x.get("detail"),x.get("owner_user_id"),x.get("step_key"),x.get("due_at"),x.get("remediation")] for x in health]
    return member_rows,activity_rows,invite_rows,task_rows,notification_rows,routing_rows,health_rows

def project_health_summary(data):
    health=data.get("health") or {};summary=health.get("summary") or {};state=str(health.get("state") or "unknown").replace("_"," ")
    return f"### Project health: **{state}**\nCritical **{summary.get('critical',0)}** · High **{summary.get('high',0)}** · Medium **{summary.get('medium',0)}** · Total active findings **{summary.get('total',0)}**"

def load_team_workspace(project):
    project=_require_project(project);data=api("GET",f"/api/projects/{project}/collaboration/workspace")
    members=data.get("members") or [];tasks=data.get("tasks") or [];invites=[item for item in data.get("invites") or [] if not item.get("accepted_at") and not item.get("revoked_at")];notifications=[item for item in data.get("notifications") or [] if not item.get("read_at")]
    rows=team_workspace_rows(data);member_choices=_member_choices(members);task_choices=[(f"{item.get('title')} · {str(item.get('id'))[:8]}",str(item.get("id"))) for item in tasks]
    permissions=data.get("permissions") or {};status=f"Signed in with project role **{permissions.get('role','unknown')}** · {len(members)} member(s) · {len(tasks)} task(s) · {len(notifications)} unread notification(s)."
    return (*rows,project_health_summary(data),gr.update(choices=member_choices,value=(member_choices[0][1] if member_choices else None)),gr.update(choices=member_choices,value=[]),gr.update(choices=task_choices,value=(task_choices[0][1] if task_choices else None)),gr.update(choices=task_choices,value=[]),_choice_update(invites,label_key="email"),_choice_update(notifications,label_key="kind"),permissions,status,gr.update(interactive=bool(permissions.get("can_manage_members"))),gr.update(interactive=bool(permissions.get("can_manage_members"))),gr.update(interactive=bool(permissions.get("can_post"))),gr.update(interactive=bool(permissions.get("can_create_tasks"))))

def poll_team_workspace(project):
    if not (project or "").strip():return tuple(gr.skip() for _ in range(9))
    data=api("GET",f"/api/projects/{project}/collaboration/workspace");rows=team_workspace_rows(data);members=data.get("members") or [];tasks=data.get("tasks") or [];notifications=[item for item in data.get("notifications") or [] if not item.get("read_at")];permissions=data.get("permissions") or {}
    status=f"Live sync · signed in as **{permissions.get('role','unknown')}** · {len(members)} member(s) · {len(tasks)} task(s) · {len(notifications)} unread notification(s)."
    return (*rows,project_health_summary(data),status)

def poll_team_channel(project,kind,subject_key):
    if not (project or "").strip():return tuple(gr.skip() for _ in range(3))
    if str(kind or "general")=="section" and not str(subject_key or "").strip():return tuple(gr.skip() for _ in range(3))
    loaded=load_team_channel(project,kind,subject_key)
    return loaded[0],loaded[1],loaded[3]

def poll_shared_artifact_versions(project,local_search_plan_version,local_manifest_version):
    if not (project or "").strip():return gr.skip()
    notices=[]
    for label,artifact_type,local_version in (("search plan","literature_search_plan",local_search_plan_version),("literature manifest","literature_manifest",local_manifest_version)):
        remote=api("GET",f"/api/projects/{project}/workflow/artifacts/{artifact_type}");remote_version=remote.get("version")
        if local_version and remote_version and int(local_version)!=int(remote_version):notices.append(f"A teammate published {label} v{remote_version} while this browser has v{int(local_version)} loaded. Reload before editing; stale saves are rejected server-side.")
    return "\n\n".join(f"> ⚠ {notice}" for notice in notices) if notices else "Live collaboration sync is current."

def poll_collaboration_snapshot(project,kind,subject_key,local_search_plan_version,local_manifest_version):
    """Fetch one collaboration snapshot and release all per-run state on return.

    The Gradio timer owns scheduling. This callback never sleeps, loops, retains a
    background task, or schedules itself, so each 12-second refresh is an isolated
    request whose locals become collectible immediately after its response.
    """
    if not (project or "").strip():return tuple(gr.skip() for _ in range(13))
    workspace=poll_team_workspace(project)
    channel=poll_team_channel(project,kind,subject_key)
    artifact_sync=poll_shared_artifact_versions(project,local_search_plan_version,local_manifest_version)
    return (*workspace,*channel,artifact_sync)

def channel_messages_html(messages):
    if not messages:return '<div class="team-chat">No messages in this channel yet.</div>'
    blocks=[]
    for item in messages:
        indent=" style=\"margin-left:32px\"" if item.get("parent_message_id") else ""
        meta=f"#{item.get('id')} · {item.get('author')} · {item.get('created_at')}"
        body=html.escape(str(item.get("body") or "")).replace(chr(10),"<br>")
        blocks.append(f'<div class="chat-message"{indent}><div class="chat-meta"><b>{html.escape(meta)}</b></div>{body}</div>')
    return '<div class="team-chat">'+"".join(blocks)+"</div>"

def load_team_channel(project,kind,subject_key):
    project=_require_project(project);kind=str(kind or "general")
    subject=str(subject_key or "").strip() or None
    if kind=="section" and not subject:raise gr.Error("Choose or enter a section key for a section channel.")
    messages=api("GET",f"/api/projects/{project}/channels/{kind}",params={"subject_key":subject} if subject else {})
    rows=[[x.get("id"),x.get("parent_message_id"),x.get("author"),x.get("author_user_id"),x.get("created_at"),x.get("body")] for x in messages]
    choices=[(f"#{x.get('id')} · {x.get('author')} · {str(x.get('body') or '')[:50]}",x.get("id")) for x in messages]
    return channel_messages_html(messages),rows,gr.update(choices=choices,value=None),f"Loaded {len(messages)} message(s) from **{kind}{': '+subject if subject else ''}**."

def post_team_channel_message(project,kind,subject_key,message,parent_id,mentioned_user_ids):
    project=_require_project(project);body=str(message or "").strip()
    if not body:raise gr.Error("Write a message before posting.")
    subject=str(subject_key or "").strip() or None
    api("POST",f"/api/projects/{project}/channels/{kind}",params={"subject_key":subject} if subject else {},json={"body":body,"parent_message_id":int(parent_id) if parent_id else None,"mentioned_user_ids":[str(value) for value in (mentioned_user_ids or [])]})
    loaded=load_team_channel(project,kind,subject_key)
    return *loaded,""

def invite_delivery_summary(result):
    link=f"{os.getenv('APP_PUBLIC_URL','http://127.0.0.1:7860').rstrip('/')}/invite?token={quote(str(result.get('token') or ''))}"
    delivery="Email accepted by the configured SMTP server." if result.get("email_sent") else f"Email was not accepted: {result.get('delivery_error')}. Send the one-time link through an approved secure channel."
    account=("The recipient has an active account and can accept this invitation."
             if result.get("account_exists") else
             "The recipient does not have an active account yet. A system administrator must create one with this exact email before the invitation can be accepted.")
    public_url=os.getenv("APP_PUBLIC_URL","http://127.0.0.1:7860").rstrip("/")
    reachability=(" The configured link uses a loopback address and is only usable from this computer."
                  if public_url.startswith(("http://127.0.0.1","http://localhost")) else "")
    return f"Invite created for **{result.get('email')}**. {delivery} {account}{reachability}\n\nOne-time link: `{link}`"

def _validated_email_address(value):
    display_name,address=parseaddr(str(value or "").strip())
    if "\r" in address or "\n" in address:raise ValueError("Email address contains invalid newline characters")
    normalized=Address(addr_spec=address).addr_spec
    if not normalized:raise ValueError("Email address is empty")
    return display_name,normalized

def _python_smtp_send(recipient,subject,body):
    host=os.getenv("SMTP_HOST","").strip()
    if not host:raise RuntimeError("SMTP_HOST is not configured")
    security=os.getenv("SMTP_SECURITY","starttls").strip().lower()
    if security not in {"none","starttls","tls"}:raise RuntimeError("SMTP_SECURITY must be none, starttls, or tls")
    default_port=465 if security=="tls" else (587 if security=="starttls" else 25)
    port=int(os.getenv("SMTP_PORT",str(default_port)) or default_port)
    timeout=float(os.getenv("SMTP_TIMEOUT_SECONDS","30") or 30)
    sender_display,sender=_validated_email_address(os.getenv("SMTP_FROM",""))
    _,recipient=_validated_email_address(recipient)
    if "\r" in subject or "\n" in subject:raise ValueError("Email subject contains invalid newline characters")
    message=EmailMessage();message["From"]=Address(display_name=sender_display,addr_spec=sender);message["To"]=recipient;message["Subject"]=subject;message.set_content(body)
    context=ssl.create_default_context();client=None
    try:
        client=smtplib.SMTP_SSL(host,port,timeout=timeout,context=context) if security=="tls" else smtplib.SMTP(host,port,timeout=timeout)
        client.ehlo()
        if security=="starttls":client.starttls(context=context);client.ehlo()
        username=os.getenv("SMTP_USERNAME","");password=os.getenv("SMTP_PASSWORD","")
        if bool(username)!=bool(password):raise RuntimeError("SMTP_USERNAME and SMTP_PASSWORD must be configured together")
        if username:client.login(username,password)
        refused=client.send_message(message,from_addr=sender,to_addrs=[recipient])
        if refused:raise RuntimeError(f"SMTP server refused the recipient: {refused}")
    finally:
        if client is not None:
            try:client.quit()
            except Exception:client.close()

def _ensure_invite_email(result,project_title,role):
    if result.get("email_sent") or not os.getenv("SMTP_HOST","").strip():return result
    public_url=os.getenv("APP_PUBLIC_URL","http://127.0.0.1:7860").rstrip("/")
    token=str(result.get("token") or "")
    body=(f'You were invited to the Grantspace project "{project_title}" with the role {PROJECT_ROLE_LABELS.get(role,role)}.\n\n'
          f'Sign in with the Grantspace account matching this email address, then accept the single-use invitation:\n{public_url}/invite?token={token}\n\n'
          'If you do not yet have an account, contact the Grantspace administrator who invited you. Do not forward this link.')
    updated=dict(result)
    try:
        _python_smtp_send(result.get("email"),"You were invited to a Grantspace project",body)
        updated["email_sent"]=True;updated["delivery_error"]=None;updated["delivery_provider"]="python_smtplib"
    except Exception as error:
        prior=str(result.get("delivery_error") or "").strip()
        updated["email_sent"]=False;updated["delivery_error"]="; ".join(value for value in [prior,f"Python SMTP fallback: {error}"] if value)
    return updated

def create_project_invite_ui(project,email,role,expires_days):
    project=_require_project(project);result=api("POST",f"/api/projects/{project}/invites",json={"email":str(email or "").strip(),"role":role,"expires_in_days":int(expires_days)})
    metadata=api("GET",f"/api/projects/{project}");result=_ensure_invite_email(result,metadata.get("title") or "Grantspace project",role)
    return invite_delivery_summary(result),*load_team_workspace(project)

def add_existing_project_member_ui(project,user_id,role):
    project=_require_project(project);user_id=str(user_id or "").strip()
    if not user_id:raise gr.Error("Enter the existing account's stable user ID.")
    api("POST",f"/api/projects/{project}/collaboration/join",json={"user_id":user_id,"role":role})
    return f"User `{user_id}` added with role **{role}**.",*load_team_workspace(project),""

def revoke_project_invite_ui(project,invite_id):
    project=_require_project(project)
    if not invite_id:raise gr.Error("Choose an active invite.")
    api("POST",f"/api/projects/{project}/invites/{invite_id}/revoke",json={})
    return "Invite revoked.",*load_team_workspace(project)

def create_project_task_ui(project,title,description,owner,priority,due_at,dependencies):
    project=_require_project(project)
    api("POST",f"/api/projects/{project}/tasks",json={"title":str(title or "").strip(),"description":str(description or "").strip(),"owner_user_id":owner,"source":"human","priority":priority,"due_at":str(due_at or "").strip() or None,"dependencies":[str(value) for value in (dependencies or [])]})
    return "Task created and the owner was notified.",*load_team_workspace(project),"",""

def update_project_task_ui(project,task_id,status):
    project=_require_project(project)
    if not task_id:raise gr.Error("Choose a task.")
    api("POST",f"/api/projects/{project}/tasks/{task_id}/status",json={"status":status})
    return f"Task moved to **{status}**.",*load_team_workspace(project)

def load_artifact_comments(project,target_type,target_key,version_id):
    project=_require_project(project);key=str(target_key or "").strip()
    if not key or not version_id:raise gr.Error("Artifact key and exact version ID are required.")
    comments=api("GET",f"/api/projects/{project}/comments/{target_type}/{key}",params={"version_id":int(version_id)})
    rows=[[x.get("id"),x.get("parent_comment_id"),x.get("author"),x.get("created_at"),x.get("start_offset"),x.get("end_offset"),x.get("quoted_text"),x.get("body"),bool(x.get("resolved_at")),x.get("resolved_at")] for x in comments]
    open_comments=[x for x in comments if not x.get("resolved_at")]
    choices=[(f"#{x.get('id')} · {x.get('author')} · {str(x.get('body') or '')[:45]}",x.get("id")) for x in open_comments]
    return rows,gr.update(choices=choices,value=None),gr.update(choices=choices,value=None),f"Loaded {len(comments)} comment(s) anchored to exact version **{int(version_id)}**."

def current_comment_target_version(project,target_type,target_key):
    project=_require_project(project);key=str(target_key or "").strip()
    if not key:raise gr.Error("Enter a section key or workflow artifact type.")
    if target_type=="section":data=api("GET",f"/api/projects/{project}/sections/{key}");version=(data.get("latest") or {}).get("version")
    else:data=api("GET",f"/api/projects/{project}/workflow/artifacts/{key}");version=data.get("id")
    if not version:raise gr.Error("The selected artifact has no stored version yet.")
    return int(version),f"Selected immutable stored version ID **{int(version)}** for commenting."

def post_artifact_comment(project,target_type,target_key,version_id,start,end,quoted_text,body,parent_id,mentions):
    project=_require_project(project);key=str(target_key or "").strip()
    api("POST",f"/api/projects/{project}/comments/{target_type}/{key}",json={"version_id":int(version_id),"start_offset":int(start) if start is not None else None,"end_offset":int(end) if end is not None else None,"quoted_text":str(quoted_text or "").strip() or None,"body":str(body or "").strip(),"parent_comment_id":int(parent_id) if parent_id else None,"mentioned_user_ids":[str(value) for value in (mentions or [])]})
    loaded=load_artifact_comments(project,target_type,key,version_id);return *loaded,""

def resolve_artifact_comment(project,target_type,target_key,version_id,comment_id):
    project=_require_project(project)
    if not comment_id:raise gr.Error("Choose an open comment.")
    api("POST",f"/api/projects/{project}/comments/{int(comment_id)}/resolve",json={})
    loaded=list(load_artifact_comments(project,target_type,target_key,version_id));loaded[3]="Comment resolved without deleting its history.";return tuple(loaded)

def mark_notification_read_ui(project,notification_id):
    if not notification_id:raise gr.Error("Choose an unread notification.")
    api("POST",f"/api/notifications/{int(notification_id)}/read",json={})
    return "Notification marked read.",*load_team_workspace(project)

def version_history(project,key):
    if not project or not key:
        empty=gr.update(choices=[],value=None)
        return [],empty,empty,empty
    items=api("GET",f"/api/projects/{project}/sections/{key}/versions")
    rows=[[x.get("version"),x.get("created_at"),x.get("editor"),x.get("source"),x.get("approved"),x.get("base_version_id"),x.get("restored_from_version_id"),x.get("characters"),x.get("preview")] for x in items]
    choices=[(f"v{x.get('version')} · {x.get('editor') or x.get('source')} · {x.get('created_at')}",x.get("version")) for x in items]
    latest=choices[0][1] if choices else None;older=choices[1][1] if len(choices)>1 else latest
    return rows,gr.update(choices=choices,value=older),gr.update(choices=choices,value=latest),gr.update(choices=choices,value=older)

def compare_versions(project,key,from_version,to_version):
    if not from_version or not to_version:raise gr.Error("Choose two stored versions.")
    if int(from_version)==int(to_version):raise gr.Error("Choose two different versions.")
    data=api("GET",f"/api/projects/{project}/sections/{key}/compare",params={"from_version":int(from_version),"to_version":int(to_version)})
    before=data["from"];after=data["to"]
    table=difflib.HtmlDiff(wrapcolumn=110).make_table((before.get("body") or "").splitlines(),(after.get("body") or "").splitlines(),fromdesc=f"v{before['version']}",todesc=f"v{after['version']}",context=True,numlines=5)
    return f'<div class="version-diff">{table}</div>',f"Comparing immutable version **{before['version']}** with **{after['version']}**. Green is added; red is removed."

def restore_version(project,project_title,section,key,current_version,restore_id):
    if not restore_id:raise gr.Error("Choose a version to restore.")
    if int(restore_id)==int(current_version):raise gr.Error("That is already the latest version.")
    api("POST",f"/api/projects/{project}/sections/{key}/restore",json={"version_id":int(restore_id),"base_version_id":int(current_version)})
    loaded=load_section(project,project_title,section);history=version_history(project,key)
    loaded=list(loaded);loaded[3]=f"Restored v{restore_id} as a new auditable version. Approval is still required."
    return *loaded,*history

def assert_section_identity(section,key):
    expected=slug(section)
    if not expected or key!=expected:
        raise gr.Error("Section state changed while an edit/approval was in progress. Reload the selected section before saving or approving.")


def refresh_projects():
    items=api("GET","/api/projects");choices=[(f"{p['title']} · {p['stage']} · {p['id'][:8]}",p["id"]) for p in items]
    return gr.update(choices=choices,value=(choices[0][1] if choices else None))

def project_catalog(include_archived=False):
    items=api("GET","/api/projects",params={"include_archived":str(bool(include_archived)).lower()})
    choices=[];rows=[]
    for project in items:
        archived=bool(project.get("archived_at"));state="Archived" if archived else "Active"
        label=f"{project['title']} · {project.get('stage') or 'not started'} · {state} · {project['id'][:8]}"
        choices.append((label,project["id"]))
        rows.append([project.get("id"),project.get("title"),project.get("sponsor") or "",project.get("mechanism") or "",project.get("stage"),project.get("role") or "system administrator",state,project.get("updated_at"),project.get("created_at")])
    active=sum(1 for item in items if not item.get("archived_at"));archived=len(items)-active
    summary=f"**{active} active grant(s)**"+(f" · {archived} archived" if include_archived else "")+" · saved on the shared server and available after sign-out/sign-in."
    return gr.update(choices=choices,value=(choices[0][1] if choices else None)),rows,summary

def update_project_manager(project,title,archive_action,include_archived):
    if not project:raise gr.Error("Choose a saved grant first.")
    payload={}
    if (title or "").strip():payload["title"]=title.strip()
    if archive_action=="Archive":payload["archived"]=True
    elif archive_action=="Restore":payload["archived"]=False
    if not payload:raise gr.Error("Enter a new title or choose Archive/Restore.")
    updated=api("PATCH",f"/api/projects/{project}",json=payload)
    selector,rows,summary=project_catalog(include_archived)
    action="archived" if archive_action=="Archive" else ("restored" if archive_action=="Restore" else "updated")
    return selector,rows,summary,f"Grant **{updated.get('title')}** was {action}. All documents, versions, approvals, research, and collaboration history remain stored.","",None

def create_project(title,sponsor,mechanism,source,source_url,source_text,supporting,brand,workflow=None,actor=None,analyze=True,progress=None):
    if progress:progress(0.03,desc="Stage 1 of 4 · Validating the project title and authoritative grant source")
    if not title.strip():raise gr.Error("Working title is required.")
    if not source and not (source_url or "").strip() and not (source_text or "").strip():raise gr.Error("Upload, link, or paste a funding opportunity.")
    if progress:progress(0.12,desc="Stage 1 of 4 · Saving the shared grant record and workflow configuration")
    # Composed grants receive a solicitation-derived model plan after their
    # authoritative source is stored. Static legacy sections are retained only
    # for the older direct-create surface.
    payload={"title":title.strip(),"sponsor":sponsor or None,"mechanism":mechanism or None,"sections":([] if workflow is not None else DEFAULT_SECTIONS),"actor":actor or None}
    if workflow is not None:payload["workflow"]=workflow
    d=api("POST","/api/projects",json=payload)
    pid=d["id"];count=0
    if progress:progress(0.28,desc="Stage 2 of 4 · Extracting and storing the authoritative grant source")
    if source:
        p=Path(file_path(source));count+=int(push_doc(pid,p.name,"funding_opportunity",extract_file(source)))
    if (source_url or "").strip():
        f=api("POST",f"/api/projects/{pid}/documents/fetch-url",json={"url":source_url.strip(),"kind":"funding_url"},timeout=120);count+=int(f.get("added",False))
    if (source_text or "").strip():
        # Preserve the pasted characters. Trimming here would make later byte
        # offsets refer to a transformed string rather than the user's source.
        count+=int(push_doc(pid,"Pasted funding opportunity","funding_paste",source_text))
    for f in supporting or []:
        p=Path(file_path(f));count+=int(push_doc(pid,p.name,"supporting",extract_file(f)))
    if progress:progress(0.62,desc="Stage 3 of 4 · Saving supporting materials and document design settings")
    assets=copy_brand_assets(pid,brand);profile=build_design_profile(pid,sponsor,assets)
    api("POST",f"/api/projects/{pid}/design-profile",json={"profile":profile})
    sections=api("GET",f"/api/projects/{pid}/sections");section_choices=[x["title"] for x in sections]
    if not analyze:
        if progress:progress(0.18,desc="Stage 1 of 4 · Authoritative grant source saved; preparing the model-derived document plan")
        status=(f"Project `{pid}` created with {count} unique source(s). The source is stored and ready. "
                "The proposal outline and complete first draft are being prepared before the editor opens.")
        return pid,status,"",[],gr.update(choices=section_choices,value=(section_choices[0] if section_choices else None))
    if progress:progress(0.76,desc="Stage 4 of 4 · Local model is extracting structured grant requirements")
    analysis=api("POST",f"/api/projects/{pid}/analyze-requirements",timeout=600)
    enabled=set((workflow or {}).get("enabled_modules") or [])
    comp=api("POST",f"/api/projects/{pid}/compliance/compile",timeout=900) if "sponsor_compliance" in enabled else {"profile":{"rules":[]}}
    rule_count=len((comp.get("profile") or {}).get("rules") or [])
    return pid,f"Project `{pid}` created. {count} unique source(s), {analysis['count']} atomic grant requirements, and {rule_count} deterministic sponsor/submission rules were compiled from the opportunity. Review and approve both before final submission.","",requirement_rows(analysis["requirements"]),gr.update(choices=section_choices,value=(section_choices[0] if section_choices else None))

def compile_grant_ask(project,progress=gr.Progress()):
    project=_require_project(project)
    progress(0.05,desc="Stage 1 of 4 · Loading the authoritative source and approved workflow context")
    progress(0.20,desc="Stage 2 of 4 · Local model is extracting requirements, eligibility, dates, attachments, and review criteria")
    analysis=api("POST",f"/api/projects/{project}/analyze-requirements",timeout=600)
    progress(0.88,desc="Stage 3 of 4 · Validating the structured response and exact source provenance")
    progress(0.96,desc="Stage 4 of 4 · Loading the saved requirements for human review")
    requirements=analysis.get("requirements") or [];preview=requirements_response_preview(requirements)
    preview_text=f"  \n**Model-output preview:** {preview}" if preview else ""
    return requirement_rows(requirements),f"Compiled **{analysis.get('count',0)}** atomic grant requirements. Review the structured solicitation profile, correct any extraction issues, and approve it before framework generation.{preview_text}"

def poll_grant_ask_compilation(project):
    """Return one 65-second compilation heartbeat; never starts a model run."""
    if not (project or "").strip():return gr.skip(),gr.skip()
    workflow=api("GET",f"/api/projects/{project}/workflow")
    runs=[run for run in (workflow.get("generation_runs") or []) if run.get("task_kind")=="requirement_decomposition"]
    if not runs:return gr.skip(),"**Stage 1 of 4 · Ready.** The authoritative grant source is stored. Select **Compile grant ask** to begin structured extraction."
    latest=runs[0];status=latest.get("status")
    if status=="running":
        provider=latest.get("provider") or "configured provider";model=latest.get("model") or "configured model"
        source=api("GET",f"/api/projects/{project}/opportunity-source");input_preview=model_response_preview(source.get("text") or "")
        contract=latest.get("output_contract_name") or "structured requirements";version=latest.get("output_contract_version")
        contract_label=f"{contract} v{version}" if version is not None else contract
        preview_text=f"  \n**Input currently in flight:** {input_preview}" if input_preview else "  \n**Input currently in flight:** Source metadata is loaded; no safe text preview was available."
        return gr.skip(),(f"**Stage 2 of 4 · Model analysis active.** **{provider} / {model}** has been processing for **{generation_elapsed(latest.get('started_at'))}** and is producing **{html.escape(str(contract_label))}**. "
                          f"No committed response is available to preview yet. The provider does not expose a trustworthy item-level percentage; this status refreshes every 1 minute 5 seconds.{preview_text}")
    if status=="failed":return gr.skip(),f"**Compilation failed.** {html.escape(str(latest.get('error') or 'No error detail was returned.'))} Correct the reported issue and start a new compilation run."
    requirements=api("GET",f"/api/projects/{project}/requirements")
    preview=requirements_response_preview(requirements);preview_text=f"  \n**Model-output preview:** {preview}" if preview else ""
    return requirement_rows(requirements),f"**Stage 4 of 4 · Compilation complete.** The validated result contains **{len(requirements)}** atomic requirements. Review, correct, and approve the structured results.{preview_text}"

def prepare_initial_grant_document(project,progress=None):
    project=_require_project(project)
    if progress:progress(0.22,desc="Stage 2 of 4 · Model is deriving the ordered proposal sections from the complete grant ask")
    plan=api("POST",f"/api/projects/{project}/editor/plan",timeout=2400)
    sections=(plan.get("sections") or [])
    if not sections:raise gr.Error("The model returned no usable grant sections, so the editor was not opened.")
    total=len(sections);drafted=0;draft_chunks=0;generation_runs=list(plan.get("generation_run_ids") or [])
    for index,section in enumerate(sections,1):
        title=str(section.get("title") or "").strip()
        key=str(section.get("section_key") or "").strip()
        if not title or not key:raise gr.Error(f"The model-derived outline contained an invalid section at position {index}.")
        if progress:
            fraction=0.30+(0.60*((index-1)/max(total,1)))
            progress(fraction,desc=f"Stage 3 of 4 · Drafting section {index} of {total}: {title} · reducing source chunks, drafting bounded parts, and assembling the section")
        result=api("POST",f"/api/projects/{project}/sections/{key}/rewrite",json={
            "title":title,"description":section.get("description") or "","base_version_id":section.get("latest_version"),"high_value":False,"initial_build":True,
        },timeout=2400)
        text=str(result.get("text") or "").strip()
        if not text:raise gr.Error(f"The model returned no assembled text for '{title}', so the editor was not opened.")
        drafted+=1;draft_chunks+=int((result.get("chunking") or {}).get("section_chunks") or 1)
        if result.get("generation_run_id"):generation_runs.append(result["generation_run_id"])
        if progress:
            preview=model_response_preview(text)
            progress(0.30+(0.60*(index/max(total,1))),desc=f"Stage 3 of 4 · Assembled and saved section {index} of {total}: {title} · response preview: {preview}")
    if progress:progress(0.94,desc="Stage 4 of 4 · Verifying every assembled section and loading the immutable saved versions")
    document=api("GET",f"/api/projects/{project}/editor/document")
    saved=document.get("sections") or []
    incomplete=[str(section.get("title") or section.get("section_key") or "Untitled") for section in saved if not str(section.get("body") or "").strip() or not section.get("version")]
    if len(saved)!=total or incomplete:
        detail=", ".join(incomplete) if incomplete else f"expected {total} sections but loaded {len(saved)}"
        raise gr.Error(f"The assembled grant did not pass final verification ({detail}), so the editor was not opened.")
    if progress:progress(1.0,desc=f"Stage 4 of 4 · Complete grant assembled: {drafted} sections from {draft_chunks} bounded model response chunks")
    return {"sections":drafted,"draft_chunks":draft_chunks,"generation_runs":len(generation_runs),"plan_chunking":plan.get("chunking") or {}}

def configured_project_creation(title,sponsor,mechanism,deadline,grant_type,source,source_url,source_text,supporting,brand,
                                preset_key,selected,required,review_mode,review_required,routing_mode,team_rows,progress=None):
    # Resolve ownership from the authenticated session at creation time; hidden
    # browser state is never authoritative for access or authorship.
    if progress:progress(0.02,desc="Stage 1 of 4 · Confirming your authenticated project ownership")
    owner=(api("GET","/api/me").get("id") or "").strip()
    if not owner:raise gr.Error("Your authenticated account could not be resolved. Sign in again.")
    workflow=build_workflow_config(preset_key,selected,required,grant_type,deadline,review_mode,review_required,routing_mode)
    pid,status,notice,requirements,sections=create_project(title,sponsor,mechanism,source,source_url,source_text,supporting,brand,workflow,owner,False,progress)
    selected_set=set(selected or [])
    if progress:progress(0.20,desc="Stage 2 of 4 · Applying team invitations, collaboration settings, and enabled tools")
    invite_reports=[]
    members=_records(team_rows,["Email","Role"])
    for row in members:
        email=_cell(row.get("Email"));role=_cell(row.get("Role"))
        if not email:continue
        if role not in PROJECT_ROLE_LABELS:raise gr.Error(f"Unsupported project role for {email}: {row.get('Role')}")
        try:
            invite=api("POST",f"/api/projects/{pid}/invites",json={"email":email,"role":role,"expires_in_days":7})
            invite=_ensure_invite_email(invite,title.strip(),role)
            invite_reports.append(invite_delivery_summary(invite))
        except Exception as error:
            invite_reports.append(f"Invitation for **{html.escape(email)}** could not be created: {html.escape(str(error))}")
    if "team_collaboration" in selected_set:
        approvers=[owner]
        routed_artifacts=["solicitation_profile","research_framework","aim_set","literature_manifest","proposal_section","proposal_snapshot"]
        if "review_simulator" in selected_set:routed_artifacts.append("review_simulation")
        routing={"schema_version":1,"project_owner_user_id":owner,"routes":[
            {"artifact_type":artifact,"owner_user_id":owner,"approver_user_ids":approvers,"minimum_approvals":1}
            for artifact in routed_artifacts
        ]}
        routing_record=api("POST",f"/api/projects/{pid}/workflow/artifacts/collaboration_record",json={"body":routing,"source":"human_wizard","author":owner,"expected_version":None})
        api("POST",f"/api/projects/{pid}/workflow/artifacts/collaboration_record/approve",json={"version":routing_record["version"],"approver":owner})
    document_report=prepare_initial_grant_document(pid,progress)
    workflow_data=api("GET",f"/api/projects/{pid}/workflow")
    workflow_status=api("GET",f"/api/projects/{pid}/workflow/status")
    summary=(f"### {html.escape(title.strip())}\nWorkflow definition **v{workflow_data.get('definition_version')}** · configuration **v{workflow_data.get('config_version')}** · "
             f"{len(selected_set)} optional capabilities · model routing **{routing_mode}**. The shared server is the source of truth for every teammate.")
    visibility={key:key in selected_set for key in WORKFLOW_MODULES}
    status=(f"Shared grant and authoritative source saved. The configured model prepared and verified "
            f"{document_report['sections']} proposal sections from {document_report['draft_chunks']} bounded response chunks.")
    return (gr.update(visible=False),pid,title.strip(),sponsor or "",mechanism or "",status,notice,requirements,sections,summary,workflow_status,
            gr.update(visible=visibility.get("investigator_interview",False)),
            gr.update(visible=True),
            gr.update(visible=visibility.get("clinical_design",False)),
            gr.update(visible=visibility.get("competitive_intelligence",False)),
            gr.update(visible=visibility.get("sponsor_compliance",False)),
            gr.update(visible=visibility.get("review_simulator",False)),
            gr.update(visible=visibility.get("advanced_workbench",False)),
            "\n\n---\n\n".join(invite_reports))

def configured_project_creation_ui(title,sponsor,mechanism,deadline,grant_type,source,source_url,source_text,supporting,brand,
                                   preset_key,selected,required,review_mode,review_required,routing_mode,team_rows,edit_project=None,progress=gr.Progress()):
    """Keep the wizard usable and expose a durable result when creation fails."""
    try:
        if (edit_project or "").strip():
            pid=edit_project.strip();current=api("GET",f"/api/projects/{pid}/workflow");proposed=build_workflow_config(preset_key,selected,required,grant_type,deadline,review_mode,review_required,routing_mode)
            impact=api("POST",f"/api/projects/{pid}/workflow/impact",json={"workflow":proposed})
            if impact.get("destructive"):raise gr.Error("The server marked this workflow change destructive; no changes were applied.")
            actor=(api("GET","/api/me").get("id") or "").strip()
            if not actor:raise gr.Error("Your authenticated account could not be resolved. Sign in again.")
            api("PATCH",f"/api/projects/{pid}/workflow",json={"workflow":proposed,"expected_config_version":int(current.get("config_version")),"actor":actor})
            if str(title or "").strip():api("PATCH",f"/api/projects/{pid}",json={"title":str(title).strip(),"archived":None})
            loaded=list(load_project(pid));workflow=project_workflow_ui(pid)
            return (gr.update(visible=False),*loaded,*workflow,
                    "Workflow configuration saved on the existing grant. Historical artifacts from removed tools remain preserved, while only selected tools remain active.",
                    gr.update(value="Save workflow changes →",interactive=True))
        creation=configured_project_creation(title,sponsor,mechanism,deadline,grant_type,source,source_url,source_text,supporting,brand,
                                             preset_key,selected,required,review_mode,review_required,routing_mode,team_rows,progress)
        pid=creation[1]
        invite_report=creation[-1]
        loaded=list(load_project(pid))
        if invite_report:
            loaded[4]=f"{loaded[4]}\n\n### Teammate invitation results\n{invite_report}"
        workflow=project_workflow_ui(pid)
        document=api("GET",f"/api/projects/{pid}/editor/document");prepared=document.get("sections") or []
        creation_status=(f"Shared grant created and loaded. **{len(prepared)} model-prepared sections are assembled, versioned, and ready for human review.** "
                         "The editor opened only after every section returned non-empty saved text. Humans can edit and save any wording before publishing.")
        if invite_report:creation_status+=f"\n\n### Teammate invitation results\n{invite_report}"
        return (gr.update(visible=False),*loaded,*workflow,
                creation_status,
                gr.update(value="Create shared grant →",interactive=True))
    except Exception as error:
        message=html.escape(str(error).strip() or "The server did not return an error description.")
        return (*(gr.skip() for _ in range(26)),f"### Grant creation failed\n{message}",gr.update(value="Try creating the grant again →",interactive=True))

def global_navigation_state(project,action):
    action=str(action or "").strip()
    field_count=11+len(WORKFLOW_REGISTRY["optional_modules"])
    if action=="projects":
        values=[gr.update(visible=True),None,gr.update(value="Create shared grant →",interactive=True),*wizard_page_updates(1),*(gr.skip() for _ in range(field_count)),gr.update(value="")]
        return tuple(values)
    if action!="workflow":return tuple(gr.skip() for _ in range(4+len(wizard_page_updates(1))+field_count))
    project=_require_project(project);metadata=api("GET",f"/api/projects/{project}");workflow=api("GET",f"/api/projects/{project}/workflow");config=workflow.get("config") or workflow
    enabled=set(config.get("enabled_modules") or []);modes=["include" if item["key"] in enabled else "skip" for item in WORKFLOW_REGISTRY["optional_modules"]]
    base=[gr.update(visible=True),project,gr.update(value="Save workflow changes →",interactive=True),*wizard_page_updates(2)]
    fields=[metadata.get("title") or "",metadata.get("sponsor") or "",metadata.get("mechanism") or "",config.get("target_deadline") or "",config.get("grant_type") or "custom",config.get("template") or "custom_configuration_v1",list(enabled),list(config.get("required_modules") or []),bool(config.get("review_required")),config.get("review_mode"),config.get("model_routing_mode") or os.getenv("MODEL_ROUTING_MODE","local_only"),*modes]
    return tuple(base+fields+[gr.update(value="")])

def wizard_to_preview(title,sponsor,mechanism,deadline,selected,required,review_mode,review_required,routing_mode,team_rows):
    return wizard_page_updates(WIZARD_PREVIEW_PAGE)+[workflow_preview_html(title,sponsor,mechanism,deadline,selected,required,review_mode,review_required,routing_mode,team_rows)]

def load_project(pid):
    if not pid:raise gr.Error("Choose a project.")
    p=api("GET",f"/api/projects/{pid}");reqs=api("GET",f"/api/projects/{pid}/requirements");sections=api("GET",f"/api/projects/{pid}/sections");choices=[x["title"] for x in sections]
    selected=choices[0] if choices else None
    state=load_section(pid,p["title"],selected) if selected else (None,"",section_preview(pid,p["title"],"Section",""),"No sections configured.","",slug("Section"),None,"")
    notice=global_competitive_update_banner(p.get("competitive_updates") or {})
    return pid,p["title"],p.get("sponsor") or "",p.get("mechanism") or "",f"Opened `{pid}` at workflow stage **{p['stage']}**.",notice,requirement_rows(reqs),gr.update(choices=choices,value=selected),*state

def load_project_if_available(pid):
    """Hydrate the workspace after creation without masking a creation error."""
    if not pid:
        return tuple(gr.skip() for _ in range(16))
    return load_project(pid)

def open_project_workspace(pid):
    """Open one saved grant and hydrate its complete persisted workflow atomically."""
    loaded=load_project(pid)
    workflow=project_workflow_ui(pid)
    return gr.update(visible=False),*loaded,*workflow

def refresh_shared_updates(project,kind,subject_key,search_plan_version,literature_version,include_archived):
    """Fetch one user-requested shared snapshot; no timer or background loop is involved."""
    wizard_projects=refresh_projects()
    catalog=project_catalog(include_archived)
    if not (project or "").strip():
        return (wizard_projects,*catalog,*(gr.skip() for _ in range(24)),
                "Saved grants refreshed. Open a grant to refresh its shared workflow and collaboration data.")
    collaboration=poll_collaboration_snapshot(project,kind,subject_key,search_plan_version,literature_version)
    compilation=poll_grant_ask_compilation(project)
    workflow=project_workflow_ui(project)
    return (wizard_projects,*catalog,*collaboration,*compilation,*workflow,
            "Shared updates loaded from the server. Teammate activity, tasks, messages, approvals, workflow progress, and compilation status are current as of this refresh.")

def _portable_archive_payload(upload):
    if not upload:raise gr.Error("Choose a Grantspace portable project archive first.")
    path=Path(file_path(upload))
    if path.suffix.lower()!=".zip" or path.stat().st_size>128*1024*1024:raise gr.Error("Portable project archives must be ZIP files no larger than 128 MiB.")
    try:
        with zipfile.ZipFile(path) as archive:
            files=[item for item in archive.infolist() if not item.is_dir()]
            if len(files)!=1 or files[0].filename!="grantspace-project.json":raise gr.Error("Archive must contain exactly one root file named grantspace-project.json.")
            info=files[0];mode=(info.external_attr>>16)&0o170000
            if mode==0o120000 or info.file_size>128*1024*1024:raise gr.Error("Links and oversized project payloads are not allowed.")
            raw=archive.read(info)
    except gr.Error:raise
    except (OSError,zipfile.BadZipFile,RuntimeError) as error:raise gr.Error(f"Portable project archive is invalid: {error}")
    try:return json.loads(raw)
    except (UnicodeDecodeError,json.JSONDecodeError) as error:raise gr.Error(f"Portable project JSON is invalid: {error}")

def import_portable_project(upload):
    package=_portable_archive_payload(upload)
    validation=api("POST","/api/project-imports/validate",json={"package":package},timeout=120)
    imported=api("POST","/api/project-imports",json={"package":package},timeout=600)
    projects=api("GET","/api/projects");choices=[(f"{item['title']} · {item.get('sponsor') or 'No sponsor'}",item["id"]) for item in projects]
    counts=validation.get("counts") or {}
    return gr.update(choices=choices,value=imported["id"]),(f"✓ Validated and transactionally imported **{imported['title']}** as `{imported['id']}`. "
        f"Restored {counts.get('documents',0)} documents, {counts.get('section_versions',0)} section versions, "
        f"{counts.get('workflow_artifacts',0)} workflow artifacts, and {counts.get('export_snapshots',0)} immutable exports.")

def export_portable_project(project):
    if not project:raise gr.Error("Create or open a project first.")
    package=api("GET",f"/api/projects/{project}/portable-export",timeout=600)
    export_root=WORKSPACE/"portable_exports";export_root.mkdir(parents=True,exist_ok=True)
    path=export_root/f"grantspace-project-{project}-{uuid.uuid4().hex[:10]}.zip"
    raw=json.dumps(package,ensure_ascii=False,separators=(",",":"),sort_keys=True).encode("utf-8")
    with zipfile.ZipFile(path,"w",compression=zipfile.ZIP_DEFLATED,compresslevel=9) as archive:
        archive.writestr("grantspace-project.json",raw)
    return str(path),f"Created validated portable project archive with payload SHA-256 `{package.get('payload_sha256')}`."

def approve_requirements(project):
    if not project:raise gr.Error("Create or open a project first.")
    d=api("POST",f"/api/projects/{project}/requirements/approve");return f"✓ Human approved {d['approved']} parsed requirements. Workflow stage: `{d['stage']}`."

def next_open(questions):return next((q for q in questions if q.get("status")=="open"),None)
def render_question(q):
    if not q:return '<div class="question-card"><b>Interview complete.</b> No unresolved investigator questions remain.</div>'
    choices=q.get("choices") or [];choice_html=("<br><b>Allowed choices:</b> "+", ".join(html.escape(str(x)) for x in choices)) if choices else "";unit=(f" <b>Unit:</b> {html.escape(str(q['unit']))}" if q.get("unit") else "")
    return f'<div class="question-card"><b>{html.escape(q.get("requirement_id",""))}</b><h3>{html.escape(q.get("question",""))}</h3><p>{html.escape(q.get("why_needed") or "")}</p><small>Type: {html.escape(q.get("answer_type","text"))}{unit}</small>{choice_html}</div>'

def generate_interview(project):
    if not project:raise gr.Error("Create or open a project first.")
    d=api("POST",f"/api/projects/{project}/interview/generate",timeout=600);q=next_open(d.get("questions",[]));return d.get("questions",[]),q,render_question(q),"",f"Generated {d['count']} unresolved question(s) using `{d['model']}`.",d.get("questions",[])

def parse_typed_answer(q,raw):
    t=q.get("answer_type","text");raw=(raw or "").strip()
    if not raw:raise gr.Error("An answer is required. Use classification 'unknown' when the fact is not known.")
    if t=="integer":return int(raw)
    if t in ("number","percentage"):
        v=float(raw)
        if t=="percentage" and not 0<=v<=100:raise gr.Error("Percentage answers must be between 0 and 100.")
        return v
    if t=="boolean":
        v=raw.lower()
        if v in ("true","yes","1"):return True
        if v in ("false","no","0"):return False
        raise gr.Error("Enter true/false or yes/no.")
    if t=="date":
        try:return date.fromisoformat(raw).isoformat()
        except ValueError:raise gr.Error("Use ISO date format YYYY-MM-DD.")
    if t=="choice":
        choices=q.get("choices") or []
        if choices and raw not in choices:raise gr.Error("Answer must exactly match an allowed choice.")
    return raw

def submit_answer(project,questions,current_q,raw,confidence,classification,notes,answered_by):
    if not current_q:return questions,None,render_question(None),"","Interview is already complete."
    value=parse_typed_answer(current_q,raw);d=api("POST",f"/api/projects/{project}/interview/answer",json={"question_id":current_q["id"],"value":value,"confidence":confidence,"classification":classification,"notes":notes or None,"answered_by":answered_by or None})
    refreshed=api("GET",f"/api/projects/{project}/interview");q=next_open(refreshed);return refreshed,q,render_question(q),"",(f"Saved answer. {d['open_questions']} question(s) remain." if q else "✓ Interview complete; research/writing gates are now open.")

def run_research(project,search_plan_version,results_per):
    if not search_plan_version:raise gr.Error("Generate, review, and approve a literature search plan before running research.")
    d=api("POST",f"/api/projects/{project}/research/run",json={"search_plan_version":int(search_plan_version),"results_per_query":int(results_per)},timeout=1200);status=f"Atomically committed research run `{d.get('run_id')}` from approved search plan v{d.get('search_plan_version')}, with {d.get('sources_saved') or 0} assessed source(s).";status+=(f" {len(d['failures'])} isolated failure(s) were recorded as unresolved risk." if d.get("failures") else "");return evidence_rows(d.get("evidence",[])),status

def reviewer_role_rows(project):
    data=api("GET",f"/api/projects/{project}/review-panel/roles")
    rows=[[role.get("key"),role.get("title"),role.get("description"),", ".join(role.get("criterion_ids") or [])] for role in data.get("roles") or []]
    return rows,data.get("synthetic_review_notice") or ""

def create_panel_plan(project,mode):
    plan=api("POST",f"/api/projects/{project}/review-panel/plan",json={"mode":mode})
    rows=[[role.get("key"),role.get("title"),role.get("description"),", ".join(role.get("criterion_ids") or [])] for role in plan.get("roles") or []]
    return plan.get("id"),rows,f"Draft panel plan `{plan.get('id')}` was derived from approved solicitation profile v{plan.get('solicitation_profile_version')}. Approval is required before execution.",plan

def approve_panel_plan(project,plan_id):
    if not plan_id:raise gr.Error("Create a panel plan first.")
    plan=api("POST",f"/api/projects/{project}/review-panel/plan/{plan_id}/approve",json={})
    return f"✓ Panel plan `{plan_id}` approved by {plan.get('approved_by_user_id')}.",plan

def review_result_views(data):
    result=data.get("result") or {}
    rows=[]
    for review in result.get("reviews") or []:
        for score in review.get("criterion_scores") or []:
            rows.append([review.get("reviewer_archetype"),score.get("criterion_id"),score.get("score"),score.get("confidence"),"\n".join(score.get("strengths") or []),"\n".join(score.get("weaknesses") or []),", ".join(score.get("proposal_anchors") or [])])
    tasks=[[index,task.get("priority"),task.get("title"),task.get("description"),task.get("rationale"),", ".join(task.get("proposal_anchors") or [])] for index,task in enumerate(result.get("revision_tasks") or [])]
    status=f"Review run `{data.get('id')}` · **{data.get('status')}** · immutable result SHA-256 `{str(data.get('result_sha256') or '')[:20]}…`"
    return data.get("id"),result.get("panel_summary") or {},rows,tasks,result.get("causal_analysis") or {},status,data

def execute_review_panel(project,plan_id):
    if not plan_id:raise gr.Error("Create and approve a solicitation-derived panel plan first.")
    data=api("POST",f"/api/projects/{project}/review-simulations",json={"panel_plan_id":plan_id},timeout=3600)
    return review_result_views(data)

def approve_review_result(project,run_id):
    if not run_id:raise gr.Error("Run and validate a review simulation first.")
    artifact=api("POST",f"/api/projects/{project}/review-simulations/{run_id}/approve",json={})
    progress=artifact.get("approval_progress") or {}
    if artifact.get("approved"):
        return f"✓ Approved immutable review run `{run_id}` as workflow artifact v{artifact.get('version')}."
    return (f"Approval recorded for immutable review run `{run_id}` "
            f"({progress.get('approvals',0)} of {progress.get('minimum_approvals','?')}); additional configured approvals remain.")

def create_revision_tasks(project,run_id,indexes,owner_id,due_at):
    if not run_id:raise gr.Error("Run a review simulation first.")
    selected=[]
    for value in _split_values(indexes):
        try:selected.append(int(value))
        except ValueError:raise gr.Error("Task indexes must be comma-separated integers.")
    data=api("POST",f"/api/projects/{project}/review-simulations/{run_id}/tasks",json={"task_indexes":selected,"owner_user_id":str(owner_id or "").strip(),"due_at":str(due_at or "").strip() or None})
    return data,f"✓ Created {len(data.get('created_tasks') or [])} assigned revision task(s) from immutable review run `{run_id}`."

def load_causal_models(project,run_id):
    if not run_id:raise gr.Error("Run a causal review first.")
    models=api("GET",f"/api/projects/{project}/review-simulations/{run_id}/causal-models")
    latest=(models or [{}])[0]
    return latest.get("body") or {},models,f"Loaded {len(models)} append-only causal model version(s)."

def save_causal_editor(project,run_id,body,confirmed):
    if not isinstance(body,dict):raise gr.Error("The causal editor must contain a structured object.")
    saved=api("POST",f"/api/projects/{project}/review-simulations/{run_id}/causal-models",json={"body":body,"confirmed":bool(confirmed)})
    return body,f"Saved causal model v{saved.get('version')} as {'methodologist-confirmed' if confirmed else 'inferred/unconfirmed'}; prior versions remain immutable."

def rebuild_index(project):
    d=api("POST",f"/api/projects/{project}/index/rebuild",timeout=1800);return f"✓ Knowledge index: {d['rows']} records × {d['dimensions']} dimensions with `{d['embedding_model']}`.",d

def index_status(project):
    d=api("GET",f"/api/projects/{project}/index/status");return (("✓ Index is current." if d.get("fresh") else "⚠ Index is missing/stale; it will rebuild before retrieval."),d)

def test_retrieval(project,query,k):
    if not (query or "").strip():raise gr.Error("Enter a retrieval query.")
    d=api("POST",f"/api/projects/{project}/retrieve",json={"query":query,"k":int(k)},timeout=1800)
    return [[x.get("score"),x.get("semantic"),x.get("lexical"),x.get("evidence"),x.get("freshness"),x.get("graph_boost"),x.get("record",{}).get("kind"),x.get("record",{}).get("source_ref"),x.get("record",{}).get("text","")[:280]] for x in d]


def _cell(v):
    if v is None:return ""
    if isinstance(v,float) and math.isnan(v):return ""
    s=str(v).strip()
    return "" if s.lower()=="nan" else s

def _opt_float(v):
    s=_cell(v)
    if not s:return None
    x=float(s)
    if not math.isfinite(x):raise gr.Error("Numeric inputs must be finite values.")
    return x

def _opt_int(v):
    x=_opt_float(v)
    if x is None:return None
    if not x.is_integer():raise gr.Error("Integer inputs cannot contain a fractional value.")
    return int(x)

def _lines(v):return [x.strip() for x in (v or "").splitlines() if x.strip()]
def _csv(v):return [x.strip() for x in _cell(v).split(",") if x.strip()]
def _truth(v):
    if isinstance(v,bool):return v
    return str(v or "").strip().lower() in {"true","yes","1","y"}

def _records(table,headers):
    if table is None:return []
    if hasattr(table,"to_dict"):
        try:
            rows=table.to_dict(orient="records")
            return [{h:r.get(h) for h in headers} for r in rows]
        except Exception:pass
    rows=table if isinstance(table,list) else []
    out=[]
    for row in rows:
        if isinstance(row,dict):out.append({h:row.get(h) for h in headers})
        elif isinstance(row,(list,tuple)):out.append({h:(row[i] if i<len(row) else None) for i,h in enumerate(headers)})
    return out

AIM_HEADERS=["Aim ID","Title","Hypothesis","Expected endpoint type","Endpoint IDs","Expected result","Risk","Alternative strategy"]
ARM_HEADERS=["Arm ID","Name","Intervention","Comparator"]
ENDPOINT_HEADERS=["Endpoint ID","Name","Type","Primary","Analysis family"]
TIMELINE_HEADERS=["Task ID","Name","Start month","Duration months","Dependencies"]
RESOURCE_HEADERS=["Resource ID","Name","Required","Available","Notes"]

def build_clinical_study(clinical_problem,knowledge_gap,central_hypothesis,disease,disease_stage,biomarker,inclusion,exclusion,
                         design_type,study_phase,randomization,allocation_ratio,blinding,follow_up_months,design_sites,
                         available_patients,eligibility_pct,biomarker_pct,consent_pct,target_enrollment,accrual_months,recruitment_sites,
                         test_type,alpha,power,attrition_pct,control_rate,treatment_rate,null_rate,alternative_rate,mean_delta,std_dev,hazard_ratio,event_probability,
                         aims_table,arms_table,endpoints_table,timeline_table,resources_table):
    aims=[]
    for r in _records(aims_table,AIM_HEADERS):
        aid=_cell(r.get("Aim ID"))
        if not aid:continue
        aims.append({"aim_id":aid,"title":_cell(r.get("Title")),"hypothesis":_cell(r.get("Hypothesis")),"expected_endpoint_type":_cell(r.get("Expected endpoint type")),"endpoint_ids":_csv(r.get("Endpoint IDs")),"expected_result":_cell(r.get("Expected result")),"risk":_cell(r.get("Risk")),"alternative_strategy":_cell(r.get("Alternative strategy"))})
    arms=[]
    for r in _records(arms_table,ARM_HEADERS):
        aid=_cell(r.get("Arm ID"));name=_cell(r.get("Name"))
        if not aid and not name:continue
        arms.append({"arm_id":aid,"name":name,"intervention":_cell(r.get("Intervention")),"comparator":_truth(r.get("Comparator"))})
    endpoints=[]
    for r in _records(endpoints_table,ENDPOINT_HEADERS):
        eid=_cell(r.get("Endpoint ID"));name=_cell(r.get("Name"))
        if not eid and not name:continue
        endpoints.append({"endpoint_id":eid,"name":name,"endpoint_type":_cell(r.get("Type")),"primary":_truth(r.get("Primary")),"analysis_family":_cell(r.get("Analysis family"))})
    timeline=[]
    for r in _records(timeline_table,TIMELINE_HEADERS):
        tid=_cell(r.get("Task ID"));name=_cell(r.get("Name"))
        if not tid and not name:continue
        timeline.append({"task_id":tid,"name":name,"start_month":float(_cell(r.get("Start month")) or 0),"duration_months":float(_cell(r.get("Duration months")) or 0),"dependencies":_csv(r.get("Dependencies"))})
    resources=[]
    for r in _records(resources_table,RESOURCE_HEADERS):
        rid=_cell(r.get("Resource ID"));name=_cell(r.get("Name"))
        if not rid and not name:continue
        resources.append({"resource_id":rid,"name":name,"required":_truth(r.get("Required")),"available":_truth(r.get("Available")),"notes":_cell(r.get("Notes"))})
    return {
      "clinical_problem":clinical_problem or "","knowledge_gap":knowledge_gap or "","central_hypothesis":central_hypothesis or "",
      "population":{"disease":disease or "","stage":disease_stage or "","biomarker_criteria":biomarker or "","inclusion_criteria":_lines(inclusion),"exclusion_criteria":_lines(exclusion)},
      "design":{"design_type":design_type or "","phase":study_phase or "","randomization":randomization or "","allocation_ratio":allocation_ratio or "","blinding":blinding or "","follow_up_months":_opt_float(follow_up_months),"sites":_opt_int(design_sites)},
      "recruitment":{"available_patients_per_site_month":_opt_float(available_patients),"eligibility_rate_pct":_opt_float(eligibility_pct),"biomarker_positive_rate_pct":_opt_float(biomarker_pct),"consent_rate_pct":_opt_float(consent_pct),"target_enrollment":_opt_int(target_enrollment),"accrual_months":_opt_float(accrual_months),"sites":_opt_int(recruitment_sites)},
      "statistics":{"test_type":test_type or "","alpha":_opt_float(alpha),"power":_opt_float(power),"attrition_pct":_opt_float(attrition_pct),"control_rate":_opt_float(control_rate),"treatment_rate":_opt_float(treatment_rate),"null_rate":_opt_float(null_rate),"alternative_rate":_opt_float(alternative_rate),"mean_delta":_opt_float(mean_delta),"std_dev":_opt_float(std_dev),"hazard_ratio":_opt_float(hazard_ratio),"event_probability":_opt_float(event_probability)},
      "aims":aims,"arms":arms,"endpoints":endpoints,"timeline":timeline,"resources":resources
    }

def clinical_summary(a):
    if not a:return "No clinical assessment available."
    rec=a.get("recruitment") or {};bits=[]
    if rec.get("complete"):
        bits.append(f"**Recruitment:** expected {rec.get('expected_enrollments_per_month',0):.2f}/month; required {rec.get('required_enrollments_per_month') if rec.get('required_enrollments_per_month') is not None else 'not configured'}/month; estimated accrual {rec.get('estimated_accrual_months') if rec.get('estimated_accrual_months') is not None else 'not configured'} months.")
        if rec.get("feasible_within_planned_window") is False:bits.append("⚠ Recruitment assumptions do not meet the planned accrual window.")
    stats=a.get("statistics") or {}
    if stats.get("complete"):bits.append(f"**Sample size:** raw N={stats.get('raw_total_n')}, attrition-adjusted N={stats.get('adjusted_total_n')} ({stats.get('method')}).")
    errs=a.get("errors") or [];warn=a.get("warnings") or [];conf=(a.get("cross_section_consistency") or {}).get("count",0)
    bits.append(f"**Deterministic checks:** {len(errs)} error(s), {len(warn)} warning(s), {conf} approved-section consistency conflict(s).")
    return "\n\n".join(bits)

def _csv_numbers(v,integer=False):
    out=[]
    for x in _csv(v):out.append(int(float(x)) if integer else float(x))
    return out

def run_feasibility_scenarios(project,sites_values,consent_values,biomarker_values):
    if not project:raise gr.Error("Create or open a project first.")
    d=api("POST",f"/api/projects/{project}/clinical/scenarios",json={"sites":_csv_numbers(sites_values,True),"consent_rates_pct":_csv_numbers(consent_values),"biomarker_positive_rates_pct":_csv_numbers(biomarker_values)})
    rows=[[x.get("sites"),x.get("consent_rate_pct"),x.get("biomarker_positive_rate_pct"),x.get("expected_enrollments_per_month"),x.get("required_enrollments_per_month"),x.get("estimated_accrual_months"),x.get("feasible"),x.get("shortfall_per_month")] for x in d.get("rows",[])]
    return rows,f"Evaluated **{d.get('combinations',0)}** recruitment scenarios in the Rust/Rayon clinical engine."

def save_clinical_study(project,*values):
    if not project:raise gr.Error("Create or open a project first.")
    study=build_clinical_study(*values)
    d=api("POST",f"/api/projects/{project}/clinical-study",json=study,timeout=300)
    a=d.get("assessment") or {}
    return a,clinical_summary(a)+f"\n\nSaved clinical study version **{d['saved']['version']}** (`{d['saved']['sha256'][:16]}…`)."

def load_clinical_study(project):
    if not project:raise gr.Error("Create or open a project first.")
    d=api("GET",f"/api/projects/{project}/clinical-study")
    if not d.get("exists"):
        return ("","","","","","","","","","","","","",None,None,None,None,None,None,None,None,None,None,"",None,None,None,None,None,None,None,None,None,None,[],[],[],[],[],{},"No saved clinical study exists yet.")
    st=d["study"];p=st.get("population") or {};des=st.get("design") or {};r=st.get("recruitment") or {};sp=st.get("statistics") or {}
    aims=[[x.get("aim_id"),x.get("title"),x.get("hypothesis"),x.get("expected_endpoint_type"),", ".join(x.get("endpoint_ids") or []),x.get("expected_result"),x.get("risk"),x.get("alternative_strategy")] for x in st.get("aims") or []]
    arms=[[x.get("arm_id"),x.get("name"),x.get("intervention"),x.get("comparator")] for x in st.get("arms") or []]
    eps=[[x.get("endpoint_id"),x.get("name"),x.get("endpoint_type"),x.get("primary"),x.get("analysis_family")] for x in st.get("endpoints") or []]
    timeline=[[x.get("task_id"),x.get("name"),x.get("start_month"),x.get("duration_months"),", ".join(x.get("dependencies") or [])] for x in st.get("timeline") or []]
    resources=[[x.get("resource_id"),x.get("name"),x.get("required"),x.get("available"),x.get("notes")] for x in st.get("resources") or []]
    a=api("GET",f"/api/projects/{project}/clinical-assessment")
    return (st.get("clinical_problem",""),st.get("knowledge_gap",""),st.get("central_hypothesis",""),p.get("disease",""),p.get("stage",""),p.get("biomarker_criteria",""),"\n".join(p.get("inclusion_criteria") or []),"\n".join(p.get("exclusion_criteria") or []),des.get("design_type",""),des.get("phase",""),des.get("randomization",""),des.get("allocation_ratio",""),des.get("blinding",""),des.get("follow_up_months"),des.get("sites"),r.get("available_patients_per_site_month"),r.get("eligibility_rate_pct"),r.get("biomarker_positive_rate_pct"),r.get("consent_rate_pct"),r.get("target_enrollment"),r.get("accrual_months"),r.get("sites"),sp.get("test_type",""),sp.get("alpha"),sp.get("power"),sp.get("attrition_pct"),sp.get("control_rate"),sp.get("treatment_rate"),sp.get("null_rate"),sp.get("alternative_rate"),sp.get("mean_delta"),sp.get("std_dev"),sp.get("hazard_ratio"),sp.get("event_probability"),aims,arms,eps,timeline,resources,a,clinical_summary(a)+f"\n\nLoaded clinical study version **{d['version']}**.")

def calculate_sample_size(project,test_type,alpha,power,attrition_pct,control_rate,treatment_rate,null_rate,alternative_rate,mean_delta,std_dev,hazard_ratio,event_probability):
    plan={"test_type":test_type or "","alpha":_opt_float(alpha),"power":_opt_float(power),"attrition_pct":_opt_float(attrition_pct),"control_rate":_opt_float(control_rate),"treatment_rate":_opt_float(treatment_rate),"null_rate":_opt_float(null_rate),"alternative_rate":_opt_float(alternative_rate),"mean_delta":_opt_float(mean_delta),"std_dev":_opt_float(std_dev),"hazard_ratio":_opt_float(hazard_ratio),"event_probability":_opt_float(event_probability)}
    d=api("POST",f"/api/projects/{project}/clinical/sample-size",json=plan)
    return d,(f"Calculated **N={d.get('adjusted_total_n')}** including attrition using {d.get('method')}." if d.get("complete") else "Statistics inputs are incomplete.")


def competitive_candidate_rows(data):
    rows=[]
    for c in data.get("candidates") or []:
        rows.append([
            c.get("rank"),c.get("name"),round(float(c.get("overall_score",0)),3),
            round(float(c.get("grant_score",0)),3),round(float(c.get("publication_score",0)),3),
            round(float(c.get("clinical_trial_score",0)),3),round(float(c.get("patent_ip_score",0)),3),
            round(float(c.get("technology_score",0)),3),round(float(c.get("breadth_score",0)),3),c.get("asset_count",0)
        ])
    return rows

def competitive_asset_rows(data,limit=250):
    names={c.get("candidate_key"):c.get("name") for c in data.get("candidates") or []}
    rows=[]
    for a in (data.get("assets") or [])[:limit]:
        rows.append([names.get(a.get("candidate_key"),a.get("candidate_key")),a.get("asset_type"),round(float(a.get("relevance",0)),3),a.get("title"),a.get("year"),a.get("provider"),a.get("url") or ""])
    return rows

def competitive_strategy_markdown(data):
    if not data.get("exists"):
        return "No competitive applicant intelligence has been run yet."
    strategy=data.get("strategy") or {}
    lines=["### Evidence-backed competitive positioning","> **Important:** organizations below are capability-overlap candidates inferred from public evidence, **not confirmed applicants**."]
    if strategy.get("market_summary"):lines.extend(["",strategy["market_summary"]])
    if strategy.get("positioning_principles"):
        lines.extend(["","#### Positioning principles"]+[f"- {x}" for x in strategy.get("positioning_principles") or []])
    if strategy.get("differentiators"):
        lines.extend(["","#### Differentiators to emphasize"])
        for x in strategy["differentiators"]:
            refs=", ".join(x.get("asset_keys") or [])
            lines.append(f"- **{x.get('theme','Differentiator')}** — {x.get('our_advantage','')}"+(f"  Public evidence: `{refs}`" if refs else ""))
    if strategy.get("gaps_to_close"):
        lines.extend(["","#### Competitive gaps to close"])
        for x in strategy["gaps_to_close"]:
            refs=", ".join(x.get("asset_keys") or [])
            lines.append(f"- **{x.get('gap','Gap')}** — {x.get('recommended_action','')}"+(f"  Public evidence: `{refs}`" if refs else ""))
    if strategy.get("candidate_notes"):
        lines.extend(["","#### Capability-matched organization notes"] )
        for x in strategy["candidate_notes"]:
            refs=", ".join(x.get("asset_keys") or [])
            lines.append(f"- `{x.get('candidate_key','candidate')}` — {x.get('why_relevant','')} **Positioning:** {x.get('how_to_outposition','')}"+(f"  Evidence: `{refs}`" if refs else ""))
    if strategy.get("section_guidance"):
        lines.extend(["","#### Section-specific positioning guidance"] )
        for x in strategy["section_guidance"]:
            refs=", ".join(x.get("asset_keys") or [])
            lines.append(f"- **{x.get('section_key','section')}** — {x.get('guidance','')}"+(f"  Evidence: `{refs}`" if refs else ""))
    if strategy.get("do_not_claim"):
        lines.extend(["","#### Do not claim"]+[f"- {x}" for x in strategy.get("do_not_claim") or []])
    return "\n".join(lines)

def generate_competitive_profile(project):
    if not project:raise gr.Error("Create or open a project first.")
    d=api("POST",f"/api/projects/{project}/competitive/profile/generate",timeout=900)
    return d,f"Generated competitive applicant capability profile version **{d.get('version')}** using `{d.get('model')}`. This describes public capabilities to search for; it does not assert who will apply."

def load_competitive(project):
    if not project:raise gr.Error("Create or open a project first.")
    # Self-healing load: the backend refreshes stale/expired/config-changed public intelligence before returning.
    data=api("GET",f"/api/projects/{project}/competitive",timeout=2400)
    profile=api("GET",f"/api/projects/{project}/competitive/profile")
    refreshed=data.get("auto_refreshed",False)
    reason=", ".join(data.get("refresh_reason") or [])
    if data.get("exists"):
        status=f"Competitive run **{data.get('run_id')}** is current."
        if refreshed: status+=f" Automatically refreshed{(' ('+reason+')') if reason else ''}."
        status+=f" Public-source age: **{data.get('age_seconds','n/a')}s**; refresh TTL: **{data.get('refresh_ttl_seconds','n/a')}s**."
    else:
        status="No competitive intelligence run exists yet."
    return profile,competitive_candidate_rows(data),competitive_asset_rows(data),data.get("provider_status") or [],competitive_strategy_markdown(data),data,status,global_competitive_update_banner(data.get("competitive_updates") or {})

def run_competitive_intelligence(project):
    if not project:raise gr.Error("Create or open a project first.")
    data=api("POST",f"/api/projects/{project}/competitive/refresh",json={},timeout=2400)
    providers=data.get("provider_status") or []
    ok=sum(1 for p in providers if p.get("ok")); total=len(providers)
    status=f"Competitive intelligence run **{data.get('run_id')}** completed with **{len(data.get('candidates') or [])}** capability-matched organizations and **{len(data.get('assets') or [])}** public assets. Providers healthy: {ok}/{total}."
    changed=((data.get("agentic_update") or {}).get("section_updates") or [])
    if changed:status+=f" **{COMPETITIVE_UPDATE_LABEL} refreshed {len(changed)} existing grant section(s); highlighted changes now require human review.**"
    return competitive_candidate_rows(data),competitive_asset_rows(data),providers,competitive_strategy_markdown(data),data,status,global_competitive_update_banner(data.get("competitive_updates") or {})

def load_section(project,project_title,section_title):
    if not project or not section_title:return None,"",section_preview(project,project_title,section_title or "Section",""),"No section selected.","",slug(section_title or "Section"),None,""
    key=slug(section_title);d=api("GET",f"/api/projects/{project}/sections/{key}",timeout=2400)
    latest=d.get("latest") if d.get("exists") else None;approved=d.get("approved") if d.get("exists") else None;update=d.get("competitive_update") or None
    body=(latest or {}).get("body","");version=(latest or {}).get("version")
    status=f"Latest version: {version or 'none'} · approved version: {(approved or {}).get('version','none')}"
    if update:status+=f" · competitive update event {update.get('event_id')} requires review"
    return version,body,section_preview(project,project_title,section_title,body,key,version,update),status,body,key,(update or {}).get("event_id"),competitive_update_banner(update,version)


COMPLIANCE_HEADERS=["Rule"]

def compliance_rule_text(rule):
    manual=str(rule.get("text_value") or "").strip()
    if rule.get("rule_type")=="manual_requirement" and manual:return manual
    target=str(rule.get("target") or "the application").replace("_"," ").strip()
    number=rule.get("numeric_value")
    values=", ".join(rule.get("list_value") or [])
    templates={
        "required_section":f"Include the required section: {target}.",
        "required_form":f"Include the required form: {target}.",
        "required_attachment":f"Include the required attachment: {target}.",
        "max_words":f"Limit {target} to {number:g} words." if number is not None else "",
        "min_words":f"Use at least {number:g} words for {target}." if number is not None else "",
        "max_pages":f"Limit {target} to {number:g} pages." if number is not None else "",
        "min_font_size_pt":f"Use a font size of at least {number:g} points." if number is not None else "",
        "min_margin_in":f"Use margins of at least {number:g} inches." if number is not None else "",
        "allowed_extensions":f"Use only these file types for {target}: {values}." if values else "",
        "deadline":f"Submit by {manual or target}.",
        "max_budget":f"Do not exceed a budget of {number:g}." if number is not None else "",
        "project_period_max_months":f"Limit the project period to {number:g} months." if number is not None else "",
        "submission_system":f"Submit through {manual or target}.",
    }
    return templates.get(rule.get("rule_type")) or str(rule.get("source_hint") or rule.get("notes") or target).strip()

def _compliance_rule_draft(rule):
    return {key:rule.get(key) for key in (
        "rule_id","category","rule_type","scope","target","severity","mandatory",
        "numeric_value","text_value","list_value","source_hint","source_document_hint",
        "source_page_hint","notes",
    )}

def compliance_rule_rows(profile):
    return [[compliance_rule_text(rule)] for rule in (profile or {}).get("rules") or []]

def compliance_rule_choices(profile):
    return [(compliance_rule_text(rule),rule.get("rule_id")) for rule in (profile or {}).get("rules") or [] if rule.get("rule_id")]

def compliance_provenance_rows(profile):
    return [[r.get("rule_id"),r.get("source_status"),r.get("source_document_id"),r.get("source_page"),r.get("source_start_offset"),r.get("source_end_offset"),r.get("source_excerpt")] for r in (profile or {}).get("rules") or []]

def build_compliance_profile(current_profile,sponsor,mechanism,submission_system,deadline,table):
    rules=[]
    prior_by_text={};used_ids=set()
    for rule in (current_profile or {}).get("rules") or []:
        prior_by_text.setdefault(compliance_rule_text(rule),[]).append(rule)
        used_ids.add(str(rule.get("rule_id") or ""))
    for r in _records(table,COMPLIANCE_HEADERS):
        text=_cell(r.get("Rule"))
        if not text:continue
        matches=prior_by_text.get(text) or []
        if matches:rules.append(_compliance_rule_draft(matches.pop(0)));continue
        rule_id=_stable_human_id("RULE",text,used_ids)
        rules.append({"rule_id":rule_id,"category":"administrative","rule_type":"manual_requirement","scope":"proposal","target":"proposal","severity":"warning","mandatory":False,"numeric_value":None,"text_value":text,"list_value":[],"source_hint":text,"source_document_hint":None,"source_page_hint":None,"notes":"Human-authored rule"})
    return {"sponsor":_cell(sponsor) or None,"mechanism":_cell(mechanism) or None,"submission_system":_cell(submission_system) or None,"deadline_iso":_cell(deadline) or None,"rules":rules}

def compliance_finding_rows(a):
    return [[x.get("rule_id"),x.get("severity"),x.get("mandatory"),x.get("status"),x.get("rule_type"),x.get("target"),x.get("detail"),x.get("source_excerpt")] for x in (a or {}).get("findings") or []]

def load_compliance(project):
    if not project:raise gr.Error("Create or open a project first.")
    d=api("GET",f"/api/projects/{project}/compliance")
    profile=d.get("profile") or {}
    assessment=api("GET",f"/api/projects/{project}/compliance/assessment")
    source=api("GET",f"/api/projects/{project}/opportunity-source")
    artifacts=api("GET",f"/api/projects/{project}/submission-artifacts")
    status=(f"Compliance profile v{d.get('version')} · {'approved' if d.get('approved') else 'awaiting human approval'} · {'fresh' if d.get('fresh') else 'STALE — recompile from the current opportunity source'}. "
            f"Hard failures: {assessment.get('hard_failures',0)}.")
    return profile,source.get("text") or "",profile.get("sponsor") or "",profile.get("mechanism") or "",profile.get("submission_system") or "",profile.get("deadline_iso") or "",compliance_rule_rows(profile),gr.update(choices=compliance_rule_choices(profile),value=None),compliance_provenance_rows(profile),compliance_finding_rows(assessment),assessment,[[x.get("slot"),x.get("filename"),x.get("extension"),x.get("sha256")] for x in artifacts],status

def compile_compliance(project):
    d=api("POST",f"/api/projects/{project}/compliance/compile",timeout=900);profile=d.get("profile") or {};a=api("GET",f"/api/projects/{project}/compliance/assessment");source=api("GET",f"/api/projects/{project}/opportunity-source")
    sections=api("GET",f"/api/projects/{project}/sections");choices=[x.get("title") for x in sections if x.get("title")]
    missing=sum(1 for r in profile.get("rules") or [] if r.get("source_status")!="located")
    return profile,source.get("text") or "",profile.get("sponsor") or "",profile.get("mechanism") or "",profile.get("submission_system") or "",profile.get("deadline_iso") or "",compliance_rule_rows(profile),gr.update(choices=compliance_rule_choices(profile),value=None),compliance_provenance_rows(profile),compliance_finding_rows(a),a,f"Recompiled {len(profile.get('rules') or [])} sponsor rules; {missing} require source-location review. Human approval is required.",gr.update(choices=choices)

def save_compliance(project,current_profile,sponsor,mechanism,submission_system,deadline,table):
    profile=build_compliance_profile(current_profile,sponsor,mechanism,submission_system,deadline,table)
    d=api("POST",f"/api/projects/{project}/compliance",json={"profile":profile},timeout=300);a=api("GET",f"/api/projects/{project}/compliance/assessment")
    sections=api("GET",f"/api/projects/{project}/sections");choices=[x.get("title") for x in sections if x.get("title")]
    saved=d.get("profile") or profile
    return saved,compliance_rule_rows(saved),gr.update(choices=compliance_rule_choices(saved),value=None),compliance_provenance_rows(saved),compliance_finding_rows(a),a,f"Saved human-reviewed compliance profile v{d.get('version')}. Approval is still required.",gr.update(choices=choices)

def approve_compliance(project):
    d=api("POST",f"/api/projects/{project}/compliance/approve");a=api("GET",f"/api/projects/{project}/compliance/assessment")
    return compliance_provenance_rows(d.get("profile") or {}),compliance_finding_rows(a),a,f"✓ Human approved sponsor compliance profile v{d.get('version')}. Deterministic hard-rule failures remaining: {a.get('hard_failures',0)}."

def resolve_compliance(project,rule_id,status,notes,resolved_by):
    if not (rule_id or "").strip():raise gr.Error("Choose the rule to resolve.")
    a=api("POST",f"/api/projects/{project}/compliance/resolve",json={"rule_id":rule_id.strip(),"status":status,"notes":notes or None,"resolved_by":resolved_by or None})
    return compliance_finding_rows(a),a,f"Saved manual resolution `{status}` for {rule_id}."

def register_submission_artifacts(project,slot,files):
    if not project:raise gr.Error("Open a project first.")
    slot=slug(slot)
    if not slot:raise gr.Error("Enter a submission slot, e.g. letters_of_support or biosketches.")
    dest=WORKSPACE/"projects"/project/"submission"/slot;dest.mkdir(parents=True,exist_ok=True)
    for item in files or []:
        src=Path(file_path(item));digest=hashlib.sha256(src.read_bytes()).hexdigest();target=dest/src.name
        if target.exists() and hashlib.sha256(target.read_bytes()).hexdigest()!=digest:target=dest/f"{src.stem}_{digest[:12]}{src.suffix.lower()}"
        if not target.exists():shutil.copy2(src,target)
        api("POST",f"/api/projects/{project}/submission-artifacts",json={"slot":slot,"filename":target.name,"path":str(target),"sha256":digest,"extension":target.suffix.lower().lstrip('.')})
    artifacts=api("GET",f"/api/projects/{project}/submission-artifacts");a=api("GET",f"/api/projects/{project}/compliance/assessment")
    return [[x.get("slot"),x.get("filename"),x.get("extension"),x.get("sha256")] for x in artifacts],compliance_finding_rows(a),a,f"Registered {len(files or [])} submission artifact(s) under `{slot}`."

def refresh_compliance_measurements(project):
    d=api("GET",f"/api/projects/{project}/approved-document",timeout=2400)
    sections=d.get("sections") or []
    if not sections:return api("GET",f"/api/projects/{project}/compliance/assessment")
    meta=d.get("project") or {}
    m=renderer_api("/measure",{"project_id":project,"title":meta.get("title") or "Grant Application","sponsor":meta.get("sponsor"),"organization_name":ORGANIZATION_NAME,"sections":[{"section_key":x.get("section_key"),"title":x.get("title") or "Section","body":x.get("body") or "","version":x.get("version")} for x in sections],"include_document_title":True,"design_profile":d.get("design_profile")},timeout=240)
    return api("POST",f"/api/projects/{project}/compliance/measurements",json={"measurements":m})

def measure_compliance(project):
    a=refresh_compliance_measurements(project)
    return compliance_finding_rows(a),a,f"Rendered compliance preflight: {a.get('passed',0)} passed, {a.get('hard_failures',0)} hard failure(s), {a.get('deferred',0)} deferred."

def draft_section(project,project_title,section,additional_context,high_value):
    key=slug(section);d=api("POST",f"/api/projects/{project}/draft-section",json={"section_key":key,"title":section,"additional_context":additional_context or None,"high_value":bool(high_value)},timeout=2400)
    state=api("GET",f"/api/projects/{project}/sections/{key}",timeout=2400);latest=state.get("latest") or {};update=state.get("competitive_update") or None
    body=latest.get("body",d["text"]);version=latest.get("version",d["version"])
    return version,body,key,section_preview(project,project_title,section,body,key,version,update),f"Draft version {version} created with `{d['model']}`. Click ✎ to edit or ✓ to approve this exact version.",gr.update(value=body,visible=False),gr.update(visible=False),gr.update(visible=False),(update or {}).get("event_id"),competitive_update_banner(update,version)

def show_editor(baseline):return gr.update(value=baseline or "",visible=True),gr.update(visible=True),gr.update(visible=True)

def save_edit(project,project_title,section,key,current_version,baseline,body):
    assert_section_identity(section,key)
    body=body or ""
    if body==baseline:
        state=api("GET",f"/api/projects/{project}/sections/{key}",timeout=2400);update=state.get("competitive_update") or None
        return current_version,baseline,section_preview(project,project_title,section,baseline,key,current_version,update),f"No changes detected; version {current_version} remains current.",gr.update(value=baseline,visible=False),gr.update(visible=False),gr.update(visible=False),(update or {}).get("event_id"),competitive_update_banner(update,current_version)
    try:
        d=api("POST",f"/api/projects/{project}/sections/{key}",json={"title":section,"body":body,"html":None,"base_version_id":int(current_version) if current_version else None},timeout=2400)
    except Exception:
        state=api("GET",f"/api/projects/{project}/sections/{key}",timeout=2400);latest=state.get("latest") or {};latest_version=latest.get("version")
        if not latest_version or int(latest_version)==int(current_version or 0):raise
        merge=api("POST",f"/api/projects/{project}/sections/{key}/merge-preview",json={"base_version_id":int(current_version),"latest_version_id":int(latest_version),"proposed_body":body},timeout=2400)
        merged=merge.get("merged_body") or "";conflicts=merge.get("conflicts") or [];update=state.get("competitive_update") or None
        status=(f"The section changed to version **{latest_version}** while you were editing. "
                + ("A clean three-way merge is ready below; review and save it as a new version." if not conflicts else f"The merge preserved both sides with **{len(conflicts)} conflict marker block(s)**. Resolve every marker before saving."))
        return latest_version,latest.get("body") or "",section_preview(project,project_title,section,merged,key,latest_version,update),status,gr.update(value=merged,visible=True),gr.update(visible=True),gr.update(visible=True),(update or {}).get("event_id"),competitive_update_banner(update,latest_version)
    v=d["version"]
    state=api("GET",f"/api/projects/{project}/sections/{key}",timeout=2400);update=state.get("competitive_update") or None
    return v,body,section_preview(project,project_title,section,body,key,v,update),f"Saved human edit as version {v}; approval is still required.",gr.update(value=body,visible=False),gr.update(visible=False),gr.update(visible=False),(update or {}).get("event_id"),competitive_update_banner(update,v)

def cancel_edit(project,project_title,section,key,current_version,baseline):
    assert_section_identity(section,key)
    state=api("GET",f"/api/projects/{project}/sections/{key}",timeout=2400);update=state.get("competitive_update") or None
    return section_preview(project,project_title,section,baseline,key,current_version,update),gr.update(value=baseline,visible=False),gr.update(visible=False),gr.update(visible=False),"Edit canceled; the persisted version was restored."

def approve_section(project,project_title,section,key,current_version,baseline,editor_body,competitive_update_event_id):
    assert_section_identity(section,key)
    if not current_version:raise gr.Error("Generate or save a section version before approval.")
    body=editor_body if editor_body is not None else baseline;version=current_version
    if body!=baseline:
        try:saved=api("POST",f"/api/projects/{project}/sections/{key}",json={"title":section,"body":body,"html":None,"base_version_id":int(current_version)},timeout=2400)
        except Exception:
            state=api("GET",f"/api/projects/{project}/sections/{key}",timeout=2400);latest=state.get("latest") or {};latest_version=latest.get("version")
            if not latest_version or int(latest_version)==int(current_version):raise
            merge=api("POST",f"/api/projects/{project}/sections/{key}/merge-preview",json={"base_version_id":int(current_version),"latest_version_id":int(latest_version),"proposed_body":body},timeout=2400)
            merged=merge.get("merged_body") or "";conflicts=merge.get("conflicts") or [];update=state.get("competitive_update") or None
            status=(f"Approval stopped because version **{latest_version}** was saved while you were editing. "
                    + ("A clean merge is ready for human review; save it before approving." if not conflicts else f"Resolve the **{len(conflicts)} conflict marker block(s)** below, then save and approve."))
            return latest_version,latest.get("body") or "",section_preview(project,project_title,section,merged,key,latest_version,update),status,gr.update(value=merged,visible=True),gr.update(visible=True),gr.update(visible=True),(update or {}).get("event_id"),competitive_update_banner(update,latest_version)
        version=saved["version"];baseline=body
    d=api("POST",f"/api/projects/{project}/sections/{key}/approve",json={"version_id":int(version),"competitive_update_event_id":int(competitive_update_event_id) if competitive_update_event_id else None},timeout=2400)
    if not d.get("approved"):
        progress=d.get("approval_progress") or {}
        return version,baseline,section_preview(project,project_title,section,baseline,key,version,None),f"Approval recorded ({progress.get('approvals',0)} of {progress.get('minimum_approvals','?')}); this exact version remains unapproved until the configured threshold is met.",gr.update(value=baseline,visible=False),gr.update(visible=False),gr.update(visible=False),None,""
    return version,baseline,section_preview(project,project_title,section,baseline,key,version,None),f"✓ Human approved exact version {d['approved_version']} for {section}. Workflow stage: `{d['stage']}`.",gr.update(value=baseline,visible=False),gr.update(visible=False),gr.update(visible=False),None,""

def return_section_for_revision(project,key,current_version,rationale):
    if not key or not current_version:raise gr.Error("Load a saved proposal section before returning it for revision.")
    rationale=str(rationale or "").strip()
    if not rationale:raise gr.Error("Explain what must change before returning the section for revision.")
    api("POST",f"/api/projects/{project}/sections/{key}/return-for-revision",json={"version":int(current_version),"rationale":rationale})
    status=api("GET",f"/api/projects/{project}/workflow/status")
    return f"↩ Returned exact section version {current_version} for revision. Contributors may now publish a corrected version; all prior work remains in history.",status,""

def preview_approved_grant(project):
    if not project:raise gr.Error("Create or open a project first.")
    d=api("GET",f"/api/projects/{project}/approved-document")
    plan=d.get("section_plan") or []
    approved_by_key={x.get("section_key"):x for x in d.get("sections") or []}
    rows=[]
    for item in plan:
        approved=approved_by_key.get(item.get("section_key"))
        rows.append([
            item.get("position"),
            item.get("title"),
            "✓ Approved" if approved else "Not approved",
            approved.get("version") if approved else None,
        ])
    counts=d.get("counts") or {}
    status=f"Aggregating **{counts.get('approved',0)} / {counts.get('configured',0)}** configured sections from human-approved versions only."
    if d.get("readiness",{}).get("ready"):
        status += " All required sections are approved; the final DOCX/PDF choice is available."
    return rows,full_document_preview(project,d),status

def readiness(project):
    try:refresh_compliance_measurements(project)
    except Exception:pass
    d=api("GET",f"/api/projects/{project}/readiness",timeout=2400)
    ready=bool(d.get("ready"))
    if ready:
        status="✓ All workflow gates passed. Choose the final format below."
    elif int(d.get("competitive_text_updates_pending") or 0)>0:
        status=f"⚡ **{COMPETITIVE_UPDATE_LABEL}:** {d.get('competitive_text_updates_pending')} refreshed section(s) are waiting for human review. Open the highlighted updates, edit if needed, and approve before export."
    elif int(d.get("competitive_refresh_processing_pending") or 0)>0:
        status="⚡ Competitive intelligence found new public information and is still self-healing affected section text. Retry shortly; export remains fail-closed until the update completes."
    elif not d.get("sponsor_compliance_ready"):
        status=f"Sponsor compliance is not ready: {d.get('sponsor_compliance_hard_failures')} hard rule failure(s). Open **Sponsor Compliance & Submission**, resolve the highlighted rules, then rerun readiness."
    else:
        status="Not export-ready: "+json.dumps(d,indent=2)
    return d,status,gr.update(visible=ready),gr.update(visible=ready)

def system_diagnostics():
    info=api("GET","/api/system/info",timeout=60)
    summary=(f"**Build:** {info.get('build_version')} · **Runtime:** {info.get('runtime_profile')} · "
             f"**Model routing:** {info.get('model_routing_mode')} · **OMP:** {info.get('omp_threads')} · "
             f"**Rayon:** {info.get('rayon_threads')} · **DB:** {int(info.get('database_bytes') or 0)/1024/1024:.1f} MB")
    return info,summary

def run_hpc_diagnostics():
    data=api("POST","/api/hpc/benchmark",timeout=600)
    summary=(f"HPC benchmark completed: normalize {float(data.get('normalize_ms') or 0):.1f} ms · "
             f"SGEMV {float(data.get('sgemv_ms') or 0):.1f} ms · mmap open {float(data.get('mmap_open_ms') or 0):.2f} ms · "
             f"mmap score {float(data.get('mmap_score_ms') or 0):.1f} ms.")
    return data,summary

def export(project,fmt,_title):
    refresh_compliance_measurements(project)
    snap=api("POST",f"/api/projects/{project}/export-snapshot");meta=snap["project"];sections=snap["sections"];fmts=["docx","pdf"] if fmt=="BOTH" else [fmt.lower()];paths=[]
    final_title=meta.get("title") or "Grant Application"
    for f in fmts:
        payload={"project_id":project,"snapshot_id":snap["snapshot_id"],"format":f,"title":final_title,"sponsor":meta.get("sponsor"),"organization_name":ORGANIZATION_NAME,"sections":[{"section_key":x.get("section_key"),"title":x["title"],"body":x["body"],"version":x.get("version")} for x in sections],"include_document_title":True,"design_profile":snap.get("design_profile")}
        paths.append(renderer_api("/render",payload,timeout=240)["path"])
    package=renderer_api("/package",{"project_id":project,"snapshot_id":snap["snapshot_id"],"title":final_title,"generated_paths":paths,"manifest":{"snapshot_sha256":snap.get("sha256"),"sponsor_compliance_profile":snap.get("sponsor_compliance_profile"),"sponsor_compliance_assessment":snap.get("sponsor_compliance_assessment"),"submission_artifacts":snap.get("submission_artifacts")}},timeout=240)["path"]
    paths.append(package)
    return paths,f"Created from immutable export snapshot {snap['snapshot_id']} ({snap['sha256'][:16]}…). Sponsor-compliant package: {package}"

def editor_outline_html(sections,selected_key=None):
    if not sections:return '<div class="editor-empty">No sections yet. Add the first section below.</div>'
    items=[]
    for section in sections:
        key=str(section.get("section_key") or "");title=str(section.get("title") or "Untitled section");description=str(section.get("description") or "")
        active=" active" if key==selected_key else ""
        items.append(f'''<div class="editor-outline-item{active}" draggable="true" data-section-key="{html.escape(key,quote=True)}">
          <span class="drag" title="Drag to reorder">⋮⋮</span><div><b>{html.escape(title)}</b><div class="editor-outline-description">{html.escape(description)}</div></div>
          <span class="outline-actions"><button type="button" data-outline-action="up" title="Move up">↑</button><button type="button" data-outline-action="down" title="Move down">↓</button></span></div>''')
    return '<div class="editor-outline"><div class="editor-outline-head"><b>Grant sections</b><span>drag to reorder</span></div>'+"".join(items)+"</div>"

def editor_guidance_html(comments):
    if not comments:return '<div class="editor-secondary">No team guidance yet. Add a comment, question, or rule for this section.</div>'
    cards=[]
    for comment in comments:
        raw=str(comment.get("body") or "");kind="COMMENT";text=raw
        match=re.match(r"^(COMMENT|QUESTION|RULE)\s*[·:]\s*(.*)$",raw,flags=re.I|re.S)
        if match:kind=match.group(1).upper();text=match.group(2)
        resolved=bool(comment.get("resolved_at"));classes="guidance-card guidance-resolved" if resolved else "guidance-card"
        state=" · resolved" if resolved else ""
        cards.append(f'<article class="{classes}"><div class="guidance-type">{kind}</div><div class="guidance-meta">{html.escape(str(comment.get("author") or "Team member"))} · {html.escape(str(comment.get("created_at") or ""))}{state}</div><div>{html.escape(text).replace(chr(10),"<br>")}</div></article>')
    return "".join(cards)

def editor_comment_state(project,key):
    if not project or not key:return '<div class="editor-secondary">Choose a section to view team guidance.</div>',gr.update(choices=[],value=None)
    comments=api("GET",f"/api/projects/{project}/comments/section/{key}")
    open_comments=[comment for comment in comments if not comment.get("resolved_at")]
    choices=[(f"{str(comment.get('body') or '')[:70]} · {comment.get('author')}",comment.get("id")) for comment in open_comments]
    return editor_guidance_html(comments),gr.update(choices=choices,value=None)

def editor_document_html(sections,selected_key=None,override=None):
    if not sections:return '<div class="editor-empty">No sections yet. Add the first section from the outline.</div>'
    override=override or {};pages=[]
    for section in sections:
        key=str(section.get("section_key") or "");display=dict(section);dirty=False
        if key==selected_key and override:
            display.update({name:value for name,value in override.items() if name in {"title","description","body"}});dirty=True
        title=html.escape(str(display.get("title") or "Untitled section"));description=html.escape(str(display.get("description") or ""));body=html.escape(str(display.get("body") or ""))
        version=display.get("version");selected=" selected" if key==selected_key else "";unsaved=" unsaved" if dirty else ""
        pages.append(f'''<article id="grant-doc-{html.escape(key,quote=True)}" class="grant-doc-section{selected}{unsaved}" data-section-key="{html.escape(key,quote=True)}" data-version="{version if version is not None else ''}" data-dirty="{'true' if dirty else 'false'}">
          <div class="grant-doc-version">{f'Version {version}' if version else 'New section'}<span class="unsaved-label">Unsaved changes</span></div>
          <h2 contenteditable="true" spellcheck="true" data-editor-field="title">{title}</h2>
          <div class="grant-doc-description" contenteditable="true" spellcheck="true" data-editor-field="description">{description}</div>
          <div class="grant-doc-body" contenteditable="true" spellcheck="true" data-editor-field="body">{body.replace(chr(10),'<br>')}</div>
        </article>''')
    return '<div class="grant-document-scroll">'+"".join(pages)+"</div>"

def selected_editor_controls(project,key):
    if not project or not key:return None,None,gr.update(choices=[],value=None),*editor_comment_state(project,None),"Choose a section from the outline."
    state=api("GET",f"/api/projects/{project}/sections/{key}");latest=state.get("latest") or {};version=latest.get("version")
    versions=api("GET",f"/api/projects/{project}/sections/{key}/versions")
    choices=[(f"v{item.get('version')} · {item.get('editor') or item.get('source')} · {item.get('created_at')}",item.get("version")) for item in versions]
    guidance_html,guidance_choices=editor_comment_state(project,key)
    status=f"Section selected · latest saved version {version}." if version else "Section selected · not saved yet. Use Rewrite selected or write directly in the document."
    return key,version,gr.update(choices=choices,value=version),guidance_html,guidance_choices,status

def load_continuous_editor(project,requested_key=None,override=None,status=None):
    if not (project or "").strip():
        return editor_outline_html([],None),editor_document_html([],None),None,None,gr.update(choices=[],value=None),'<div class="editor-secondary">Open a saved grant to view guidance.</div>',gr.update(choices=[],value=None),"Open a saved grant to begin."
    ensure_document_editor(project);document=api("GET",f"/api/projects/{project}/editor/document");sections=document.get("sections") or []
    keys=[str(section.get("section_key")) for section in sections];key=str(requested_key or "") if str(requested_key or "") in keys else (keys[0] if keys else None)
    controls=selected_editor_controls(project,key)
    return editor_outline_html(sections,key),editor_document_html(sections,key,override),*controls[:-1],status or controls[-1]

def prefill_continuous_editor(project,requested_key=None,progress=gr.Progress()):
    if not (project or "").strip():return load_continuous_editor(project,requested_key)
    try:
        ensure_document_editor(project);document=api("GET",f"/api/projects/{project}/editor/document");sections=document.get("sections") or []
        missing=[section for section in sections if not str(section.get("body") or "").strip()]
        if not missing:return load_continuous_editor(project,requested_key,status="The AI-prepared grant draft is loaded. Edit any section directly or add team guidance.")
        failures=[];total=len(missing);drafted_chunks=0
        for index,section in enumerate(missing,1):
            title=section.get("title") or "Untitled section"
            progress((index-1)/total,desc=f"Drafting {index} of {total} · {title} · compiling grant sources and evidence")
            try:
                result=api("POST",f"/api/projects/{project}/sections/{section.get('section_key')}/rewrite",json={"title":title,"description":section.get("description") or "","base_version_id":section.get("version"),"high_value":False},timeout=2400)
                drafted_chunks+=int((result.get("chunking") or {}).get("section_chunks") or 1)
            except Exception as error:failures.append(f"{title}: {error}")
        progress(1.0,desc="Assembling the complete editable grant document")
        note=f"AI prepared {total-len(failures)} of {total} missing section(s) in {drafted_chunks} bounded drafting chunk(s), then assembled and saved each complete section from the saved grant ask, retrieved evidence, and configured workflow. Review and edit every claim before publishing."
        if failures:note+=f" Drafting paused for {len(failures)} section(s): {'; '.join(failures)}"
        return load_continuous_editor(project,requested_key,status=note)
    except Exception as error:
        # Opening a project must never fail merely because model generation is unavailable.
        try:return load_continuous_editor(project,requested_key,status=f"The saved document opened, but automatic drafting could not start: {error}. Use Draft missing sections to retry.")
        except Exception:return load_continuous_editor(None,None)

def _editor_payload_sections(raw):
    try:payload=json.loads(raw or "{}")
    except Exception:raise gr.Error("The browser could not serialize the document. Your visible edits remain in place; retry Save document.")
    sections=payload.get("sections") if isinstance(payload,dict) else None
    if not isinstance(sections,list):raise gr.Error("The browser returned an invalid document payload. Your visible edits remain in place.")
    return sections

def _save_continuous_changes(project,raw):
    sections=_editor_payload_sections(raw)
    if not sections:return []
    response=api("POST",f"/api/projects/{project}/editor/document",json={"sections":sections},timeout=2400)
    return response.get("saved") or []

def save_continuous_editor(project,raw,current_key):
    project=_require_project(project)
    try:saved=_save_continuous_changes(project,raw)
    except Exception as error:raise gr.Error(f"The document was not saved because a teammate changed one of these sections. Your edits remain in this browser. Refresh only after copying or reconciling them. {error}")
    note=(f"Saved {len(saved)} changed section(s) as one atomic document update. Teammates can use Refresh shared changes." if saved else "No unsaved document changes were detected.")
    return load_continuous_editor(project,current_key,status=note)

def rewrite_continuous_section(project,key,raw):
    project=_require_project(project);key=str(key or "").strip()
    if not key:raise gr.Error("Choose a section from the outline first.")
    _save_continuous_changes(project,raw)
    document=api("GET",f"/api/projects/{project}/editor/document");section=next((item for item in document.get("sections") or [] if str(item.get("section_key"))==key),None)
    if not section:raise gr.Error("The selected section no longer exists. Refresh shared changes.")
    result=api("POST",f"/api/projects/{project}/sections/{key}/rewrite",json={"title":section.get("title") or "Untitled section","description":section.get("description") or "","base_version_id":section.get("version"),"high_value":False},timeout=2400)
    preview=" ".join(str(result.get("text") or "").split()[:20]);research=result.get("research") or {};failures=research.get("failures") or [];chunking=result.get("chunking") or {}
    research_note=f"Research saved {research.get('sources_saved',0)} source(s)"+(f"; {' '.join(str(value) for value in failures)}" if failures else "")
    chunk_note=(f" Ollama used {chunking.get('section_chunks',1)} section chunk(s) and {chunking.get('context_reduction_rounds',0)} reduction round(s)." if chunking.get("enabled") else "")
    note=f"Rewrote {section.get('title')} as version {result.get('version')} with {result.get('model')}. {research_note}.{chunk_note} Preview: {preview}{'…' if len(str(result.get('text') or '').split())>20 else ''}"
    return load_continuous_editor(project,key,status=note)

def load_continuous_historical_version(project,key,version):
    project=_require_project(project)
    if not key or not version:raise gr.Error("Choose a stored version first.")
    item=api("GET",f"/api/projects/{project}/sections/{key}/versions/{int(version)}")
    document=api("GET",f"/api/projects/{project}/editor/document");sections=document.get("sections") or []
    note=f"Historical version {version} is now an unsaved working copy. Save document to make it the newest version; publishing is unchanged until then."
    return editor_document_html(sections,key,{"title":item.get("title") or "","body":item.get("body") or ""}),note

def add_continuous_editor_section(project,title,raw,current_key):
    project=_require_project(project);title=str(title or "").strip()
    if not title:raise gr.Error("Enter a section title.")
    _save_continuous_changes(project,raw)
    sections=ensure_document_editor(project);existing={str(section.get("section_key")) for section in sections};base=slug(title) or "section";key=base;counter=2
    while key in existing:key=f"{base}_{counter}";counter+=1
    api("POST",f"/api/projects/{project}/editor/initialize",json={"sections":[{"key":key,"title":title,"description":""}]})
    return (*load_continuous_editor(project,key,status=f"Added {title}. Write directly in the new page or choose Rewrite selected."),"")

def handle_continuous_outline_command(project,current_key,raw):
    if not raw:return current_key,None,gr.update(),gr.skip(),gr.update(),"",gr.update(value="")
    try:command=json.loads(raw)
    except Exception:raise gr.Error("The section outline command was invalid. Refresh shared changes and try again.")
    action=command.get("action");keys=[str(value) for value in command.get("keys") or []];selected=str(command.get("key") or current_key or "")
    if action=="reorder" and keys:api("POST",f"/api/projects/{project}/sections/reorder",json={"section_keys":keys})
    controls=selected_editor_controls(project,selected)
    status=("Section order saved for every collaborator." if action=="reorder" else controls[-1])
    return *controls[:-1],status,gr.update(value="")

def post_continuous_guidance(project,key,kind,body,raw):
    project=_require_project(project);key=str(key or "").strip();text=str(body or "").strip()
    if not key:raise gr.Error("Choose a section from the outline first.")
    if not text:raise gr.Error("Write a comment, question, or rule first.")
    _save_continuous_changes(project,raw)
    state=api("GET",f"/api/projects/{project}/sections/{key}");latest=state.get("latest") or {};version=latest.get("version")
    if not version:
        document=api("GET",f"/api/projects/{project}/editor/document");section=next(item for item in document.get("sections") or [] if str(item.get("section_key"))==key)
        api("POST",f"/api/projects/{project}/sections/{key}/rewrite",json={"title":section.get("title") or "Untitled section","description":section.get("description") or "","base_version_id":None,"high_value":False},timeout=2400)
        version=(api("GET",f"/api/projects/{project}/sections/{key}").get("latest") or {}).get("version")
    api("POST",f"/api/projects/{project}/comments/section/{key}",json={"version_id":int(version),"start_offset":None,"end_offset":None,"quoted_text":None,"body":f"{str(kind or 'COMMENT').upper()} · {text}","parent_comment_id":None,"mentioned_user_ids":[]})
    if str(kind or "").lower() in {"question","rule"}:
        loaded=rewrite_continuous_section(project,key,json.dumps({"sections":[]}));return (*loaded,"",f"The {kind} triggered evidence research and a new saved rewrite.")
    loaded=load_continuous_editor(project,key,status="Shared comment saved. It will remain visible until a teammate resolves it.")
    return (*loaded,"","Shared comment added without rewriting the section.")

def ensure_document_editor(project):
    project=_require_project(project);sections=api("GET",f"/api/projects/{project}/sections")
    if not sections:
        planned=api("POST",f"/api/projects/{project}/editor/plan",timeout=2400)
        sections=planned.get("sections") or []
    if not sections:raise gr.Error("No sponsor-derived document sections are available for this grant.")
    return sections

def load_document_editor(project,requested_key=None):
    if not (project or "").strip():
        return editor_outline_html([],None),None,"","","",None,"",gr.update(choices=[],value=None),'<div class="editor-secondary">Open a saved grant to begin editing.</div>',gr.update(choices=[],value=None),"Open a saved grant."
    sections=ensure_document_editor(project)
    keys=[str(section.get("section_key")) for section in sections]
    key=str(requested_key or "") if requested_key in keys else (keys[0] if keys else "")
    if not key:return editor_outline_html([],None),None,"","","",None,"",gr.update(choices=[],value=None),*editor_comment_state(project,None),"Add a section to begin."
    meta=next(section for section in sections if str(section.get("section_key"))==key)
    state=api("GET",f"/api/projects/{project}/sections/{key}");latest=state.get("latest") or {}
    body=latest.get("body") or "";version=latest.get("version")
    versions=api("GET",f"/api/projects/{project}/sections/{key}/versions")
    version_choices=[(f"v{item.get('version')} · {item.get('editor') or item.get('source')} · {item.get('created_at')}",item.get("version")) for item in versions]
    guidance_html,guidance_choices=editor_comment_state(project,key)
    status=(f"Loaded version {version}. Refresh before editing if a teammate may have changed this section." if version else "This section is ready for its first evidence-grounded draft. Choose Rewrite with team guidance to prefill it.")
    return editor_outline_html(sections,key),key,meta.get("title") or "",meta.get("description") or "",body,version,body,gr.update(choices=version_choices,value=version),guidance_html,guidance_choices,status

def prepare_document_editor(project,requested_key=None,progress=gr.Progress()):
    loaded=load_document_editor(project,requested_key)
    if project and loaded[1] and not loaded[5]:
        progress(0.08,desc=f"Researching the grant ask for {loaded[2]}")
        progress(0.28,desc="Compiling saved sources, evidence, workflow options, and team guidance")
        rewritten=rewrite_document_section(project,loaded[1],loaded[2],loaded[3],None)
        progress(0.96,desc="Saving the first collaborative section version")
        return rewritten
    return loaded

def load_editor_historical_version(project,key,version,current_version):
    project=_require_project(project)
    if not key or not version:raise gr.Error("Choose a stored version first.")
    item=api("GET",f"/api/projects/{project}/sections/{key}/versions/{int(version)}")
    note=(f"Loaded historical version {version} into the editable working copy. Your published grant is unchanged. "
          f"Select Save changes to preserve your modifications as a new version based on current head {current_version}.")
    return item.get("title") or "",item.get("body") or "",note

def save_document_section(project,key,title,description,base_version,body):
    project=_require_project(project);key=str(key or "").strip();title=str(title or "").strip();body=str(body or "")
    if not key:raise gr.Error("Choose a section first.")
    if not title:raise gr.Error("A section title is required.")
    if not body.strip():raise gr.Error("Section text cannot be empty. Use Rewrite with team guidance to create its first draft.")
    try:
        result=api("POST",f"/api/projects/{project}/sections/{key}",json={"title":title,"description":str(description or ""),"body":body,"html":None,"base_version_id":int(base_version) if base_version else None},timeout=2400)
    except Exception:
        latest=api("GET",f"/api/projects/{project}/sections/{key}").get("latest") or {}
        if latest.get("version") and int(latest.get("version"))!=int(base_version or 0):
            raise gr.Error(f"A teammate saved version {latest.get('version')} while you were editing. Your text remains in this browser. Copy it if needed, refresh the section, then reconcile the changes.")
        raise
    loaded=list(load_document_editor(project,key));loaded[-1]=f"Saved as collaborative version {result.get('version')}. Teammates will see it when they refresh."
    return tuple(loaded)

def rewrite_document_section(project,key,title,description,base_version):
    project=_require_project(project);key=str(key or "").strip()
    if not key:raise gr.Error("Choose a section first.")
    result=api("POST",f"/api/projects/{project}/sections/{key}/rewrite",json={"title":str(title or "").strip(),"description":str(description or ""),"base_version_id":int(base_version) if base_version else None,"high_value":False},timeout=2400)
    loaded=list(load_document_editor(project,key));preview=" ".join(str(result.get("text") or "").split()[:20]);research=result.get("research") or {};failures=research.get("failures") or []
    research_note=(f"saved {research.get('sources_saved',0)} new research source(s)" if not failures else f"saved {research.get('sources_saved',0)} new research source(s); {' '.join(str(value) for value in failures)}")
    chunking=result.get("chunking") or {};chunk_note=(f" Ollama processed {chunking.get('section_chunks',1)} section chunk(s) after {chunking.get('context_reduction_rounds',0)} context-reduction round(s)." if chunking.get("enabled") else "")
    loaded[-1]=f"Rewrote this section as version {result.get('version')} with {result.get('model')}. Applied {result.get('guidance_count',0)} open team item(s); {research_note}.{chunk_note} Preview: {preview}{'…' if len(str(result.get('text') or '').split())>20 else ''}"
    return tuple(loaded)

def post_editor_guidance(project,key,version,kind,body):
    project=_require_project(project);key=str(key or "").strip();text=str(body or "").strip()
    if not version:raise gr.Error("Create the section's first draft before adding version-anchored guidance.")
    if not text:raise gr.Error("Write a comment, question, or rule first.")
    api("POST",f"/api/projects/{project}/comments/section/{key}",json={"version_id":int(version),"start_offset":None,"end_offset":None,"quoted_text":None,"body":f"{str(kind or 'COMMENT').upper()} · {text}","parent_comment_id":None,"mentioned_user_ids":[]})
    guidance_html,guidance_choices=editor_comment_state(project,key)
    return guidance_html,guidance_choices,"",f"Shared {str(kind or 'comment').lower()} added. Use Rewrite with team guidance when the team is ready to apply it."

def post_editor_guidance_and_rewrite(project,key,version,kind,body,title,description):
    if not version:
        initial=rewrite_document_section(project,key,title,description,None)
        version=initial[5]
    post_editor_guidance(project,key,version,kind,body)
    if str(kind or "").lower() in {"question","rule"}:
        loaded=list(rewrite_document_section(project,key,title,description,version))
        research_status=loaded[-1]
        return (*loaded,"",f"The {kind} triggered evidence research and a new section rewrite. {research_status}")
    loaded=load_document_editor(project,key)
    return (*loaded,"",f"Shared comment added. It remains open until a teammate applies or resolves it.")

def resolve_editor_guidance(project,key,comment_id):
    project=_require_project(project)
    if not comment_id:raise gr.Error("Choose an open team item to resolve.")
    api("POST",f"/api/projects/{project}/comments/{int(comment_id)}/resolve",json={})
    guidance_html,guidance_choices=editor_comment_state(project,key)
    return guidance_html,guidance_choices,"Team item resolved; its audit history was preserved."

def add_editor_section(project,title):
    project=_require_project(project);title=str(title or "").strip()
    if not title:raise gr.Error("Enter a section title.")
    sections=ensure_document_editor(project);existing={str(section.get("section_key")) for section in sections};base=slug(title) or "section";key=base;counter=2
    while key in existing:key=f"{base}_{counter}";counter+=1
    api("POST",f"/api/projects/{project}/editor/initialize",json={"sections":[{"key":key,"title":title,"description":""}]})
    return (*load_document_editor(project,key),"")

def handle_editor_outline_command(project,current_key,raw):
    if not raw:return load_document_editor(project,current_key)
    try:command=json.loads(raw)
    except Exception:raise gr.Error("The section outline command was invalid. Refresh the editor and try again.")
    action=command.get("action");keys=[str(value) for value in command.get("keys") or []];selected=str(command.get("key") or current_key or "")
    if action=="reorder" and keys:api("POST",f"/api/projects/{project}/sections/reorder",json={"section_keys":keys})
    loaded=load_document_editor(project,selected)
    if action=="select" and loaded[1] and not loaded[5]:return rewrite_document_section(project,loaded[1],loaded[2],loaded[3],None)
    return loaded

def publish_document_editor(project):
    project=_require_project(project);response=api("POST",f"/api/projects/{project}/publish-snapshot",json={},timeout=300);snapshot=response.get("snapshot") or {};meta=snapshot.get("project") or {};sections=snapshot.get("sections") or [];paths=[]
    payload_base={"project_id":project,"snapshot_id":response["snapshot_id"],"title":meta.get("title") or "Grant Application","sponsor":meta.get("sponsor"),"organization_name":ORGANIZATION_NAME,"sections":[{"section_key":item.get("section_key"),"title":item.get("title"),"body":item.get("body"),"version":item.get("version")} for item in sections],"include_document_title":True,"design_profile":snapshot.get("design_profile")}
    for output_format in ("docx","pdf"):
        rendered=renderer_api("/render",{**payload_base,"format":output_format},timeout=300);paths.append(rendered["path"])
    return paths,f"Published exact collaborative snapshot {response['snapshot_id']} as DOCX and PDF. Snapshot SHA-256: {response['sha256']}."

EDITOR_OUTLINE_JS="""() => {
  if (window.__grantspaceEditorOutlineBound) return [];
  window.__grantspaceEditorOutlineBound = true;
  let dragged = null;
  window.grantspaceCollectEditorDocument = () => ({sections:[...document.querySelectorAll('.grant-doc-section[data-dirty="true"]')].map(section => ({
    key:section.dataset.sectionKey,
    title:(section.querySelector('[data-editor-field="title"]')?.innerText || '').trim(),
    description:(section.querySelector('[data-editor-field="description"]')?.innerText || '').trim(),
    body:section.querySelector('[data-editor-field="body"]')?.innerText || '',
    base_version_id:section.dataset.version ? Number(section.dataset.version) : null
  }))});
  const emit = (action, key) => {
    const items = [...document.querySelectorAll('.editor-outline-item')];
    const host = document.querySelector('#editor-outline-command textarea, #editor-outline-command input');
    if (!host) return;
    const value = JSON.stringify({action, key, keys: items.map(item => item.dataset.sectionKey)});
    const prototype = host.tagName === 'TEXTAREA' ? HTMLTextAreaElement.prototype : HTMLInputElement.prototype;
    const setter = Object.getOwnPropertyDescriptor(prototype, 'value')?.set;
    if (setter) setter.call(host, value); else host.value = value;
    host.dispatchEvent(new Event('input', {bubbles:true}));
    host.dispatchEvent(new Event('change', {bubbles:true}));
  };
  document.addEventListener('dragstart', event => { const item=event.target.closest('.editor-outline-item'); if(!item)return; dragged=item;item.classList.add('dragging');event.dataTransfer.effectAllowed='move'; });
  document.addEventListener('dragover', event => { const item=event.target.closest('.editor-outline-item'); if(!dragged||!item||item===dragged)return;event.preventDefault();const box=item.getBoundingClientRect();item.parentElement.insertBefore(dragged,event.clientY<box.top+box.height/2?item:item.nextSibling); });
  const reorderDocument = () => { const host=document.querySelector('.grant-document-scroll');if(!host)return;[...document.querySelectorAll('.editor-outline-item')].forEach(item=>{const page=document.getElementById(`grant-doc-${item.dataset.sectionKey}`);if(page)host.appendChild(page);}); };
  document.addEventListener('dragend', event => { if(!dragged)return;const key=dragged.dataset.sectionKey;dragged.classList.remove('dragging');dragged=null;reorderDocument();emit('reorder',key); });
  document.addEventListener('input', event => { const field=event.target.closest('[contenteditable="true"][data-editor-field]');if(!field)return;const section=field.closest('.grant-doc-section');if(section){section.dataset.dirty='true';section.classList.add('unsaved');} });
  document.addEventListener('click', event => { const actionButton=event.target.closest('[data-outline-action]');const item=event.target.closest('.editor-outline-item');if(!item)return;event.preventDefault();document.querySelectorAll('.editor-outline-item,.grant-doc-section').forEach(node=>node.classList.remove('active','selected'));item.classList.add('active');const page=document.getElementById(`grant-doc-${item.dataset.sectionKey}`);if(page)page.classList.add('selected');if(actionButton){const direction=actionButton.dataset.outlineAction;const sibling=direction==='up'?item.previousElementSibling:item.nextElementSibling;if(sibling&&sibling.classList.contains('editor-outline-item'))item.parentElement.insertBefore(direction==='up'?item:sibling,direction==='up'?sibling:item);reorderDocument();emit('reorder',item.dataset.sectionKey);}else{page?.scrollIntoView({behavior:'smooth',block:'start'});emit('select',item.dataset.sectionKey);} });
  return [];
}"""

EDITOR_SAVE_INPUT_JS="""(project,payload,currentKey) => [project,JSON.stringify((window.grantspaceCollectEditorDocument||(()=>({sections:[]})))()),currentKey]"""
EDITOR_REWRITE_INPUT_JS="""(project,currentKey,payload) => [project,currentKey,JSON.stringify((window.grantspaceCollectEditorDocument||(()=>({sections:[]})))())]"""
EDITOR_GUIDANCE_INPUT_JS="""(project,currentKey,kind,body,payload) => [project,currentKey,kind,body,JSON.stringify((window.grantspaceCollectEditorDocument||(()=>({sections:[]})))())]"""
EDITOR_ADD_INPUT_JS="""(project,title,payload,currentKey) => [project,title,JSON.stringify((window.grantspaceCollectEditorDocument||(()=>({sections:[]})))()),currentKey]"""
GLOBAL_NAVIGATION_JS="""() => {
  if (window.__grantspaceGlobalNavigationBound) return [];
  window.__grantspaceGlobalNavigationBound=true;
  document.addEventListener('click',event=>{
    const control=event.target.closest('[data-global-nav]');if(!control)return;
    const action=control.dataset.globalNav;
    if(action==='editor'||action==='admin'){
      const overlay=document.getElementById('wizard-overlay');if(overlay){overlay.style.setProperty('display','none','important');overlay.setAttribute('aria-hidden','true');}
      const wanted=action==='editor'?'Grant editor':'System administration';
      const tab=[...document.querySelectorAll('button[role="tab"]')].find(button=>button.textContent.trim()===wanted);
      if(tab)tab.click();return;
    }
    const overlay=document.getElementById('wizard-overlay');if(overlay){overlay.style.setProperty('display','block','important');overlay.setAttribute('aria-hidden','false');}
    const host=document.querySelector('#global-navigation-command textarea,#global-navigation-command input');if(!host)return;
    const prototype=host.tagName==='TEXTAREA'?HTMLTextAreaElement.prototype:HTMLInputElement.prototype;
    const setter=Object.getOwnPropertyDescriptor(prototype,'value')?.set;if(setter)setter.call(host,action);else host.value=action;
    host.dispatchEvent(new Event('input',{bubbles:true}));host.dispatchEvent(new Event('change',{bubbles:true}));
  });
  return [];
}"""

with gr.Blocks(title="Grantspace") as demo:
    with gr.Column(visible=True,elem_classes="wizard-lightbox",elem_id="wizard-overlay") as wizard_shell:
        gr.HTML('''<nav class="global-nav" aria-label="Global navigation"><span class="brand">Grantspace</span><button type="button" data-global-nav="projects">Grants &amp; wizard</button><button type="button" data-global-nav="workflow">Workflow setup</button><button type="button" data-global-nav="editor">Grant editor</button><button type="button" data-global-nav="admin">Administration</button><a href="/logout">Sign out</a></nav>''',container=False)
        with gr.Row(elem_classes="wizard-shell"):
            wizard_rail=gr.HTML(wizard_rail_html(1),container=False)
            with gr.Column(elem_classes="wizard-main"):
                wizard_preset=gr.State("custom_configuration_v1")
                wizard_edit_project=gr.State(None)
                wizard_modules=gr.State([])
                wizard_required_modules=gr.State([])
                wizard_review_required=gr.State(False)
                with gr.Column(visible=True,elem_id="wizard-page-1") as wizard_page_1:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Composable grant workflow</div><h1>Start with the grant.<br><span class="accent">Build only what you need.</span></h1><p>Grantspace turns each solicitation into a shared, approval-ready workflow grounded in evidence, owned by your team, and reproducible at export.</p></div>')
                    with gr.Row():
                        with gr.Column(scale=3,elem_classes=["wizard-panel","wizard-hero"]):
                            gr.Markdown("## Create a new shared grant\nCompose a workflow from a new RFA, RFI, NOFO, or application form.")
                            wizard_new_btn=gr.Button("Start configuration →",variant="primary")
                        with gr.Column(scale=2):
                            with gr.Group(elem_classes="wizard-panel"):
                                gr.Markdown("### Open an existing grant\nRestore its persisted workflow, files, versions, and decisions for this teammate.")
                                wizard_existing=gr.Dropdown(label="Shared grant",choices=[])
                                with gr.Row():
                                    wizard_refresh_btn=gr.Button("Refresh")
                                    wizard_open_btn=gr.Button("Open",variant="secondary")
                            with gr.Group(elem_classes="wizard-panel"):
                                gr.Markdown("### Import a local project\nMigration accepts a Grantspace project export; validation runs before any shared records are written.")
                                wizard_import=gr.File(label="Grantspace project export",file_types=[".zip"],type="filepath")
                                wizard_import_btn=gr.Button("Validate and import")
                                wizard_import_status=gr.Markdown()
                    gr.HTML('<div class="privacy-note"><b>Your research stays governed.</b><br>Provider routing and the exact cloud-bound task categories are shown before the project is created.</div>')
                with gr.Column(visible=False,elem_id="wizard-page-2") as wizard_page_2:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Grant details</div><h2>Name this grant</h2><p>Enter only the identifying information known today. Sponsor and mechanism can be corrected later from the authoritative source.</p></div>')
                    with gr.Column(elem_classes="wizard-panel"):
                        wizard_title=gr.Textbox(label="Working title")
                        with gr.Row():
                            wizard_sponsor=gr.Textbox(label="Sponsor")
                            wizard_mechanism=gr.Textbox(label="Mechanism")
                        with gr.Row():
                            wizard_grant_type=gr.Dropdown(["research","clinical_trial","implementation","training","center_program","foundation","rfi_response","custom"],value="custom",label="Grant type",interactive=True)
                            wizard_deadline=gr.Textbox(label="Sponsor deadline",placeholder="YYYY-MM-DD")
                    with gr.Row():wizard_details_back=gr.Button("← Back");wizard_details_next=gr.Button("Continue →",variant="primary")
                with gr.Column(visible=False,elem_id="wizard-page-3") as wizard_page_3:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Grant source</div><h2>Add the authoritative grant ask</h2><p>Upload a file, enter a public URL, or paste the solicitation text. Only one primary source is required.</p></div>')
                    with gr.Column(elem_classes="wizard-panel"):
                        wizard_source=gr.File(label="RFA, RFI, NOFO, or application form",file_count="single",type="filepath")
                        wizard_source_url=gr.Textbox(label="Or import a public URL",placeholder="https://…")
                        wizard_source_text=gr.Textbox(label="Or paste the grant ask",lines=7)
                        wizard_supporting=gr.File(label="Supporting guidance (optional)",file_count="multiple",type="filepath")
                        wizard_brand=gr.File(label="Brand or layout references (optional)",file_count="multiple",type="filepath")
                    with gr.Row():wizard_source_back=gr.Button("← Back");wizard_source_next=gr.Button("Continue →",variant="primary")

                with gr.Column(visible=False,elem_id=f"wizard-page-{WIZARD_CORE_PAGE}") as wizard_core_page:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Core workflow</div><h2>Five outcomes, one clear workflow</h2><p>These five outcomes create the editable proposal. Review them together and continue once.</p></div>')
                    with gr.Column(elem_classes="wizard-summary-table"):
                        for core_index,step in enumerate(WORKFLOW_REGISTRY["core_steps"],1):
                            with gr.Row(elem_classes="wizard-summary-row"):
                                gr.HTML(f'<div class="wizard-summary-copy"><h3>{core_index:02d} · {html.escape(step["title"])}</h3><p>{html.escape(step["description"])}</p><small>Produces: {html.escape(step["output"])}</small></div>',container=False)
                                gr.HTML('<span class="core-included">Included in core workflow</span>',container=False)
                    with gr.Row():wizard_core_back=gr.Button("← Back");wizard_core_next=gr.Button("Continue →",variant="primary")

                with gr.Column(visible=False,elem_id=f"wizard-page-{WIZARD_MODULE_PAGE}") as wizard_module_page:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Optional tools</div><h2>Add only the tools you want</h2><p>Every tool defaults to Skip. Included tools remain advisory and cannot block another outcome or final export.</p></div>')
                    with gr.Row(elem_classes="skip-all-panel"):
                        gr.Markdown("**Fast path:** use the five core outcomes without specialist tools.")
                        wizard_optional_skip_all=gr.Button("Skip all optional tools →",variant="secondary")
                    wizard_module_modes=[]
                    with gr.Column(elem_classes="wizard-summary-table"):
                        for module in WORKFLOW_REGISTRY["optional_modules"]:
                            with gr.Row(elem_classes="wizard-summary-row"):
                                gr.HTML(f'<div class="wizard-summary-copy"><h3>{html.escape(module["title"])}</h3><p>{html.escape(module["description"])}</p><small>{html.escape(module["placement"].replace("_"," ").title())} · Produces: {html.escape(module["output"])}</small></div>',container=False)
                                mode=gr.Radio(
                                    [("Skip","skip"),("Add optionally","include")],value="skip",label=None,
                                    interactive=True,container=False,elem_classes="wizard-choice",
                                )
                                wizard_module_modes.append(mode)
                    with gr.Row():wizard_module_back=gr.Button("← Back");wizard_module_next=gr.Button("Continue →",variant="primary")

                with gr.Column(visible=False,elem_id=f"wizard-page-{WIZARD_REVIEW_PAGE}") as wizard_review_page:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Review setup</div><h2>Choose the depth of critique</h2><p>This choice is used only when the Review simulator tool is included. Review remains advisory and never blocks the grant.</p></div>')
                    wizard_review_mode=gr.Radio(REVIEW_MODE_CHOICES,value=REVIEW_MODE_CHOICES[0][1],label="Review mode",interactive=True,elem_classes="wizard-option-grid")
                    gr.HTML('<div class="privacy-note">Reviewer roles are derived later from the approved solicitation and the versioned registry. Synthetic reviewers never represent named people or predict an award decision.</div>')
                    with gr.Row():wizard_review_back=gr.Button("← Back");wizard_review_next=gr.Button("Continue →",variant="primary")

                with gr.Column(visible=False,elem_id=f"wizard-page-{WIZARD_TEAM_PAGE}") as wizard_team_page:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Team</div><h2>Invite collaborators</h2><p>Invitations change access only. They never add grant-completion rules.</p></div>')
                    with gr.Column(elem_classes="wizard-panel"):
                        wizard_owner_id=gr.State("")
                        wizard_identity_status=gr.Markdown()
                        gr.Markdown("Invite an existing account holder by email, or skip this screen and invite people later from the project workspace.")
                        with gr.Row():
                            wizard_team_email=gr.Textbox(label="Teammate email",placeholder="name@organization.org",scale=2)
                            wizard_team_role=gr.Dropdown(PROJECT_ROLE_CHOICES,value="contributor",label="Project role",scale=1)
                            wizard_team_add=gr.Button("Add teammate",variant="secondary",scale=1)
                        wizard_team=gr.Dataframe(value=[],headers=["Email","Role"],datatype=["str","str"],column_count=(2,"fixed"),interactive=False,label="Pending invitations")
                        with gr.Row():
                            wizard_team_remove_choice=gr.Dropdown([],label="Remove teammate",scale=3)
                            wizard_team_remove=gr.Button("Remove",variant="secondary",scale=1)
                        wizard_team_status=gr.Markdown()
                    with gr.Row():wizard_team_back=gr.Button("← Back");wizard_team_next=gr.Button("Continue →",variant="primary")

                with gr.Column(visible=False,elem_id=f"wizard-page-{WIZARD_ROUTING_PAGE}") as wizard_routing_page:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Model routing</div><h2>Choose where model work runs</h2><p>This policy is stored on the grant and enforced for its model calls.</p></div>')
                    with gr.Column(elem_classes="wizard-panel"):
                        wizard_routing=gr.Dropdown(WORKFLOW_REGISTRY["model_routing_modes"],value=os.getenv("MODEL_ROUTING_MODE","local_only"),label="Model routing policy",interactive=True)
                        gr.Markdown(f"**Local provider:** `{os.getenv('LOCAL_LLM_PROVIDER','not configured')}`  \n**Local model:** `{os.getenv('LOCAL_LLM_API_MODEL',os.getenv('LOCAL_LLM_MODEL','not configured'))}`  \n**Cloud model:** `{os.getenv('CLAUDE_MODEL','not configured')}`")
                        gr.HTML('<div class="privacy-note">Local-only prevents proposal content from being sent to Claude. Hybrid routing shows and records the provider and model for every generated artifact.</div>')
                    with gr.Row():wizard_routing_back=gr.Button("← Back");wizard_routing_next=gr.Button("Preview workflow →",variant="primary")

                with gr.Column(visible=False,elem_id=f"wizard-page-{WIZARD_PREVIEW_PAGE}") as wizard_preview_page:
                    gr.HTML('<div class="wizard-title"><div class="wizard-kicker">Workflow preview</div><h2>Your grant, composed</h2><p>This exact configuration becomes the server-side source of truth for every teammate.</p></div>')
                    wizard_preview=gr.HTML()
                    gr.HTML('<div class="privacy-note"><b>What is enforced</b><br>The five grant outcomes remain available. Every tool you included is optional and cannot block another outcome or final export.</div>')
                    gr.HTML('<div id="wizard-create-progress" role="status" aria-live="polite"><b class="label">Ready to create the shared grant.</b><div class="track"><div class="bar"></div></div></div>')
                    wizard_create_status=gr.Markdown(elem_id="wizard-create-status")
                    with gr.Row():wizard_preview_back=gr.Button("← Back");wizard_create_btn=gr.Button("Create shared grant →",variant="primary",elem_id="wizard-create-button")
                with gr.Row(elem_classes="wizard-footer"):
                    gr.Markdown("Configuration is not persisted until you create the grant.")
                    wizard_progress=gr.Markdown(f"1 of {WIZARD_PAGE_COUNT}",elem_id="wizard-progress")
    global_navigation_command=gr.Textbox(visible=True,elem_id="global-navigation-command")
    gr.HTML('''<nav class="global-nav" aria-label="Global navigation"><span class="brand">Grantspace</span><button type="button" data-global-nav="projects">Grants &amp; wizard</button><button type="button" data-global-nav="workflow">Workflow setup</button><button type="button" data-global-nav="editor">Grant editor</button><button type="button" data-global-nav="admin">Administration</button><a href="/logout">Sign out</a></nav>''',container=False)
    gr.Markdown(f"# Grantspace\n{ORGANIZATION_NAME} · Shared grant writing · Build {GRANT_BUILD_VERSION}",visible=False)
    refresh_shared_updates_btn=gr.Button("↻ Refresh shared updates",variant="secondary",elem_id="refresh-shared-updates",visible=False)
    manual_refresh_status=gr.Markdown(visible=False)
    project_id=gr.State("");interview_questions=gr.State([]);current_question=gr.State(None);current_version=gr.State(None);baseline_body=gr.State("");current_section_key=gr.State("");current_competitive_update_event=gr.State(None)
    workspace_workflow_summary=gr.Markdown(visible=False)
    workspace_workflow_status=gr.JSON(label="Composable workflow status",visible=False)
    with gr.Accordion("Switch grant",open=False):
        with gr.Row():
            recent=gr.Dropdown(label="Saved grant",choices=[]);refresh_projects_btn=gr.Button("Refresh list");open_project_btn=gr.Button("Open",variant="secondary")
    with gr.Accordion("Manage saved grants",open=False):
        gr.Markdown("Grants are persisted on the shared server, scoped to your account memberships, and remain available after every sign-out or restart. Archiving is reversible and never deletes grant records.")
        with gr.Row():
            include_archived_projects=gr.Checkbox(label="Include archived grants",value=False)
            refresh_project_manager_btn=gr.Button("Refresh saved grants")
        project_catalog_table=gr.Dataframe(headers=["Project ID","Title","Sponsor","Mechanism","Workflow stage","My role","State","Last updated","Created"],interactive=False,label="Saved grants available to this account")
        project_catalog_summary=gr.Markdown()
        with gr.Row():
            managed_project_title=gr.Textbox(label="New title (optional)")
            managed_project_action=gr.Dropdown(["Keep active","Archive","Restore"],value="Keep active",label="Lifecycle action")
            apply_project_management_btn=gr.Button("Apply to selected grant")
        project_management_status=gr.Markdown()
    agentic_global_notice=gr.Markdown(visible=False)
    with gr.Tabs() as workspace_tabs:
        with gr.Tab("Grant editor") as document_editor_tab:
            with gr.Column(elem_classes="grant-editor-shell"):
                gr.HTML('<div class="grant-editor-header"><div><h2>Collaborative grant editor</h2><div class="editor-secondary">Edit the complete proposal in one continuous document. Saves create immutable versions; Refresh shared changes pulls teammate updates.</div></div></div>')
                editor_current_key=gr.State(None);editor_version=gr.State(None)
                editor_outline_command=gr.Textbox(visible=True,elem_id="editor-outline-command")
                editor_document_payload=gr.Textbox(visible=True,elem_id="editor-document-payload")
                with gr.Row(elem_classes="grant-editor-layout"):
                    with gr.Column(scale=2,elem_classes="grant-outline-panel"):
                        with gr.Accordion("Sections",open=True):
                            editor_outline=gr.HTML(editor_outline_html([],None),container=False)
                            with gr.Accordion("Add a section",open=False):
                                editor_new_section_title=gr.Textbox(label="Section title",placeholder="Enter a new section title")
                                editor_add_section_btn=gr.Button("Add section",variant="secondary")
                            editor_refresh_btn=gr.Button("↻ Refresh shared changes",variant="secondary")
                    with gr.Column(scale=9,elem_classes="grant-document-panel"):
                        with gr.Row(elem_classes="grant-editor-toolbar"):
                            editor_save_btn=gr.Button("Save document",variant="primary",scale=1)
                            editor_prefill_btn=gr.Button("Draft missing sections",variant="secondary",scale=1)
                            editor_rewrite_btn=gr.Button("Rewrite selected",variant="secondary",scale=1)
                            editor_version_choice=gr.Dropdown(label="Selected section version",choices=[],scale=2,container=False)
                            editor_load_version_btn=gr.Button("Load version",variant="secondary",scale=1)
                        editor_document=gr.HTML(editor_document_html([],None),container=False)
                        editor_status=gr.Markdown("Open a saved grant to begin.")
                    with gr.Column(scale=2,elem_classes="grant-guidance-panel"):
                        gr.Markdown("### Comments & guidance\nQuestions and rules trigger evidence research and a rewrite. Comments remain shared notes.")
                        editor_guidance_view=gr.HTML(editor_guidance_html([]),container=False)
                        editor_guidance_kind=gr.Radio([("Comment","comment"),("Question","question"),("Rule","rule")],value="comment",label="Add")
                        editor_guidance_body=gr.Textbox(label="Message",lines=3,placeholder="Write one clear item")
                        editor_post_guidance_btn=gr.Button("Share",variant="primary")
                        with gr.Accordion("Resolve completed guidance",open=False):
                            editor_resolve_choice=gr.Dropdown(label="Open item")
                            editor_resolve_btn=gr.Button("Mark resolved",variant="secondary")
                        editor_guidance_status=gr.Markdown()
                publish_grant_btn=gr.Button("Publish grant · DOCX + PDF",elem_id="publish-grant-floating")
                editor_publish_files=gr.File(label="Published grant files",file_count="multiple")
                editor_publish_status=gr.Markdown()
        with gr.Tab("1 · Grant Ask",visible=False) as intake_tab:
            with gr.Row():
                with gr.Column(scale=2):
                    project_title=gr.Textbox(label="Working title");sponsor=gr.Textbox(label="Sponsor");mechanism=gr.Textbox(label="Mechanism")
                    gr.Markdown("""### Grant Opportunity
Use whichever input is easiest. Upload, URL, and pasted text are normalized into the same grant-opportunity source before requirements and deterministic submission rules are compiled.""")
                    with gr.Tabs():
                        with gr.Tab("Upload"):
                            source=gr.File(label="Upload RFA / NOFO / grant opportunity (PDF, DOCX, TXT, HTML)",file_count="single",type="filepath")
                        with gr.Tab("URL"):
                            source_url=gr.Textbox(label="Public grant opportunity URL",placeholder="https://...")
                        with gr.Tab("Paste Text"):
                            source_text=gr.Textbox(label="Paste the grant opportunity",lines=16,placeholder="Paste the full funding opportunity, sponsor instructions, or relevant grant announcement text here…")
                    supporting=gr.File(label="Relevant institutional/project materials",file_count="multiple",type="filepath");brand=gr.File(label="Branding/layout inspiration",file_count="multiple",type="filepath")
                    create=gr.Button("Create / Analyze Grant",variant="primary");project_status=gr.Markdown()
                with gr.Column(scale=5):
                    requirements_table=gr.Dataframe(headers=["Requirement","Priority","Evidence or action","Status"],interactive=False,label="Grant requirements")
                    with gr.Row():
                        compile_grant_ask_btn=gr.Button("Compile grant ask",variant="primary")
                        approve_req=gr.Button("✓ Approve Requirements",variant="secondary")
                    req_status=gr.Markdown()
                    gr.Markdown("### Versioned solicitation profile\nThe profile normalizes eligibility, requirements, rubric, deadlines, budget, attachments, and unresolved human questions. Exact source offsets and hashes are validated by the server.")
                    solicitation_version=gr.State(None)
                    with gr.Row():
                        load_solicitation_btn=gr.Button("Load profile")
                        save_solicitation_btn=gr.Button("Save corrected profile")
                        approve_solicitation_btn=gr.Button("✓ Approve solicitation profile",variant="primary")
                    with gr.Row():
                        solicitation_working_title=gr.Textbox(label="Solicitation working title")
                        solicitation_sponsor=gr.Textbox(label="Sponsor")
                        solicitation_mechanism=gr.Textbox(label="Mechanism")
                    solicitation_purpose=gr.Textbox(label="Purpose",lines=4)
                    solicitation_facts=gr.Dataframe(headers=SOLICITATION_FACT_HEADERS,column_count=(1,"fixed"),row_count=(8,"dynamic"),interactive=True,label="Rules and requirements · one per row")
                    solicitation_criteria=gr.Dataframe(headers=SOLICITATION_CRITERION_HEADERS,column_count=(1,"fixed"),row_count=(4,"dynamic"),interactive=True,label="Review criteria · one per row")
                    solicitation_questions=gr.Dataframe(headers=["Question"],column_count=(1,"fixed"),row_count=(3,"dynamic"),interactive=True,label="Questions · one per row")
                    solicitation_artifact_json=gr.JSON(label="Solicitation artifact metadata",visible=False)
                    solicitation_status=gr.Markdown()
                    with gr.Row():
                        solicitation_return_reason=gr.Textbox(label="Return-for-revision rationale",placeholder="State the exact correction required")
                        return_solicitation_btn=gr.Button("↩ Return approved profile for revision")
        with gr.Tab("2 · Research Plan",visible=False) as framework_tab:
            gr.Markdown("### Sponsor-mapped research plan\nGenerate from the exact approved solicitation profile, then edit the argument, mappings, evidence gaps, ownership, dependencies, and word allocations before approval.")
            framework_version=gr.State(None)
            with gr.Row():
                load_framework_btn=gr.Button("Load framework")
                generate_framework_btn=gr.Button("Generate from approved solicitation",variant="primary")
                save_framework_btn=gr.Button("Save framework")
                approve_framework_btn=gr.Button("✓ Approve framework",variant="primary")
            framework_argument=gr.Textbox(label="Overall proposal argument",lines=5)
            framework_references=gr.Dataframe(headers=REFERENCE_HEADERS,interactive=False,label="Approved solicitation and active-member reference catalog")
            framework_nodes=gr.Dataframe(headers=FRAMEWORK_HEADERS,column_count=(len(FRAMEWORK_HEADERS),"fixed"),row_count=(8,"dynamic"),interactive=True,label="Ordered sponsor-mapped framework nodes")
            framework_artifact_json=gr.JSON(label="Framework artifact metadata",visible=False)
            framework_status=gr.Markdown()
            with gr.Row():
                framework_return_reason=gr.Textbox(label="Return-for-revision rationale",placeholder="State the exact correction required")
                return_framework_btn=gr.Button("↩ Return approved framework for revision")
        with gr.Tab("3 · Aims",visible=False) as aims_tab:
            gr.Markdown("### Versioned aims\nKeep objectives, central thesis, rationale, approach, outcomes, impact, innovation, dependencies, classifications, and supporting evidence explicit.")
            aims_version=gr.State(None)
            with gr.Row():
                load_aims_btn=gr.Button("Load aims")
                generate_aims_btn=gr.Button("Generate from approved framework",variant="primary")
                save_aims_btn=gr.Button("Save aims")
                approve_aims_btn=gr.Button("✓ Approve aims",variant="primary")
            with gr.Row():
                aims_objective=gr.Textbox(label="Overall objective",lines=4)
                aims_hypothesis=gr.Textbox(label="Central hypothesis or thesis",lines=4)
            aims_references=gr.Dataframe(headers=REFERENCE_HEADERS,interactive=False,label="Approved framework and project evidence reference catalog")
            core_aims=gr.Dataframe(headers=CORE_AIM_HEADERS,column_count=(len(CORE_AIM_HEADERS),"fixed"),row_count=(3,"dynamic"),interactive=True,label="Structured aims · classification must be fact, estimate, or assumption")
            aims_artifact_json=gr.JSON(label="Aims artifact metadata",visible=False)
            aims_status=gr.Markdown()
            with gr.Row():
                aims_return_reason=gr.Textbox(label="Return-for-revision rationale",placeholder="State the exact correction required")
                return_aims_btn=gr.Button("↩ Return approved aims for revision")
        with gr.Tab("Tools · Investigator Interview",visible=False) as interview_tab:
            with gr.Row():
                with gr.Column(scale=3):
                    generate_q=gr.Button("Generate / Recompute Missing Questions",variant="primary");question_card=gr.HTML(render_question(None));answer=gr.Textbox(label="Answer")
                    with gr.Row():confidence=gr.Dropdown(["high","medium","low"],value="high",label="Confidence");classification=gr.Dropdown(["verified_fact","investigator_estimate","assumption","unknown"],value="verified_fact",label="Classification")
                    answer_notes=gr.Textbox(label="Supporting explanation / notes",lines=4);answered_by=gr.Textbox(label="Answered by / role");submit=gr.Button("Save Answer & Continue",variant="primary");interview_status=gr.Markdown()
                with gr.Column(scale=2):interview_table=gr.JSON(label="Interview state")
        with gr.Tab("4 · Evidence",visible=False) as research_tab:
            gr.Markdown("### Approved search plan\nGenerate a solicitation-, framework-, and aims-grounded query plan, let contributors correct it, and require named approval before the application performs any external search.")
            search_plan_version=gr.State(None)
            with gr.Row():
                max_queries=gr.Slider(1,24,value=8,step=1,label="Maximum planned queries")
                generate_search_plan_btn=gr.Button("Generate versioned plan",variant="primary")
                load_search_plan_btn=gr.Button("Load latest plan")
                save_search_plan_btn=gr.Button("Save plan corrections")
                approve_search_plan_btn=gr.Button("✓ Approve search plan")
            search_plan_upstream=gr.Markdown()
            search_plan_references=gr.Dataframe(headers=REFERENCE_HEADERS,interactive=False,label="Approved upstream requirement, criterion, and aim catalog")
            search_plan_queries=gr.Dataframe(headers=LITERATURE_QUERY_HEADERS,column_count=(len(LITERATURE_QUERY_HEADERS),"fixed"),row_count=(8,"dynamic"),interactive=True,label="Reviewable external research queries")
            search_plan_artifact_json=gr.JSON(label="Search-plan artifact metadata",visible=False)
            search_plan_status=gr.Markdown()
            with gr.Row():
                search_plan_return_reason=gr.Textbox(label="Return-for-revision rationale",placeholder="State the exact query-plan correction required")
                return_search_plan_btn=gr.Button("↩ Return approved search plan for revision")
            shared_literature_sync_status=gr.Markdown("Live collaboration sync starts after a project is opened.")
            with gr.Row():results_per=gr.Slider(1,10,value=5,step=1,label="Results per approved query");research_btn=gr.Button("Run approved plan atomically",variant="primary")
            evidence_table=gr.Dataframe(headers=["Evidence ID","Requirement","Source type","Evidence purpose","Status","Confidence","URL"],interactive=False);research_status=gr.Markdown()
            gr.Markdown("### Reproducible literature manifest\nReview the exact approved upstream versions, queries, evidence dispositions, citations, contradictions, and unresolved risks before locking the research run.")
            literature_version=gr.State(None)
            with gr.Row():
                load_literature_btn=gr.Button("Load latest manifest")
                save_literature_btn=gr.Button("Save manifest corrections")
                approve_literature_btn=gr.Button("✓ Approve literature manifest",variant="primary")
            literature_manifest_summary=gr.Markdown()
            literature_references=gr.Dataframe(headers=REFERENCE_HEADERS,interactive=False,label="Approved upstream artifacts and project evidence/source/citation catalog")
            literature_queries=gr.Dataframe(headers=LITERATURE_QUERY_HEADERS,column_count=(len(LITERATURE_QUERY_HEADERS),"fixed"),row_count=(8,"dynamic"),interactive=True,label="Approved solicitation- and aim-grounded queries")
            literature_needs=gr.Dataframe(headers=EVIDENCE_NEED_HEADERS,column_count=(len(EVIDENCE_NEED_HEADERS),"fixed"),row_count=(8,"dynamic"),interactive=True,label="Evidence-need dispositions")
            literature_contradictions=gr.Dataframe(headers=["Contradiction"],column_count=(1,"fixed"),row_count=(3,"dynamic"),interactive=True,label="Contradictions and unresolved tensions")
            literature_artifact_json=gr.JSON(label="Literature artifact metadata",visible=False)
            literature_artifact_status=gr.Markdown()
            with gr.Row():
                literature_return_reason=gr.Textbox(label="Return-for-revision rationale",placeholder="State the evidence, disposition, or contradiction correction required")
                return_literature_btn=gr.Button("↩ Return approved literature manifest for revision")
            gr.Markdown("### Compiled HPC Knowledge Index")
            with gr.Row():rebuild_idx=gr.Button("Build / Refresh MMAP Index");status_idx=gr.Button("Check Index Status")
            index_message=gr.Markdown();index_json=gr.JSON(label="Index manifest / status")
            with gr.Row():retrieval_query=gr.Textbox(label="Test hybrid retrieval query");retrieval_k=gr.Slider(1,50,value=12,step=1,label="Top K");retrieval_btn=gr.Button("Run Retrieval")
            retrieval_table=gr.Dataframe(headers=["Score","Semantic","BM25","Evidence","Freshness","Graph boost","Kind","Source","Excerpt"],interactive=False)
        with gr.Tab("Team",visible=False) as team_tab:
            gr.Markdown("### Shared project collaboration\nEvery message, comment, task, invitation, notification, and approval decision is tied to an authenticated user and the persisted project workflow.")
            team_permissions=gr.State({})
            with gr.Row():
                refresh_team_workspace_btn=gr.Button("Refresh team workspace",variant="primary")
                team_workspace_status=gr.Markdown()
            with gr.Tabs():
                with gr.Tab("Project health"):
                    project_health_summary_md=gr.Markdown("Project health loads from the shared server after a project is opened.")
                    project_health_table=gr.Dataframe(headers=["Severity","Kind","Finding","Details","Owner user ID","Workflow step","Due","Required action"],interactive=False,label="Authoritative active risks · recalculated from persisted project state")
                with gr.Tab("Members & invitations"):
                    team_members=gr.Dataframe(headers=["User ID","Name","Email","Project role","Online","Joined","Last seen"],interactive=False,label="Authenticated project members · presence is advisory; edit safety uses version checks")
                    with gr.Row():
                        existing_member_user_id=gr.Textbox(label="Existing account user ID")
                        existing_member_role=gr.Dropdown(PROJECT_ROLE_CHOICES,value="contributor",label="Project role")
                        add_existing_member_btn=gr.Button("Add existing account")
                    existing_member_status=gr.Markdown()
                    gr.Markdown("#### Invite by verified email\nInvitation links are single-use, expire, and can only be accepted by an authenticated account with the same email address.")
                    with gr.Row():
                        project_invite_email=gr.Textbox(label="Invite email")
                        project_invite_role=gr.Dropdown(PROJECT_ROLE_CHOICES,value="contributor",label="Role")
                        project_invite_days=gr.Slider(1,30,value=7,step=1,label="Expires in days")
                    with gr.Row():
                        create_project_invite_btn=gr.Button("Create and email invitation",variant="primary")
                        active_project_invite=gr.Dropdown(label="Active invitation")
                        revoke_project_invite_btn=gr.Button("Revoke invitation")
                    project_invite_status=gr.Markdown()
                    project_invites=gr.Dataframe(headers=["Invite ID","Email","Role","Status","Expires","Created"],interactive=False,label="Invitation audit history")
                with gr.Tab("Channels & threads"):
                    with gr.Row():
                        team_channel_kind=gr.Dropdown([("General","general"),("Framework","framework"),("Aims","aims"),("Proposal section","section")],value="general",label="Channel")
                        team_channel_subject=gr.Textbox(label="Section key (section channels only)")
                        load_team_channel_btn=gr.Button("Load channel")
                    team_channel_html=gr.HTML('<div class="team-chat">Load a channel to begin.</div>')
                    team_messages=gr.Dataframe(headers=["Message ID","Parent ID","Author","Author user ID","Created","Message"],interactive=False,label="Threaded message audit")
                    with gr.Row():
                        team_reply_message_id=gr.Dropdown(label="Reply to message (optional)")
                        team_mentions=gr.Dropdown(multiselect=True,label="Mention project members")
                    team_message_body=gr.Textbox(label="Message",lines=4)
                    post_team_channel_btn=gr.Button("Post authenticated message",variant="primary")
                    team_channel_status=gr.Markdown()
                with gr.Tab("Version comments"):
                    gr.Markdown("Comments are anchored to an exact immutable artifact or section version. Resolving a comment preserves its complete history.")
                    with gr.Row():
                        comment_target_type=gr.Dropdown([("Proposal section","section"),("Workflow artifact","workflow_artifact")],value="section",label="Target type")
                        comment_target_key=gr.Textbox(label="Section key or workflow artifact type")
                        comment_version_id=gr.Number(label="Exact stored version ID",precision=0)
                        current_comment_version_btn=gr.Button("Use current version")
                        load_comments_btn=gr.Button("Load comments")
                    artifact_comments=gr.Dataframe(headers=["Comment ID","Parent ID","Author","Created","Start","End","Quoted text","Comment","Resolved","Resolved at"],interactive=False,label="Version-anchored comments")
                    with gr.Row():
                        comment_start=gr.Number(label="Start character offset",precision=0)
                        comment_end=gr.Number(label="End character offset",precision=0)
                        comment_parent_id=gr.Dropdown(label="Reply to comment (optional)")
                        comment_resolve_id=gr.Dropdown(label="Open comment to resolve")
                    comment_quote=gr.Textbox(label="Quoted text",lines=2)
                    comment_body=gr.Textbox(label="Comment",lines=4)
                    with gr.Row():
                        post_comment_btn=gr.Button("Post version comment",variant="primary")
                        resolve_comment_btn=gr.Button("Resolve selected comment")
                    comment_status=gr.Markdown()
                with gr.Tab("Tasks & deadlines"):
                    team_tasks=gr.Dataframe(headers=["Task ID","Priority","Status","Title","Owner user ID","Due","Source","Dependencies","Updated"],interactive=False,label="Shared project tasks")
                    with gr.Row():
                        task_title=gr.Textbox(label="Task title")
                        task_owner=gr.Dropdown(label="Owner")
                        task_priority=gr.Dropdown(["low","normal","high","critical"],value="normal",label="Priority")
                        task_due=gr.Textbox(label="Due date/time (ISO-8601)")
                    task_description=gr.Textbox(label="Task description",lines=3)
                    task_dependencies=gr.Dropdown(multiselect=True,label="Dependencies")
                    create_project_task_btn=gr.Button("Create and notify owner",variant="primary")
                    with gr.Row():
                        active_project_task=gr.Dropdown(label="Task")
                        project_task_status=gr.Dropdown(["open","in_progress","blocked","complete","cancelled"],value="in_progress",label="New status")
                        update_project_task_btn=gr.Button("Update task status")
                    task_action_status=gr.Markdown()
                with gr.Tab("Notifications, approvals & activity"):
                    team_notifications=gr.Dataframe(headers=["Notification ID","Unread","Kind","Details","Created","Read at"],interactive=False,label="Your project notifications")
                    with gr.Row():
                        unread_notification_id=gr.Dropdown(label="Unread notification")
                        mark_notification_read_btn=gr.Button("Mark selected notification read")
                    notification_status=gr.Markdown()
                    approval_routing_table=gr.Dataframe(headers=["Artifact","Current version","Owner user ID","Approver user IDs","Approvals","Required","Threshold met","Artifact approved"],interactive=False,label="Configured approval-routing status")
                    team_activity=gr.Dataframe(headers=["Created","Actor","Kind","Detail"],interactive=False,label="Append-only project activity")
        with gr.Tab("Tools · Clinical Design",visible=False) as clinical_tab:
            gr.Markdown("### Authoritative clinical study model\nThese structured values are the source of truth injected into every grant section. Deterministic checks flag feasibility and consistency issues; they never rewrite human-approved prose.")
            with gr.Row():
                load_clinical_btn=gr.Button("Load Saved Study")
                save_clinical_btn=gr.Button("Save & Analyze Clinical Study",variant="primary")
            with gr.Row():
                with gr.Column():
                    clinical_problem=gr.Textbox(label="Clinical problem",lines=3)
                    knowledge_gap=gr.Textbox(label="Knowledge gap",lines=3)
                    central_hypothesis=gr.Textbox(label="Central hypothesis",lines=3)
                with gr.Column():
                    disease=gr.Textbox(label="Disease / clinical population")
                    disease_stage=gr.Textbox(label="Stage / subtype")
                    biomarker=gr.Textbox(label="Biomarker / molecular criteria")
                    inclusion=gr.Textbox(label="Inclusion criteria · one per line",lines=4)
                    exclusion=gr.Textbox(label="Exclusion criteria · one per line",lines=4)
            gr.Markdown("#### Study design")
            with gr.Row():
                design_type=gr.Textbox(label="Design type")
                study_phase=gr.Textbox(label="Phase")
                randomization=gr.Textbox(label="Randomization")
                allocation_ratio=gr.Textbox(label="Allocation ratio")
                blinding=gr.Textbox(label="Blinding")
                follow_up_months=gr.Number(label="Follow-up months",value=None)
                design_sites=gr.Number(label="Sites",value=None,precision=0)
            gr.Markdown("#### Recruitment feasibility")
            with gr.Row():
                available_patients=gr.Number(label="Available patients / site / month",value=None)
                eligibility_pct=gr.Number(label="Eligibility rate %",value=None)
                biomarker_pct=gr.Number(label="Biomarker-positive rate %",value=None)
                consent_pct=gr.Number(label="Consent rate %",value=None)
                target_enrollment=gr.Number(label="Target enrollment",value=None,precision=0)
                accrual_months=gr.Number(label="Planned accrual months",value=None)
                recruitment_sites=gr.Number(label="Recruitment sites",value=None,precision=0)
            with gr.Row():
                scenario_sites=gr.Textbox(label="Scenario site counts · comma-separated",placeholder="e.g., 1,2,3")
                scenario_consent=gr.Textbox(label="Scenario consent rates % · comma-separated",placeholder="e.g., 50,65,80")
                scenario_biomarker=gr.Textbox(label="Scenario biomarker-positive rates % · comma-separated",placeholder="e.g., 25,40,55")
                scenario_btn=gr.Button("Run Recruitment Scenario Sweep")
            scenario_table=gr.Dataframe(headers=["Sites","Consent %","Biomarker %","Expected enrollments/mo","Required/mo","Estimated accrual months","Feasible","Shortfall/mo"],interactive=False)
            scenario_status=gr.Markdown()
            gr.Markdown("#### Deterministic statistics")
            with gr.Row():
                test_type=gr.Dropdown(["two_proportions","one_proportion","two_means","log_rank"],label="Sample-size method",allow_custom_value=False)
                alpha=gr.Number(label="Two-sided alpha",value=None)
                power=gr.Number(label="Power (0–1)",value=None)
                attrition_pct=gr.Number(label="Attrition %",value=None)
            with gr.Row():
                control_rate=gr.Number(label="Control rate (0–1)",value=None)
                treatment_rate=gr.Number(label="Treatment rate (0–1)",value=None)
                null_rate=gr.Number(label="Null rate (0–1)",value=None)
                alternative_rate=gr.Number(label="Alternative rate (0–1)",value=None)
                mean_delta=gr.Number(label="Mean difference",value=None)
                std_dev=gr.Number(label="Standard deviation",value=None)
                hazard_ratio=gr.Number(label="Hazard ratio",value=None)
                event_probability=gr.Number(label="Expected event probability (0–1)",value=None)
            calculate_n_btn=gr.Button("Calculate Sample Size")
            sample_size_json=gr.JSON(label="Sample-size calculation")
            sample_size_status=gr.Markdown()
            gr.Markdown("#### Specific Aims")
            aims_table=gr.Dataframe(headers=AIM_HEADERS,datatype=["str"]*8,row_count=(3,"dynamic"),column_count=(8,"fixed"),interactive=True)
            gr.Markdown("#### Study arms")
            arms_table=gr.Dataframe(headers=ARM_HEADERS,datatype=["str","str","str","bool"],row_count=(3,"dynamic"),column_count=(4,"fixed"),interactive=True)
            gr.Markdown("#### Endpoints")
            gr.Markdown("Analysis family values: binary — `chi_square`, `fisher_exact`, `logistic_regression`, `two_proportions`; continuous — `t_test`, `anova`, `linear_regression`, `mixed_model`; count — `poisson`, `negative_binomial`; time-to-event — `log_rank`, `cox`, `cox_regression`; ordinal — `ordinal_logistic`, `wilcoxon`.")
            endpoints_table=gr.Dataframe(headers=ENDPOINT_HEADERS,datatype=["str","str","str","bool","str"],row_count=(3,"dynamic"),column_count=(5,"fixed"),interactive=True)
            gr.Markdown("#### Timeline")
            timeline_table=gr.Dataframe(headers=TIMELINE_HEADERS,datatype=["str","str","number","number","str"],row_count=(4,"dynamic"),column_count=(5,"fixed"),interactive=True)
            gr.Markdown("#### Resources")
            resources_table=gr.Dataframe(headers=RESOURCE_HEADERS,datatype=["str","str","bool","bool","str"],row_count=(4,"dynamic"),column_count=(5,"fixed"),interactive=True)
            clinical_assessment=gr.JSON(label="Deterministic clinical assessment")
            clinical_status=gr.Markdown()
        with gr.Tab("Tools · Competitive Intelligence",visible=False) as competitive_tab:
            gr.Markdown("### Public competitive applicant intelligence\nThe system builds a capability profile from this grant and clinical design, then identifies organizations whose **public** grants, publications, clinical trials, patents/IP signals, and disclosed technologies overlap with that profile. These are potential capability-matched competitors—not confirmed applicants. The resulting strategy is injected into grant drafting to emphasize defensible superiority and differentiation.")
            with gr.Row():
                generate_comp_profile_btn=gr.Button("Generate Likely Strong-Applicant Profile",variant="secondary")
                load_competitive_btn=gr.Button("Open / Auto-Refresh Competitive Intelligence")
                run_competitive_btn=gr.Button("Refresh Public Competitive Intelligence Now",variant="primary")
            competitive_profile_json=gr.JSON(label="Strong-applicant capability profile")
            competitive_status=gr.Markdown()
            competitor_table=gr.Dataframe(headers=["Rank","Potential competitor","Overall","Prior grants","Publications","Clinical trials","Patent/IP signal","Technology","Breadth","Public assets"],interactive=False,label="Capability-matched organizations · public evidence only")
            provider_status_json=gr.JSON(label="Public provider status")
            asset_table=gr.Dataframe(headers=["Potential competitor","Asset type","Relevance","Public asset","Year","Provider","URL"],interactive=False,label="Top public grants / publications / trials / IP signals / technology evidence")
            competitive_strategy=gr.Markdown("No competitive positioning run yet.")
            competitive_raw=gr.JSON(label="Auditable competitive run",visible=False)
        with gr.Tab("Tools · Sponsor Compliance",visible=False) as compliance_tab:
            gr.Markdown("""### Deterministic Sponsor Compliance
The approved funding opportunity is compiled into atomic submission rules. You can correct the parsed rules before approval. Hard sponsor rules—not AI opinion—control final export readiness.""")
            with gr.Row():
                load_compliance_btn=gr.Button("Open Compliance Profile")
                compile_compliance_btn=gr.Button("Recompile Rules from Current Opportunity")
                approve_compliance_btn=gr.Button("✓ Approve Compliance Profile",variant="primary")
                measure_compliance_btn=gr.Button("Run Rendered Preflight")
            compliance_profile_state=gr.State({})
            opportunity_source_preview=gr.Textbox(label="Authoritative stored funding-opportunity source used for deterministic rule compilation",lines=18,interactive=False)
            with gr.Row():
                compliance_sponsor=gr.Textbox(label="Sponsor")
                compliance_mechanism=gr.Textbox(label="Mechanism")
                submission_system=gr.Textbox(label="Submission system / portal")
                compliance_deadline=gr.Textbox(label="Deadline (YYYY-MM-DD)")
            compliance_rules=gr.Dataframe(headers=COMPLIANCE_HEADERS,row_count=(8,"dynamic"),column_count=(1,"fixed"),interactive=True,label="Rules · one plain-language rule per row")
            save_compliance_btn=gr.Button("Save rules")
            with gr.Accordion("Advanced audit details and manual resolutions",open=False):
                gr.Markdown("Internal IDs and exact source offsets are retained for auditability. Users select a rule by its wording; IDs never need to be entered.")
                compliance_provenance=gr.Dataframe(headers=["Internal rule ID","Source status","Document ID","Page","Start byte","End byte","Exact source excerpt"],interactive=False,label="Exact source provenance")
                compliance_findings=gr.Dataframe(headers=["Internal rule ID","Severity","Mandatory","Status","Type","Target","Detail","Source excerpt"],interactive=False,label="Deterministic assessment")
                compliance_json=gr.JSON(label="Machine-readable assessment")
                gr.Markdown("#### Resolve a rule requiring authoritative human confirmation")
                with gr.Row():
                    compliance_rule_id=gr.Dropdown(label="Rule",choices=[])
                    compliance_resolution=gr.Dropdown(["satisfied","not_applicable","waived","unresolved"],value="satisfied",label="Resolution")
                    compliance_resolved_by=gr.Textbox(label="Resolved by / role")
                compliance_resolution_notes=gr.Textbox(label="Resolution rationale / evidence note",lines=3)
                resolve_compliance_btn=gr.Button("Save manual resolution")
            gr.Markdown("""#### Submission attachments
Register the actual files that must travel with the proposal. Use a stable slot such as `letters_of_support`, `biosketches`, `data_management_plan`, or the sponsor's attachment name.""")
            with gr.Row():
                artifact_slot=gr.Textbox(label="Submission slot")
                artifact_files=gr.File(label="Attach files",file_count="multiple",type="filepath")
                register_artifact_btn=gr.Button("Register Attachment(s)")
            artifact_table=gr.Dataframe(headers=["Slot","Filename","Extension","SHA-256"],interactive=False,label="Registered submission artifacts")
            compliance_status=gr.Markdown()
        with gr.Tab("5 · Proposal",visible=False) as writing_tab:
            with gr.Row():section=gr.Dropdown(DEFAULT_SECTIONS,value=(DEFAULT_SECTIONS[0] if DEFAULT_SECTIONS else None),label="Section",allow_custom_value=True);high=gr.Checkbox(label="Escalate this draft to Claude (sends this section’s compiled context to the configured cloud API)",value=False)
            additional=gr.Textbox(lines=4,label="Optional additional human context");gen=gr.Button("Compile Context & Draft Section",variant="primary")
            section_update_banner=gr.Markdown()
            preview_box=gr.HTML('<div class="page-frame"><div style="background:white;padding:32px">Open a project and select a section.</div></div>')
            with gr.Row():edit=gr.Button("✎ Edit");approve_btn=gr.Button("✓ Approve Section",variant="primary")
            with gr.Row():
                section_return_reason=gr.Textbox(label="Return-for-revision rationale",placeholder="State the exact section correction required")
                return_section_btn=gr.Button("↩ Return approved section for revision")
            editor=gr.Textbox(lines=20,label="Section text",visible=False)
            with gr.Row():save_btn=gr.Button("Save Edit",visible=False);cancel=gr.Button("Cancel Edit",visible=False)
            write_status=gr.Markdown("No section loaded.")
            with gr.Accordion("Version history, comparison & recovery",open=False):
                gr.Markdown("Every save and restore creates an immutable version. Restoring never deletes later work, and stale edits are reconciled against their exact base version.")
                refresh_version_history_btn=gr.Button("Refresh version history")
                version_history_table=gr.Dataframe(headers=["Version","Created","Author","Source","Approved","Base version","Restored from","Characters","Preview"],interactive=False,label="Append-only section history")
                with gr.Row():
                    compare_from_version=gr.Dropdown(label="Compare from")
                    compare_to_version=gr.Dropdown(label="Compare to")
                    compare_versions_btn=gr.Button("Compare exact versions")
                version_diff=gr.HTML('<div class="version-diff">Choose two versions to compare.</div>')
                version_history_status=gr.Markdown()
                with gr.Row():
                    restore_version_id=gr.Dropdown(label="Version to restore")
                    restore_version_btn=gr.Button("Restore as a new version",variant="secondary")
        with gr.Tab("Tools · Review & Causal Critique",visible=False) as review_tab:
            gr.Markdown("### Solicitation-grounded synthetic review panel\nFeedback is generated from role archetypes derived from the approved solicitation. It does not represent real reviewers or predict a sponsor decision.")
            review_plan_id=gr.State(None);review_run_id=gr.State(None);review_plan_json=gr.JSON(visible=False);review_run_json=gr.JSON(visible=False)
            with gr.Row():
                refresh_reviewer_roles_btn=gr.Button("Derive roles from approved solicitation")
                review_mode=gr.Dropdown(REVIEW_MODE_CHOICES,value=(REVIEW_MODE_CHOICES[0][1] if REVIEW_MODE_CHOICES else None),label="Review mode")
                create_review_plan_btn=gr.Button("Create panel plan")
                approve_review_plan_btn=gr.Button("Approve panel plan",variant="secondary")
                run_review_panel_btn=gr.Button("Freeze snapshot & run panel",variant="primary")
                approve_review_result_btn=gr.Button("Approve validated review",variant="secondary")
            review_notice=gr.Markdown()
            review_roles=gr.Dataframe(headers=["Role key","Role","Responsibility","Solicitation criterion IDs"],interactive=False,label="Derived and approved reviewer roles")
            review_status=gr.Markdown()
            panel_summary=gr.JSON(label="Synthetic panel summary and disagreement map")
            reviewer_findings=gr.Dataframe(headers=["Role","Criterion","Score","Confidence","Strengths","Weaknesses","Proposal anchors"],interactive=False,label="Validated independent reviews")
            revision_findings=gr.Dataframe(headers=["Index","Priority","Title","Description","Rationale","Proposal anchors"],interactive=False,label="Grounded revision backlog")
            with gr.Row():
                revision_task_indexes=gr.Textbox(label="Revision task indexes",placeholder="Blank imports all; or 0, 2, 4")
                revision_task_owner=gr.Textbox(label="Assigned owner user ID")
                revision_task_due=gr.Textbox(label="Due date/time (optional)")
                create_revision_tasks_btn=gr.Button("Create assigned tasks")
            revision_task_result=gr.JSON(label="Created task records",visible=False);revision_task_status=gr.Markdown()
            gr.Markdown("#### Human-editable causal model\nModel-proposed nodes and edges remain inferred until a PI/methodologist confirms an append-only corrected version.")
            with gr.Row():
                load_causal_btn=gr.Button("Load causal model history")
                confirm_causal=gr.Checkbox(label="Confirm as PI/methodologist-reviewed",value=False)
                save_causal_btn=gr.Button("Save new causal model version")
            causal_editor=gr.JSON(label="Causal graph, assumptions, threats, and claim checks")
            causal_history=gr.JSON(label="Append-only causal model history",visible=False);causal_status=gr.Markdown()
        with gr.Tab("Proposal · Readiness & Export",visible=False) as export_tab:
            gr.Markdown("### Human-Approved Grant Assembly\nOnly exact section versions approved by the human are included below or in final exports. AI drafts and unapproved edits are excluded.")
            preview_approved_btn=gr.Button("Refresh Approved Grant Preview",variant="secondary")
            approved_sections_table=gr.Dataframe(headers=["Order","Section","Status","Approved version"],datatype=["number","str","str","number"],interactive=False,label="Approved section assembly")
            approved_grant_preview=gr.HTML('<div class="page-frame"><div style="background:white;padding:32px">Approve sections, then refresh this preview to see the grant assembled in final document order.</div></div>')
            approved_preview_status=gr.Markdown()
            with gr.Row():
                portable_export_btn=gr.Button("Create portable project archive",variant="secondary")
                portable_export_file=gr.File(label="Portable project archive")
            portable_export_status=gr.Markdown()
            check_ready=gr.Button("Check Submission Readiness");readiness_json=gr.JSON(label="Backend readiness gates");readiness_status=gr.Markdown()
            fmt=gr.Radio(["DOCX","PDF","BOTH"],value="DOCX",label="Would you like me to produce a professionally formatted DOCX, PDF, or both?",visible=False);export_btn=gr.Button("Compile Approved Grant",variant="primary",visible=False);export_file=gr.File(label="Generated file(s)",file_count="multiple");export_status=gr.Markdown()
        with gr.Tab("Tools · Diagnostics",visible=False) as diagnostics_tab:
            gr.Markdown("""### Production runtime diagnostics
This view exposes non-secret runtime/build information and a local HPC benchmark. It does not display API keys or uploaded grant content.""")
            with gr.Row():
                system_info_btn=gr.Button("Refresh Runtime Information")
                hpc_bench_btn=gr.Button("Run HPC Benchmark",variant="secondary")
            diagnostics_status=gr.Markdown()
            with gr.Row():
                system_info_json=gr.JSON(label="Runtime / build information")
                hpc_benchmark_json=gr.JSON(label="MMAP / OpenMP / BLAS benchmark")
        with gr.Tab("System administration",visible=os.getenv("AUTH_MODE","local_single_user")=="internal_accounts") as account_admin_tab:
            gr.Markdown("### Account administration\nOnly the bootstrap system administrator can use these controls. Accounts are created with a username, verified email destination, and temporary password; the first login is restricted to changing that password.")
            with gr.Row():
                account_username=gr.Textbox(label="Username")
                account_email=gr.Textbox(label="Email")
                account_display_name=gr.Textbox(label="Display name")
                account_temp_password=gr.Textbox(label="Temporary password",type="password")
            with gr.Row():
                create_account_btn=gr.Button("Create account and email credentials",variant="primary")
                refresh_accounts_btn=gr.Button("Refresh accounts")
            account_admin_status=gr.Markdown()
            accounts_table=gr.Dataframe(headers=["User ID","Username","Email","Display name","System role","Must change password","Active","Last seen","Locked until"],interactive=False,label="Organization accounts")
            gr.Markdown("#### Account recovery and access")
            account_target_id=gr.Textbox(label="Target user ID",placeholder="Copy the stable user ID from the table")
            with gr.Row():
                send_account_reset_btn=gr.Button("Email password-reset link")
                disable_account_btn=gr.Button("Disable account")
                enable_account_btn=gr.Button("Enable account")

    wizard_pages=[wizard_page_1,wizard_page_2,wizard_page_3,wizard_core_page,wizard_module_page,wizard_review_page,wizard_team_page,wizard_routing_page,wizard_preview_page]
    wizard_nav_outputs=wizard_pages+[wizard_rail,wizard_progress]
    global_nav_fields=[wizard_title,wizard_sponsor,wizard_mechanism,wizard_deadline,wizard_grant_type,wizard_preset,wizard_modules,wizard_required_modules,wizard_review_required,wizard_review_mode,wizard_routing]+wizard_module_modes
    global_nav_outputs=[wizard_shell,wizard_edit_project,wizard_create_btn]+wizard_nav_outputs+global_nav_fields+[global_navigation_command]
    global_navigation_event=global_navigation_command.change(gateway_callback(global_navigation_state),[project_id,global_navigation_command],global_nav_outputs,show_progress="hidden")
    global_navigation_event.success(fn=None,js=wizard_nav_from_progress_js(),queue=False)
    wizard_new_btn.click(wizard_go(2),outputs=wizard_nav_outputs,js=wizard_nav_js(2)).then(gateway_callback(authenticated_identity),outputs=[wizard_owner_id,wizard_identity_status])
    wizard_new_btn.click(lambda:(None,gr.update(value="Create shared grant →",interactive=True)),outputs=[wizard_edit_project,wizard_create_btn],queue=False)
    wizard_details_back.click(wizard_go(1),outputs=wizard_nav_outputs,js=wizard_nav_js(1))
    wizard_details_event=wizard_details_next.click(validate_grant_details_and_continue,[wizard_title,wizard_deadline],wizard_nav_outputs)
    wizard_details_event.success(fn=None,js=wizard_nav_js(3),queue=False)
    wizard_source_back.click(wizard_go(2),outputs=wizard_nav_outputs,js=wizard_nav_js(2))
    wizard_source_event=wizard_source_next.click(validate_grant_source_and_continue,[wizard_source,wizard_source_url,wizard_source_text,wizard_edit_project],wizard_nav_outputs)
    wizard_source_event.success(fn=None,js=wizard_nav_js(WIZARD_CORE_PAGE),queue=False)
    wizard_core_back.click(wizard_go(3),outputs=wizard_nav_outputs,js=wizard_nav_js(3))
    wizard_core_next.click(wizard_go(WIZARD_MODULE_PAGE),outputs=wizard_nav_outputs,js=wizard_nav_js(WIZARD_MODULE_PAGE))
    wizard_module_back.click(wizard_go(WIZARD_CORE_PAGE),outputs=wizard_nav_outputs,js=wizard_nav_js(WIZARD_CORE_PAGE))
    wizard_optional_skip_event=wizard_optional_skip_all.click(optional_tools_skip_all,outputs=wizard_module_modes+wizard_nav_outputs)
    wizard_optional_skip_event.success(fn=None,js=wizard_nav_js(WIZARD_TEAM_PAGE),queue=False)
    wizard_optional_continue_event=wizard_module_next.click(optional_tools_continue,wizard_module_modes,wizard_nav_outputs)
    wizard_optional_continue_event.success(fn=None,js=wizard_nav_from_progress_js(),queue=False)
    wizard_review_back.click(wizard_go(WIZARD_MODULE_PAGE),outputs=wizard_nav_outputs,js=wizard_nav_js(WIZARD_MODULE_PAGE))
    wizard_review_next.click(wizard_go(WIZARD_TEAM_PAGE),outputs=wizard_nav_outputs,js=wizard_nav_js(WIZARD_TEAM_PAGE))
    wizard_team_back_event=wizard_team_back.click(team_back_from_optional_tools,wizard_module_modes,wizard_nav_outputs)
    wizard_team_back_event.success(fn=None,js=wizard_nav_from_progress_js(),queue=False)
    wizard_team_next.click(wizard_go(WIZARD_ROUTING_PAGE),outputs=wizard_nav_outputs,js=wizard_nav_js(WIZARD_ROUTING_PAGE))
    wizard_routing_back.click(wizard_go(WIZARD_TEAM_PAGE),outputs=wizard_nav_outputs,js=wizard_nav_js(WIZARD_TEAM_PAGE))
    wizard_team_add.click(add_wizard_team_invitation,[wizard_team,wizard_team_email,wizard_team_role],[wizard_team,wizard_team_email,wizard_team_remove_choice,wizard_team_status])
    wizard_team_remove.click(remove_wizard_team_invitation,[wizard_team,wizard_team_remove_choice],[wizard_team,wizard_team_remove_choice,wizard_team_status])
    wizard_routing_event=wizard_routing_next.click(
        validate_routing_and_preview,
        [wizard_title,wizard_sponsor,wizard_mechanism,wizard_deadline,wizard_review_mode,wizard_routing,wizard_team]+wizard_module_modes,
        wizard_nav_outputs+[wizard_preview,wizard_modules,wizard_required_modules,wizard_review_required],
    )
    wizard_routing_event.success(fn=None,js=wizard_nav_js(WIZARD_PREVIEW_PAGE),queue=False)
    wizard_preview_back.click(wizard_go(WIZARD_ROUTING_PAGE),outputs=wizard_nav_outputs,js=wizard_nav_js(WIZARD_ROUTING_PAGE))
    wizard_refresh_btn.click(gateway_callback(refresh_projects),outputs=[wizard_existing])
    wizard_import_event=wizard_import_btn.click(gateway_callback(import_portable_project),[wizard_import],[wizard_existing,wizard_import_status]).then(lambda:gr.update(visible=False),outputs=[wizard_shell]).then(gateway_callback(load_project),[wizard_existing],[project_id,project_title,sponsor,mechanism,project_status,agentic_global_notice,requirements_table,section,current_version,baseline_body,preview_box,write_status,editor,current_section_key,current_competitive_update_event,section_update_banner]).then(gateway_callback(project_workflow_ui),[project_id],[workspace_workflow_summary,workspace_workflow_status,interview_tab,team_tab,clinical_tab,competitive_tab,compliance_tab,review_tab,diagnostics_tab]).then(gateway_callback(authenticated_identity),outputs=[wizard_owner_id,wizard_identity_status])
    wizard_import_event.success(fn=None,js=WIZARD_HIDE_JS,queue=False)
    wizard_open_event=wizard_open_btn.click(gateway_callback(open_project_workspace),[wizard_existing],
        [wizard_shell,project_id,project_title,sponsor,mechanism,project_status,agentic_global_notice,requirements_table,section,current_version,baseline_body,preview_box,write_status,editor,current_section_key,current_competitive_update_event,section_update_banner,workspace_workflow_summary,workspace_workflow_status,interview_tab,team_tab,clinical_tab,competitive_tab,compliance_tab,review_tab,diagnostics_tab])
    wizard_open_event.success(fn=None,js=WIZARD_HIDE_JS,queue=False)
    wizard_create_event=wizard_create_btn.click(gateway_callback(configured_project_creation_ui),
        [wizard_title,wizard_sponsor,wizard_mechanism,wizard_deadline,wizard_grant_type,wizard_source,wizard_source_url,wizard_source_text,wizard_supporting,wizard_brand,wizard_preset,wizard_modules,wizard_required_modules,wizard_review_mode,wizard_review_required,wizard_routing,wizard_team,wizard_edit_project],
        [wizard_shell,project_id,project_title,sponsor,mechanism,project_status,agentic_global_notice,requirements_table,section,current_version,baseline_body,preview_box,write_status,editor,current_section_key,current_competitive_update_event,section_update_banner,workspace_workflow_summary,workspace_workflow_status,interview_tab,team_tab,clinical_tab,competitive_tab,compliance_tab,review_tab,diagnostics_tab,wizard_create_status,wizard_create_btn])
    wizard_create_event.then(fn=None,inputs=[wizard_create_status],js=WIZARD_HIDE_AFTER_CREATE_JS,queue=False)
    editor_outputs=[editor_outline,editor_document,editor_current_key,editor_version,editor_version_choice,editor_guidance_view,editor_resolve_choice,editor_status]
    project_id.change(gateway_callback(prefill_continuous_editor),[project_id,editor_current_key],editor_outputs)
    editor_refresh_btn.click(gateway_callback(load_continuous_editor),[project_id,editor_current_key],editor_outputs)
    editor_prefill_btn.click(gateway_callback(prefill_continuous_editor),[project_id,editor_current_key],editor_outputs)
    editor_outline_command.change(gateway_callback(handle_continuous_outline_command),[project_id,editor_current_key,editor_outline_command],[editor_current_key,editor_version,editor_version_choice,editor_guidance_view,editor_resolve_choice,editor_status,editor_outline_command],show_progress="hidden")
    editor_save_btn.click(gateway_callback(save_continuous_editor),[project_id,editor_document_payload,editor_current_key],editor_outputs,js=EDITOR_SAVE_INPUT_JS)
    editor_rewrite_btn.click(gateway_callback(rewrite_continuous_section),[project_id,editor_current_key,editor_document_payload],editor_outputs,js=EDITOR_REWRITE_INPUT_JS)
    editor_add_section_btn.click(gateway_callback(add_continuous_editor_section),[project_id,editor_new_section_title,editor_document_payload,editor_current_key],editor_outputs+[editor_new_section_title],js=EDITOR_ADD_INPUT_JS)
    editor_load_version_btn.click(gateway_callback(load_continuous_historical_version),[project_id,editor_current_key,editor_version_choice],[editor_document,editor_status])
    editor_post_guidance_btn.click(gateway_callback(post_continuous_guidance),[project_id,editor_current_key,editor_guidance_kind,editor_guidance_body,editor_document_payload],editor_outputs+[editor_guidance_body,editor_guidance_status],js=EDITOR_GUIDANCE_INPUT_JS)
    editor_resolve_btn.click(gateway_callback(resolve_editor_guidance),[project_id,editor_current_key,editor_resolve_choice],[editor_guidance_view,editor_resolve_choice,editor_guidance_status])
    publish_grant_btn.click(gateway_callback(publish_document_editor),[project_id],[editor_publish_files,editor_publish_status])
    create.click(gateway_callback(create_project),[project_title,sponsor,mechanism,source,source_url,source_text,supporting,brand],[project_id,project_status,agentic_global_notice,requirements_table,section])
    compile_grant_ask_btn.click(gateway_callback(compile_grant_ask),[project_id],[requirements_table,req_status])
    approve_req.click(gateway_callback(approve_requirements),[project_id],[req_status])
    solicitation_form=[solicitation_working_title,solicitation_sponsor,solicitation_mechanism,solicitation_purpose,solicitation_facts,solicitation_criteria,solicitation_questions]
    solicitation_outputs=solicitation_form+[solicitation_version,solicitation_status,solicitation_artifact_json]
    load_solicitation_btn.click(gateway_callback(load_solicitation_form),[project_id],solicitation_outputs)
    save_solicitation_btn.click(gateway_callback(save_solicitation_form),[project_id,wizard_owner_id,solicitation_version,solicitation_artifact_json]+solicitation_form,solicitation_outputs)
    approve_solicitation_btn.click(gateway_callback(approve_solicitation_form),[project_id,wizard_owner_id,solicitation_version,solicitation_artifact_json]+solicitation_form,solicitation_outputs+[workspace_workflow_status])
    return_solicitation_btn.click(gateway_callback(lambda project,version,rationale:return_artifact_for_revision(project,"solicitation_profile",version,rationale)),[project_id,solicitation_version,solicitation_return_reason],[solicitation_status,solicitation_artifact_json,workspace_workflow_status,solicitation_return_reason])
    framework_outputs=[framework_argument,framework_nodes,framework_version,framework_status,framework_artifact_json,framework_references]
    load_framework_btn.click(gateway_callback(load_framework_form),[project_id],framework_outputs)
    generate_framework_btn.click(gateway_callback(generate_framework_form),[project_id,wizard_owner_id],framework_outputs)
    save_framework_btn.click(gateway_callback(lambda project,actor,version,metadata,argument,rows:save_framework_form(project,actor,version,metadata,argument,rows,False)),[project_id,wizard_owner_id,framework_version,framework_artifact_json,framework_argument,framework_nodes],framework_outputs)
    approve_framework_btn.click(gateway_callback(lambda project,actor,version,metadata,argument,rows:save_framework_form(project,actor,version,metadata,argument,rows,True)),[project_id,wizard_owner_id,framework_version,framework_artifact_json,framework_argument,framework_nodes],framework_outputs).then(gateway_callback(lambda project:api("GET",f"/api/projects/{project}/workflow/status")),[project_id],[workspace_workflow_status])
    return_framework_btn.click(gateway_callback(lambda project,version,rationale:return_artifact_for_revision(project,"research_framework",version,rationale)),[project_id,framework_version,framework_return_reason],[framework_status,framework_artifact_json,workspace_workflow_status,framework_return_reason])
    aims_outputs=[aims_objective,aims_hypothesis,core_aims,aims_version,aims_status,aims_artifact_json,aims_references]
    load_aims_btn.click(gateway_callback(load_aims_form),[project_id],aims_outputs)
    generate_aims_btn.click(gateway_callback(generate_aims_form),[project_id,wizard_owner_id],aims_outputs)
    save_aims_btn.click(gateway_callback(lambda project,actor,version,metadata,objective,hypothesis,rows:save_aims_form(project,actor,version,metadata,objective,hypothesis,rows,False)),[project_id,wizard_owner_id,aims_version,aims_artifact_json,aims_objective,aims_hypothesis,core_aims],aims_outputs)
    approve_aims_btn.click(gateway_callback(lambda project,actor,version,metadata,objective,hypothesis,rows:save_aims_form(project,actor,version,metadata,objective,hypothesis,rows,True)),[project_id,wizard_owner_id,aims_version,aims_artifact_json,aims_objective,aims_hypothesis,core_aims],aims_outputs).then(gateway_callback(lambda project:api("GET",f"/api/projects/{project}/workflow/status")),[project_id],[workspace_workflow_status])
    return_aims_btn.click(gateway_callback(lambda project,version,rationale:return_artifact_for_revision(project,"aim_set",version,rationale)),[project_id,aims_version,aims_return_reason],[aims_status,aims_artifact_json,workspace_workflow_status,aims_return_reason])
    generate_q.click(gateway_callback(generate_interview),[project_id],[interview_questions,current_question,question_card,answer,interview_status,interview_table])
    submit.click(gateway_callback(submit_answer),[project_id,interview_questions,current_question,answer,confidence,classification,answer_notes,answered_by],[interview_questions,current_question,question_card,answer,interview_status])
    search_plan_outputs=[search_plan_upstream,search_plan_queries,search_plan_version,search_plan_status,search_plan_artifact_json,search_plan_references]
    generate_search_plan_btn.click(gateway_callback(generate_search_plan_form),[project_id,max_queries],search_plan_outputs)
    load_search_plan_btn.click(gateway_callback(load_search_plan_form),[project_id],search_plan_outputs)
    save_search_plan_btn.click(gateway_callback(lambda project,actor,version,metadata,queries:save_search_plan_form(project,actor,version,metadata,queries,False)),[project_id,wizard_owner_id,search_plan_version,search_plan_artifact_json,search_plan_queries],search_plan_outputs)
    approve_search_plan_btn.click(gateway_callback(lambda project,actor,version,metadata,queries:save_search_plan_form(project,actor,version,metadata,queries,True)),[project_id,wizard_owner_id,search_plan_version,search_plan_artifact_json,search_plan_queries],search_plan_outputs)
    return_search_plan_btn.click(gateway_callback(lambda project,version,rationale:return_artifact_for_revision(project,"literature_search_plan",version,rationale)),[project_id,search_plan_version,search_plan_return_reason],[search_plan_status,search_plan_artifact_json,workspace_workflow_status,search_plan_return_reason])
    literature_outputs=[literature_manifest_summary,literature_queries,literature_needs,literature_contradictions,literature_version,literature_artifact_status,literature_artifact_json,literature_references]
    research_btn.click(gateway_callback(run_research),[project_id,search_plan_version,results_per],[evidence_table,research_status]).then(gateway_callback(load_literature_form),[project_id],literature_outputs)
    load_literature_btn.click(gateway_callback(load_literature_form),[project_id],literature_outputs)
    save_literature_btn.click(gateway_callback(lambda project,actor,version,metadata,queries,needs,contradictions:save_literature_form(project,actor,version,metadata,queries,needs,contradictions,False)),[project_id,wizard_owner_id,literature_version,literature_artifact_json,literature_queries,literature_needs,literature_contradictions],literature_outputs)
    approve_literature_btn.click(gateway_callback(lambda project,actor,version,metadata,queries,needs,contradictions:save_literature_form(project,actor,version,metadata,queries,needs,contradictions,True)),[project_id,wizard_owner_id,literature_version,literature_artifact_json,literature_queries,literature_needs,literature_contradictions],literature_outputs).then(gateway_callback(lambda project:api("GET",f"/api/projects/{project}/workflow/status")),[project_id],[workspace_workflow_status])
    return_literature_btn.click(gateway_callback(lambda project,version,rationale:return_artifact_for_revision(project,"literature_manifest",version,rationale)),[project_id,literature_version,literature_return_reason],[literature_artifact_status,literature_artifact_json,workspace_workflow_status,literature_return_reason])
    team_workspace_outputs=[team_members,team_activity,project_invites,team_tasks,team_notifications,approval_routing_table,project_health_table,project_health_summary_md,task_owner,team_mentions,active_project_task,task_dependencies,active_project_invite,unread_notification_id,team_permissions,team_workspace_status,create_project_invite_btn,add_existing_member_btn,post_team_channel_btn,create_project_task_btn]
    refresh_team_workspace_btn.click(gateway_callback(load_team_workspace),[project_id],team_workspace_outputs)
    load_team_channel_btn.click(gateway_callback(load_team_channel),[project_id,team_channel_kind,team_channel_subject],[team_channel_html,team_messages,team_reply_message_id,team_channel_status])
    post_team_channel_btn.click(gateway_callback(post_team_channel_message),[project_id,team_channel_kind,team_channel_subject,team_message_body,team_reply_message_id,team_mentions],[team_channel_html,team_messages,team_reply_message_id,team_channel_status,team_message_body])
    create_project_invite_btn.click(gateway_callback(create_project_invite_ui),[project_id,project_invite_email,project_invite_role,project_invite_days],[project_invite_status]+team_workspace_outputs)
    revoke_project_invite_btn.click(gateway_callback(revoke_project_invite_ui),[project_id,active_project_invite],[project_invite_status]+team_workspace_outputs)
    add_existing_member_btn.click(gateway_callback(add_existing_project_member_ui),[project_id,existing_member_user_id,existing_member_role],[existing_member_status]+team_workspace_outputs+[existing_member_user_id])
    create_project_task_btn.click(gateway_callback(create_project_task_ui),[project_id,task_title,task_description,task_owner,task_priority,task_due,task_dependencies],[task_action_status]+team_workspace_outputs+[task_title,task_description])
    update_project_task_btn.click(gateway_callback(update_project_task_ui),[project_id,active_project_task,project_task_status],[task_action_status]+team_workspace_outputs)
    load_comments_btn.click(gateway_callback(load_artifact_comments),[project_id,comment_target_type,comment_target_key,comment_version_id],[artifact_comments,comment_parent_id,comment_resolve_id,comment_status])
    current_comment_version_btn.click(gateway_callback(current_comment_target_version),[project_id,comment_target_type,comment_target_key],[comment_version_id,comment_status])
    post_comment_btn.click(gateway_callback(post_artifact_comment),[project_id,comment_target_type,comment_target_key,comment_version_id,comment_start,comment_end,comment_quote,comment_body,comment_parent_id,team_mentions],[artifact_comments,comment_parent_id,comment_resolve_id,comment_status,comment_body])
    resolve_comment_btn.click(gateway_callback(resolve_artifact_comment),[project_id,comment_target_type,comment_target_key,comment_version_id,comment_resolve_id],[artifact_comments,comment_parent_id,comment_resolve_id,comment_status])
    mark_notification_read_btn.click(gateway_callback(mark_notification_read_ui),[project_id,unread_notification_id],[notification_status]+team_workspace_outputs)
    refresh_reviewer_roles_btn.click(gateway_callback(reviewer_role_rows),[project_id],[review_roles,review_notice])
    create_review_plan_btn.click(gateway_callback(create_panel_plan),[project_id,review_mode],[review_plan_id,review_roles,review_status,review_plan_json])
    approve_review_plan_btn.click(gateway_callback(approve_panel_plan),[project_id,review_plan_id],[review_status,review_plan_json])
    run_review_panel_btn.click(gateway_callback(execute_review_panel),[project_id,review_plan_id],[review_run_id,panel_summary,reviewer_findings,revision_findings,causal_editor,review_status,review_run_json])
    approve_review_result_btn.click(gateway_callback(approve_review_result),[project_id,review_run_id],[review_status]).then(gateway_callback(lambda project:api("GET",f"/api/projects/{project}/workflow/status")),[project_id],[workspace_workflow_status])
    create_revision_tasks_btn.click(gateway_callback(create_revision_tasks),[project_id,review_run_id,revision_task_indexes,revision_task_owner,revision_task_due],[revision_task_result,revision_task_status])
    load_causal_btn.click(gateway_callback(load_causal_models),[project_id,review_run_id],[causal_editor,causal_history,causal_status])
    save_causal_btn.click(gateway_callback(save_causal_editor),[project_id,review_run_id,causal_editor,confirm_causal],[causal_editor,causal_status])
    rebuild_idx.click(gateway_callback(rebuild_index),[project_id],[index_message,index_json]);status_idx.click(gateway_callback(index_status),[project_id],[index_message,index_json]);retrieval_btn.click(gateway_callback(test_retrieval),[project_id,retrieval_query,retrieval_k],[retrieval_table])
    load_clinical_btn.click(gateway_callback(load_clinical_study),[project_id],[clinical_problem,knowledge_gap,central_hypothesis,disease,disease_stage,biomarker,inclusion,exclusion,design_type,study_phase,randomization,allocation_ratio,blinding,follow_up_months,design_sites,available_patients,eligibility_pct,biomarker_pct,consent_pct,target_enrollment,accrual_months,recruitment_sites,test_type,alpha,power,attrition_pct,control_rate,treatment_rate,null_rate,alternative_rate,mean_delta,std_dev,hazard_ratio,event_probability,aims_table,arms_table,endpoints_table,timeline_table,resources_table,clinical_assessment,clinical_status])
    save_clinical_btn.click(gateway_callback(save_clinical_study),[project_id]+[clinical_problem,knowledge_gap,central_hypothesis,disease,disease_stage,biomarker,inclusion,exclusion,design_type,study_phase,randomization,allocation_ratio,blinding,follow_up_months,design_sites,available_patients,eligibility_pct,biomarker_pct,consent_pct,target_enrollment,accrual_months,recruitment_sites,test_type,alpha,power,attrition_pct,control_rate,treatment_rate,null_rate,alternative_rate,mean_delta,std_dev,hazard_ratio,event_probability,aims_table,arms_table,endpoints_table,timeline_table,resources_table],[clinical_assessment,clinical_status])
    calculate_n_btn.click(gateway_callback(calculate_sample_size),[project_id,test_type,alpha,power,attrition_pct,control_rate,treatment_rate,null_rate,alternative_rate,mean_delta,std_dev,hazard_ratio,event_probability],[sample_size_json,sample_size_status])
    scenario_btn.click(gateway_callback(run_feasibility_scenarios),[project_id,scenario_sites,scenario_consent,scenario_biomarker],[scenario_table,scenario_status])
    generate_comp_profile_btn.click(gateway_callback(generate_competitive_profile),[project_id],[competitive_profile_json,competitive_status])
    load_competitive_btn.click(gateway_callback(load_competitive),[project_id],[competitive_profile_json,competitor_table,asset_table,provider_status_json,competitive_strategy,competitive_raw,competitive_status,agentic_global_notice])
    run_competitive_btn.click(gateway_callback(run_competitive_intelligence),[project_id],[competitor_table,asset_table,provider_status_json,competitive_strategy,competitive_raw,competitive_status,agentic_global_notice])
    load_compliance_btn.click(gateway_callback(load_compliance),[project_id],[compliance_profile_state,opportunity_source_preview,compliance_sponsor,compliance_mechanism,submission_system,compliance_deadline,compliance_rules,compliance_rule_id,compliance_provenance,compliance_findings,compliance_json,artifact_table,compliance_status])
    compile_compliance_btn.click(gateway_callback(compile_compliance),[project_id],[compliance_profile_state,opportunity_source_preview,compliance_sponsor,compliance_mechanism,submission_system,compliance_deadline,compliance_rules,compliance_rule_id,compliance_provenance,compliance_findings,compliance_json,compliance_status,section])
    save_compliance_btn.click(gateway_callback(save_compliance),[project_id,compliance_profile_state,compliance_sponsor,compliance_mechanism,submission_system,compliance_deadline,compliance_rules],[compliance_profile_state,compliance_rules,compliance_rule_id,compliance_provenance,compliance_findings,compliance_json,compliance_status,section])
    approve_compliance_btn.click(gateway_callback(approve_compliance),[project_id],[compliance_provenance,compliance_findings,compliance_json,compliance_status])
    resolve_compliance_btn.click(gateway_callback(resolve_compliance),[project_id,compliance_rule_id,compliance_resolution,compliance_resolution_notes,compliance_resolved_by],[compliance_findings,compliance_json,compliance_status])
    register_artifact_btn.click(gateway_callback(register_submission_artifacts),[project_id,artifact_slot,artifact_files],[artifact_table,compliance_findings,compliance_json,compliance_status])
    measure_compliance_btn.click(gateway_callback(measure_compliance),[project_id],[compliance_findings,compliance_json,compliance_status])
    section.change(gateway_callback(load_section),[project_id,project_title,section],[current_version,baseline_body,preview_box,write_status,editor,current_section_key,current_competitive_update_event,section_update_banner]).then(gateway_callback(version_history),[project_id,current_section_key],[version_history_table,compare_from_version,compare_to_version,restore_version_id])
    gen.click(gateway_callback(draft_section),[project_id,project_title,section,additional,high],[current_version,baseline_body,current_section_key,preview_box,write_status,editor,save_btn,cancel,current_competitive_update_event,section_update_banner]).then(gateway_callback(version_history),[project_id,current_section_key],[version_history_table,compare_from_version,compare_to_version,restore_version_id])
    edit.click(show_editor,[baseline_body],[editor,save_btn,cancel])
    save_btn.click(gateway_callback(save_edit),[project_id,project_title,section,current_section_key,current_version,baseline_body,editor],[current_version,baseline_body,preview_box,write_status,editor,save_btn,cancel,current_competitive_update_event,section_update_banner]).then(gateway_callback(version_history),[project_id,current_section_key],[version_history_table,compare_from_version,compare_to_version,restore_version_id])
    cancel.click(gateway_callback(cancel_edit),[project_id,project_title,section,current_section_key,current_version,baseline_body],[preview_box,editor,save_btn,cancel,write_status])
    approve_btn.click(gateway_callback(approve_section),[project_id,project_title,section,current_section_key,current_version,baseline_body,editor,current_competitive_update_event],[current_version,baseline_body,preview_box,write_status,editor,save_btn,cancel,current_competitive_update_event,section_update_banner]).then(gateway_callback(version_history),[project_id,current_section_key],[version_history_table,compare_from_version,compare_to_version,restore_version_id])
    return_section_btn.click(gateway_callback(return_section_for_revision),[project_id,current_section_key,current_version,section_return_reason],[write_status,workspace_workflow_status,section_return_reason]).then(gateway_callback(version_history),[project_id,current_section_key],[version_history_table,compare_from_version,compare_to_version,restore_version_id])
    refresh_version_history_btn.click(gateway_callback(version_history),[project_id,current_section_key],[version_history_table,compare_from_version,compare_to_version,restore_version_id])
    compare_versions_btn.click(gateway_callback(compare_versions),[project_id,current_section_key,compare_from_version,compare_to_version],[version_diff,version_history_status])
    restore_version_btn.click(gateway_callback(restore_version),[project_id,project_title,section,current_section_key,current_version,restore_version_id],[current_version,baseline_body,preview_box,write_status,editor,current_section_key,current_competitive_update_event,section_update_banner,version_history_table,compare_from_version,compare_to_version,restore_version_id])
    preview_approved_btn.click(gateway_callback(preview_approved_grant),[project_id],[approved_sections_table,approved_grant_preview,approved_preview_status])
    portable_export_btn.click(gateway_callback(export_portable_project),[project_id],[portable_export_file,portable_export_status])
    check_ready.click(gateway_callback(readiness),[project_id],[readiness_json,readiness_status,fmt,export_btn]);export_btn.click(gateway_callback(export),[project_id,fmt,project_title],[export_file,export_status])
    refresh_projects_btn.click(gateway_callback(refresh_projects),outputs=[recent])
    open_project_btn.click(gateway_callback(load_project),[recent],[project_id,project_title,sponsor,mechanism,project_status,agentic_global_notice,requirements_table,section,current_version,baseline_body,preview_box,write_status,editor,current_section_key,current_competitive_update_event,section_update_banner]).then(gateway_callback(project_workflow_ui),[project_id],[workspace_workflow_summary,workspace_workflow_status,interview_tab,team_tab,clinical_tab,competitive_tab,compliance_tab,review_tab,diagnostics_tab])
    refresh_project_manager_btn.click(gateway_callback(project_catalog),[include_archived_projects],[recent,project_catalog_table,project_catalog_summary])
    include_archived_projects.change(gateway_callback(project_catalog),[include_archived_projects],[recent,project_catalog_table,project_catalog_summary])
    apply_project_management_btn.click(gateway_callback(update_project_manager),[recent,managed_project_title,managed_project_action,include_archived_projects],[recent,project_catalog_table,project_catalog_summary,project_management_status,managed_project_title,managed_project_action])
    system_info_btn.click(gateway_callback(system_diagnostics),outputs=[system_info_json,diagnostics_status])
    hpc_bench_btn.click(gateway_callback(run_hpc_diagnostics),outputs=[hpc_benchmark_json,diagnostics_status])
    refresh_accounts_btn.click(gateway_callback(account_rows),outputs=[accounts_table])
    create_account_btn.click(gateway_callback(create_account),[account_username,account_email,account_display_name,account_temp_password],[accounts_table,account_admin_status,account_temp_password])
    send_account_reset_btn.click(gateway_callback(send_account_reset),[account_target_id],[account_admin_status])
    disable_account_btn.click(gateway_callback(lambda user_id:set_account_status(user_id,False)),[account_target_id],[accounts_table,account_admin_status])
    enable_account_btn.click(gateway_callback(lambda user_id:set_account_status(user_id,True)),[account_target_id],[accounts_table,account_admin_status])
    refresh_shared_updates_btn.click(gateway_callback(refresh_shared_updates),[project_id,team_channel_kind,team_channel_subject,search_plan_version,literature_version,include_archived_projects],
        [wizard_existing,recent,project_catalog_table,project_catalog_summary,team_members,team_activity,project_invites,team_tasks,team_notifications,approval_routing_table,project_health_table,project_health_summary_md,team_workspace_status,team_channel_html,team_messages,team_channel_status,shared_literature_sync_status,requirements_table,req_status,workspace_workflow_summary,workspace_workflow_status,interview_tab,team_tab,clinical_tab,competitive_tab,compliance_tab,review_tab,diagnostics_tab,manual_refresh_status])
    session_bootstrap_event=demo.load(fn=None,js=SESSION_STORAGE_AUTH_JS,queue=False)
    demo.load(fn=None,js=WIZARD_CREATE_CLICK_JS,queue=False)
    demo.load(fn=None,js=EDITOR_OUTLINE_JS,queue=False)
    demo.load(fn=None,js=GLOBAL_NAVIGATION_JS,queue=False)
    session_bootstrap_event.then(gateway_callback(project_catalog),[include_archived_projects],[recent,project_catalog_table,project_catalog_summary],show_progress="hidden")
    session_bootstrap_event.then(gateway_callback(refresh_projects),outputs=[wizard_existing],show_progress="hidden")
    session_bootstrap_event.then(gateway_callback(authenticated_identity),outputs=[wizard_owner_id,wizard_identity_status],show_progress="hidden")

APP_THEME=gr.themes.Base(primary_hue=gr.themes.colors.purple,secondary_hue=gr.themes.colors.fuchsia,neutral_hue=gr.themes.colors.gray)

def _auth_api(method,path,token=None,payload=None):
    headers={"Idempotency-Key":str(uuid.uuid4())} if method.upper() not in {"GET","HEAD","OPTIONS"} else {}
    if token:headers["Authorization"]=f"Bearer {token}"
    response=requests.request(method,f"{CORE}{path}",json=payload,headers=headers,timeout=20)
    data=response.json() if response.headers.get("content-type","").startswith("application/json") else {}
    if not response.ok:raise RuntimeError(data.get("error") or "Authentication request failed.")
    return data

AUTH_STYLE="""
body{margin:0;background:#0d0b14;color:#f7f3fb;font:15px/1.5 system-ui,-apple-system,sans-serif}.auth{max-width:460px;margin:8vh auto;padding:34px;background:#17131f;border:1px solid #342a40;border-radius:18px;box-shadow:0 24px 70px #0008}h1{font:38px/1.1 Georgia,serif;margin:0 0 12px}.muted{color:#aaa0b5}.error{background:#421e2b;border:1px solid #82374e;padding:12px;border-radius:9px}label{display:block;margin:15px 0 6px}input{box-sizing:border-box;width:100%;padding:12px;border:1px solid #463951;border-radius:8px;background:#100d17;color:#fff}button{width:100%;margin-top:20px;padding:12px;border:0;border-radius:8px;background:linear-gradient(100deg,#7c3aed,#d946ef);color:#fff;font-weight:700;cursor:pointer}a{color:#c77dff}.links{display:flex;justify-content:space-between;margin-top:20px}.brand{color:#c06cf4;font-size:12px;letter-spacing:.16em;text-transform:uppercase}
"""

def _auth_page(title,fields,action,submit,error=None,help_text=None,csrf=None):
    csrf=csrf or secrets.token_urlsafe(32)
    message=f'<p class="error">{html.escape(str(error))}</p>' if error else ""
    help_html=f'<p class="muted">{html.escape(help_text)}</p>' if help_text else ""
    body=f'''<!doctype html><html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width"><title>{html.escape(title)} · Grantspace</title><style>{AUTH_STYLE}</style></head><body><main class="auth"><div class="brand">Grantspace</div><h1>{html.escape(title)}</h1>{help_html}{message}<form method="post" action="{action}"><input type="hidden" name="csrf" value="{html.escape(csrf)}">{fields}<button type="submit">{html.escape(submit)}</button></form><div class="links"><a href="/login">Sign in</a><a href="/forgot-password">Forgot password?</a></div></main></body></html>'''
    response=HTMLResponse(body)
    response.set_cookie("grantspace_csrf",csrf,httponly=True,secure=_secure_cookie(),samesite="lax",max_age=3600,path="/")
    return response

def _field(name,label,password=False,autocomplete=None,value=""):
    kind="password" if password else "text"
    return f'<label for="{name}">{html.escape(label)}</label><input required id="{name}" name="{name}" type="{kind}" autocomplete="{autocomplete or ("current-password" if password else "off")}" value="{html.escape(value)}">'

def _secure_cookie():
    explicit=os.getenv("SESSION_COOKIE_SECURE")
    if explicit is not None:return explicit.strip().lower() in {"1","true","yes","on"}
    return os.getenv("APP_PUBLIC_URL","").startswith("https://")

async def _form(request):
    raw=(await request.body()).decode("utf-8",errors="strict")
    return {key:values[-1] for key,values in parse_qs(raw,keep_blank_values=True).items()}

def _csrf_valid(request,form):
    supplied=form.get("csrf","");stored=request.cookies.get("grantspace_csrf","")
    return bool(supplied and stored and secrets.compare_digest(supplied,stored))

def _session_user(request):
    token=request.cookies.get("grantspace_session")
    if not token:return None
    try:return _auth_api("GET","/api/me",token=token)
    except Exception:return None

def _set_session(response,token):
    response.set_cookie("grantspace_session",token,httponly=True,secure=_secure_cookie(),samesite="lax",max_age=int(os.getenv("AUTH_SESSION_TTL_SECONDS","43200")),path="/")
    return response

def _login_fields(username=""):
    return _field("username","Username",value=username,autocomplete="username")+_field("password","Password",password=True,autocomplete="current-password")

def build_internal_account_app():
    web=FastAPI(title="Grantspace")

    @web.get("/manifest.json",include_in_schema=False)
    async def web_manifest():
        return JSONResponse({"name":"Grantspace","short_name":"Grantspace","start_url":"/","scope":"/","display":"standalone","background_color":"#0d0b14","theme_color":"#0d0b14","icons":[]})

    @web.get("/favicon.ico",include_in_schema=False)
    async def favicon():return Response(status_code=204)

    @web.get("/session-token",include_in_schema=False)
    async def session_token(request:Request):
        token=request.cookies.get("grantspace_session")
        if not token or not _session_user(request):return JSONResponse({"error":"login session is missing or expired"},status_code=401)
        return JSONResponse({"access_token":token},headers={"Cache-Control":"no-store"})

    @web.get("/")
    async def root(request:Request):
        try:
            if _auth_api("GET","/api/auth/bootstrap/status").get("bootstrap_required"):return RedirectResponse("/setup",303)
        except Exception:return RedirectResponse("/login",303)
        user=_session_user(request)
        if not user:return RedirectResponse("/login",303)
        if user.get("must_change_password"):return RedirectResponse("/change-password",303)
        if request.cookies.get("grantspace_pending_invite"):return RedirectResponse("/invite",303)
        return RedirectResponse("/app/",303)

    @web.get("/setup")
    async def setup_get(request:Request):
        if not _auth_api("GET","/api/auth/bootstrap/status").get("bootstrap_required"):return RedirectResponse("/login",303)
        fields=_field("setup_token","Initial setup token",password=True,autocomplete="off")+_field("username","Administrator username",autocomplete="username")+_field("email","Administrator email",autocomplete="email")+_field("display_name","Display name")+_field("temporary_password","Temporary password",password=True,autocomplete="new-password")
        return _auth_page("Create the first administrator",fields,"/setup","Create the only bootstrap account",help_text="This one-time screen closes permanently after the database transaction succeeds. You must replace the temporary password on first login.")

    @web.post("/setup")
    async def setup_post(request:Request):
        form=await _form(request)
        if not _csrf_valid(request,form):return _auth_page("Create the first administrator","","/setup","Try again",error="The form expired. Reload the setup page.")
        try:
            payload={key:form.get(key,"") for key in ("setup_token","username","email","display_name","temporary_password")}
            _auth_api("POST","/api/auth/bootstrap",payload=payload)
            login=_auth_api("POST","/api/auth/login",payload={"username":form.get("username","") ,"password":form.get("temporary_password","")})
            return _set_session(RedirectResponse("/change-password",303),login["access_token"])
        except Exception as error:
            fields=_field("setup_token","Initial setup token",password=True)+_field("username","Administrator username",value=form.get("username",""))+_field("email","Administrator email",value=form.get("email",""))+_field("display_name","Display name",value=form.get("display_name",""))+_field("temporary_password","Temporary password",password=True)
            return _auth_page("Create the first administrator",fields,"/setup","Create the only bootstrap account",error=error)

    @web.get("/login")
    async def login_get(request:Request):
        if _auth_api("GET","/api/auth/bootstrap/status").get("bootstrap_required"):return RedirectResponse("/setup",303)
        user=_session_user(request)
        if user:return RedirectResponse("/change-password" if user.get("must_change_password") else "/app/",303)
        return _auth_page("Sign in",_login_fields(),"/login","Sign in",help_text="Use the username assigned by your administrator.")

    @web.post("/login")
    async def login_post(request:Request):
        form=await _form(request)
        if not _csrf_valid(request,form):return _auth_page("Sign in",_login_fields(form.get("username","")),"/login","Sign in",error="The form expired. Reload and try again.")
        try:
            result=_auth_api("POST","/api/auth/login",payload={"username":form.get("username","") ,"password":form.get("password","")})
            destination="/change-password" if result.get("user",{}).get("must_change_password") else ("/invite" if request.cookies.get("grantspace_pending_invite") else "/app/")
            return _set_session(RedirectResponse(destination,303),result["access_token"])
        except Exception as error:return _auth_page("Sign in",_login_fields(form.get("username","")),"/login","Sign in",error=error)

    @web.get("/change-password")
    async def change_get(request:Request):
        if not _session_user(request):return RedirectResponse("/login",303)
        fields=_field("current_password","Current or temporary password",password=True)+_field("new_password","New password",password=True,autocomplete="new-password")+_field("confirm_password","Confirm new password",password=True,autocomplete="new-password")
        return _auth_page("Choose a new password",fields,"/change-password","Save password",help_text="At least 14 characters and three character classes. It cannot contain your username or email name.")

    @web.post("/change-password")
    async def change_post(request:Request):
        token=request.cookies.get("grantspace_session")
        if not token:return RedirectResponse("/login",303)
        form=await _form(request);fields=_field("current_password","Current or temporary password",password=True)+_field("new_password","New password",password=True)+_field("confirm_password","Confirm new password",password=True)
        if not _csrf_valid(request,form):return _auth_page("Choose a new password",fields,"/change-password","Save password",error="The form expired. Reload and try again.")
        if form.get("new_password")!=form.get("confirm_password"):return _auth_page("Choose a new password",fields,"/change-password","Save password",error="New passwords do not match.")
        try:_auth_api("POST","/api/auth/change-password",token=token,payload={"current_password":form.get("current_password","") ,"new_password":form.get("new_password","")})
        except Exception as error:return _auth_page("Choose a new password",fields,"/change-password","Save password",error=error)
        return RedirectResponse("/invite" if request.cookies.get("grantspace_pending_invite") else "/app/",303)

    @web.get("/invite")
    async def invite_get(request:Request,token:str=""):
        pending=(token or request.cookies.get("grantspace_pending_invite") or "").strip()
        if not pending:return _auth_page("Project invitation","","/invite","Accept invitation",error="The invitation token is missing. Reopen the complete invitation link.")
        user=_session_user(request)
        if not user:
            response=RedirectResponse("/login",303);response.set_cookie("grantspace_pending_invite",pending,httponly=True,secure=_secure_cookie(),samesite="lax",max_age=30*86400,path="/");return response
        if user.get("must_change_password"):
            response=RedirectResponse("/change-password",303);response.set_cookie("grantspace_pending_invite",pending,httponly=True,secure=_secure_cookie(),samesite="lax",max_age=30*86400,path="/");return response
        fields=f'<input type="hidden" name="token" value="{html.escape(pending)}"><p class="muted">Signed in as {html.escape(str(user.get("email") or user.get("username") or user.get("id")))}. The invitation will only succeed when this authenticated email matches the invited address.</p>'
        response=_auth_page("Accept project invitation",fields,"/invite","Accept invitation")
        response.set_cookie("grantspace_pending_invite",pending,httponly=True,secure=_secure_cookie(),samesite="lax",max_age=30*86400,path="/");return response

    @web.post("/invite")
    async def invite_post(request:Request):
        token=request.cookies.get("grantspace_session")
        if not token:return RedirectResponse("/login",303)
        form=await _form(request);raw_token=(form.get("token") or request.cookies.get("grantspace_pending_invite") or "").strip()
        if not _csrf_valid(request,form):return RedirectResponse(f"/invite?token={quote(raw_token)}",303)
        try:_auth_api("POST","/api/invites/accept",token=token,payload={"token":raw_token})
        except Exception as error:
            fields=f'<input type="hidden" name="token" value="{html.escape(raw_token)}">';return _auth_page("Accept project invitation",fields,"/invite","Accept invitation",error=error)
        response=RedirectResponse("/app/",303);response.delete_cookie("grantspace_pending_invite",path="/");return response

    @web.get("/forgot-password")
    async def forgot_get(request:Request):return _auth_page("Reset your password",_field("login","Username or email",autocomplete="username"),"/forgot-password","Email reset link")

    @web.post("/forgot-password")
    async def forgot_post(request:Request):
        form=await _form(request)
        if not _csrf_valid(request,form):return _auth_page("Reset your password",_field("login","Username or email"),"/forgot-password","Email reset link",error="The form expired. Reload and try again.")
        try:_auth_api("POST","/api/auth/password-reset/request",payload={"login":form.get("login","")})
        except Exception as error:return _auth_page("Reset your password",_field("login","Username or email"),"/forgot-password","Email reset link",error=error)
        return _auth_page("Check your email","","/forgot-password","Send another link",help_text="If an active account matched, a single-use reset link was sent.")

    @web.get("/password-reset")
    async def reset_get(request:Request,token:str=""):
        fields=f'<input type="hidden" name="token" value="{html.escape(token)}">'+_field("new_password","New password",password=True,autocomplete="new-password")+_field("confirm_password","Confirm new password",password=True,autocomplete="new-password")
        return _auth_page("Set a new password",fields,"/password-reset","Reset password")

    @web.post("/password-reset")
    async def reset_post(request:Request):
        form=await _form(request);token=form.get("token","");fields=f'<input type="hidden" name="token" value="{html.escape(token)}">'+_field("new_password","New password",password=True)+_field("confirm_password","Confirm new password",password=True)
        if not _csrf_valid(request,form):return _auth_page("Set a new password",fields,"/password-reset","Reset password",error="The form expired. Reopen the email link.")
        if form.get("new_password")!=form.get("confirm_password"):return _auth_page("Set a new password",fields,"/password-reset","Reset password",error="New passwords do not match.")
        try:_auth_api("POST","/api/auth/password-reset/confirm",payload={"token":token,"new_password":form.get("new_password","")})
        except Exception as error:return _auth_page("Set a new password",fields,"/password-reset","Reset password",error=error)
        return RedirectResponse("/login",303)

    @web.get("/logout")
    async def logout(request:Request):
        token=request.cookies.get("grantspace_session")
        if token:
            try:_auth_api("POST","/api/auth/logout",token=token,payload={})
            except Exception:pass
        response=RedirectResponse("/login",303);response.delete_cookie("grantspace_session",path="/");return response

    def authorize_gradio(request:Request):
        user=_session_user(request)
        return (user.get("username") or user.get("id")) if user and not user.get("must_change_password") else None

    return gr.mount_gradio_app(web,demo,path="/app",auth_dependency=authorize_gradio,css=CSS,theme=APP_THEME)

if __name__=="__main__":
    if os.getenv("AUTH_MODE","local_single_user")=="internal_accounts":
        import uvicorn
        uvicorn.run(build_internal_account_app(),host="0.0.0.0",port=7860,proxy_headers=True,forwarded_allow_ips=os.getenv("TRUSTED_PROXY_IPS","127.0.0.1"))
    else:demo.launch(server_name="0.0.0.0",server_port=7860,css=CSS,theme=APP_THEME)
