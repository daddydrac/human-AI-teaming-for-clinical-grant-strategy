import base64
import difflib
import hashlib
import html as htmlmod
import json
import os
import re
import io
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from fastapi import FastAPI, HTTPException
from PIL import Image
from pydantic import BaseModel, Field
from pypdf import PdfReader
from weasyprint import HTML

app = FastAPI()
PROJECT_ROOT = Path(os.getenv("PROJECT_ROOT", "/workspace/projects"))
OUTPUT_ROOT = Path(os.getenv("OUTPUT_ROOT", "/exports"))
CONFIG_ROOT = Path(os.getenv("RENDERER_CONFIG_ROOT", str(Path(__file__).resolve().parents[1] / "config")))
COMPETITIVE_UPDATE_LABEL = os.getenv("COMPETITIVE_UPDATE_LABEL", "Competitive Edge Auto-Update").strip() or "Competitive Edge Auto-Update"


class Section(BaseModel):
    section_key: Optional[str] = None
    title: str
    body: str
    version: Optional[int] = None


class RenderRequest(BaseModel):
    project_id: str
    snapshot_id: Optional[int] = None
    format: Optional[str] = None
    title: str
    sponsor: Optional[str] = None
    organization_name: Optional[str] = None
    sections: List[Section]
    include_document_title: bool = True
    design_profile: Optional[Dict[str, Any]] = None


class DesignProfileRequest(BaseModel):
    project_id: str
    sponsor: Optional[str] = None
    organization_name: Optional[str] = None
    asset_paths: List[str] = Field(default_factory=list)


class PackageRequest(BaseModel):
    project_id: str
    snapshot_id: int
    title: str
    generated_paths: List[str] = Field(default_factory=list)
    manifest: Dict[str, Any] = Field(default_factory=dict)

class DiffPreviewRequest(BaseModel):
    project_id: str
    title: str
    organization_name: Optional[str] = None
    section: Section
    baseline_body: str
    update_summary: Optional[str] = None
    update_reason: Optional[str] = None
    includes_human_edits: bool = False
    design_profile: Optional[Dict[str, Any]] = None


def safe(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", value).strip("_") or "grant"


def project_dir(project_id: str) -> Path:
    # Renderer is reachable independently of the Rust core, so it must not trust
    # path-bearing input even though normal project IDs are UUIDs.
    if not re.fullmatch(r"[A-Za-z0-9_-]{8,128}", project_id or ""):
        raise HTTPException(400, "invalid project_id")
    return PROJECT_ROOT / project_id


def load_json(path: Path, default):
    try:
        return json.loads(path.read_text())
    except Exception:
        return default


def default_profile() -> Dict[str, Any]:
    built_in = {
        "organization_name": "Organization", "body_font": "Arial", "body_size_pt": 11.0,
        "heading_font": "Arial", "heading_size_pt": 16.0, "document_title_size_pt": 20.0,
        "brand_line_size_pt": 9.0, "primary_color": "#333333", "text_color": "#222222",
        "page_width_in": 8.5, "page_height_in": 11.0, "margin_top_in": 0.5,
        "margin_bottom_in": 0.5, "margin_left_in": 0.5, "margin_right_in": 0.5,
        "line_height": 1.18, "paragraph_spacing_pt": 6.0, "logo_path": None,
    }
    profile = {**built_in, **load_json(CONFIG_ROOT / "default_design.json", {})}
    profile["organization_name"] = os.getenv(
        "ORGANIZATION_NAME", profile.get("organization_name", "Organization")
    )
    return profile


def norm_sponsor(value: Optional[str]) -> str:
    return (value or "").strip().lower()


def css_font(value: Any) -> str:
    value = re.sub(r"[^A-Za-z0-9 _.-]", "", str(value or "Arial")).strip()
    return value or "Arial"


def css_color(value: Any, default: str = "#333333") -> str:
    value = str(value or "")
    return value if re.fullmatch(r"#[0-9A-Fa-f]{6}", value) else default


def hex_rgb(value: str) -> RGBColor:
    value = css_color(value).lstrip("#")
    return RGBColor(int(value[:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def dominant_hex(path: Path) -> Optional[str]:
    try:
        image = Image.open(path).convert("RGB")
        image.thumbnail((128, 128))
        colors = image.getcolors(maxcolors=128 * 128) or []
        if not colors:
            return None
        # White/black backgrounds are common in logos and should not silently
        # become the primary brand color when a chromatic mark is present.
        usable = [
            (count, rgb)
            for count, rgb in colors
            if not all(channel >= 245 for channel in rgb)
            and not all(channel <= 12 for channel in rgb)
        ]
        _, rgb = max(usable or colors, key=lambda item: item[0])
        return "#%02x%02x%02x" % rgb
    except Exception:
        return None


def apply_docx_style(profile: Dict[str, Any], path: Path) -> None:
    try:
        doc = Document(path)
        try:
            normal = doc.styles["Normal"]
        except KeyError:
            normal = None
        if normal:
            if normal.font.name:
                profile["body_font"] = normal.font.name
            if normal.font.size:
                profile["body_size_pt"] = float(normal.font.size.pt)
        try:
            heading = doc.styles["Heading 1"]
        except KeyError:
            heading = None
        if heading:
            if heading.font.name:
                profile["heading_font"] = heading.font.name
            if heading.font.size:
                profile["heading_size_pt"] = float(heading.font.size.pt)
            if heading.font.color and heading.font.color.rgb:
                profile["primary_color"] = "#" + str(heading.font.color.rgb)
        if doc.sections:
            section = doc.sections[0]
            profile.update(
                {
                    "margin_top_in": section.top_margin.inches,
                    "margin_bottom_in": section.bottom_margin.inches,
                    "margin_left_in": section.left_margin.inches,
                    "margin_right_in": section.right_margin.inches,
                    "page_width_in": section.page_width.inches,
                    "page_height_in": section.page_height.inches,
                }
            )
    except Exception:
        # A design reference is advisory. A malformed reference must not prevent
        # the project from using the safe default profile.
        pass


def apply_pdf_layout(profile: Dict[str, Any], path: Path) -> None:
    try:
        reader = PdfReader(str(path))
        page = reader.pages[0]
        profile["page_width_in"] = float(page.mediabox.width) / 72.0
        profile["page_height_in"] = float(page.mediabox.height) / 72.0
    except Exception:
        pass


def normalize_profile(profile: Dict[str, Any]) -> Dict[str, Any]:
    default = default_profile()
    merged = {**default, **profile}
    merged["body_font"] = css_font(merged.get("body_font"))
    merged["heading_font"] = css_font(merged.get("heading_font"))
    merged["primary_color"] = css_color(merged.get("primary_color"))
    merged["text_color"] = css_color(merged.get("text_color"), "#222222")
    for key in (
        "body_size_pt",
        "heading_size_pt",
        "document_title_size_pt",
        "brand_line_size_pt",
        "page_width_in",
        "page_height_in",
        "margin_top_in",
        "margin_bottom_in",
        "margin_left_in",
        "margin_right_in",
        "line_height",
        "paragraph_spacing_pt",
    ):
        merged[key] = float(merged[key])
    return merged


def build_design_profile(req: DesignProfileRequest) -> Dict[str, Any]:
    project_root = project_dir(req.project_id).resolve()
    profile = default_profile()
    if req.organization_name:
        profile["organization_name"] = req.organization_name.strip()
    valid_assets: List[str] = []
    for raw in req.asset_paths:
        path = Path(raw)
        try:
            resolved = path.resolve()
            if project_root not in resolved.parents:
                continue
        except Exception:
            continue
        if not path.exists() or not path.is_file():
            continue
        valid_assets.append(str(path))
        ext = path.suffix.lower()
        if ext == ".docx":
            apply_docx_style(profile, path)
        elif ext == ".pdf":
            apply_pdf_layout(profile, path)
        elif ext in {".png", ".jpg", ".jpeg", ".webp"}:
            if not profile.get("logo_path"):
                profile["logo_path"] = str(path)
            color = dominant_hex(path)
            if color:
                profile["primary_color"] = color

    # Sponsor constraints are applied last so compliance overrides inspiration.
    sponsor_profiles = load_json(CONFIG_ROOT / "sponsor_formats.json", {})
    profile.update(sponsor_profiles.get(norm_sponsor(req.sponsor), {}))
    profile["sponsor"] = req.sponsor
    profile["assets"] = valid_assets
    profile = normalize_profile(profile)
    output = project_dir(req.project_id) / "design_profile.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(profile, indent=2))
    return profile


def load_profile(req: RenderRequest) -> Dict[str, Any]:
    # Export snapshots pass the exact approved design profile; do not reapply current
    # config/sponsor files or environment values, because that would make old snapshots
    # render differently after a software/configuration change.
    if req.design_profile is not None:
        return normalize_profile(dict(req.design_profile))
    path = project_dir(req.project_id) / "design_profile.json"
    profile = load_json(path, default_profile())
    if req.organization_name:
        profile["organization_name"] = req.organization_name.strip()
    if req.sponsor:
        overrides = load_json(CONFIG_ROOT / "sponsor_formats.json", {}).get(
            norm_sponsor(req.sponsor), {}
        )
        profile.update(overrides)
        profile["sponsor"] = req.sponsor
    return normalize_profile(profile)


def markdown_to_text(value: str) -> str:
    """Remove the small Markdown subset commonly emitted by drafting models."""
    value = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", value)
    value = re.sub(r"(\*\*|__)(.*?)\1", r"\2", value)
    value = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", value)
    value = value.replace("`", "")
    return value.strip()


def build_ast(req: RenderRequest) -> Dict[str, Any]:
    blocks: List[Dict[str, Any]] = []
    if req.include_document_title:
        blocks.append({"type": "document_title", "text": req.title})
    for section in req.sections:
        blocks.append(
            {
                "type": "heading",
                "level": 1,
                "text": section.title,
                "section_key": section.section_key,
                "version": section.version,
            }
        )
        first_body_block = True
        for paragraph in re.split(r"\n\s*\n", section.body or ""):
            paragraph = paragraph.strip()
            if paragraph:
                lines = paragraph.splitlines()
                if first_body_block and lines:
                    candidate = re.sub(r"^#{1,6}\s*", "", lines[0]).strip()
                    if candidate.casefold().rstrip(":") == section.title.strip().casefold().rstrip(":"):
                        lines = lines[1:]
                first_body_block = False
                paragraph = "\n".join(lines).strip()
                if not paragraph:
                    continue
                heading_match = re.fullmatch(r"#{1,6}\s+(.+)", paragraph)
                block_type = "subheading" if heading_match else "paragraph"
                text = markdown_to_text(heading_match.group(1) if heading_match else paragraph)
                blocks.append(
                    {
                        "type": block_type,
                        "text": text,
                        "section_key": section.section_key,
                        "version": section.version,
                    }
                )
    return {
        "version": 1,
        "project_id": req.project_id,
        "snapshot_id": req.snapshot_id,
        "title": req.title,
        "blocks": blocks,
    }


def logo_data_uri(path: Optional[str]) -> Optional[str]:
    if not path:
        return None
    file = Path(path)
    if not file.exists():
        return None
    mime = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }.get(file.suffix.lower())
    if not mime:
        return None
    return f"data:{mime};base64,{base64.b64encode(file.read_bytes()).decode()}"


def html_from_ast(ast: Dict[str, Any], profile: Dict[str, Any]) -> str:
    logo = logo_data_uri(profile.get("logo_path"))
    logo_html = f'<img class="logo" src="{logo}">' if logo else ""
    blocks: List[str] = []
    for block in ast["blocks"]:
        text = htmlmod.escape(block["text"])
        if block["type"] == "document_title":
            blocks.append(f'<h2 class="doc-title">{text}</h2>')
        elif block["type"] == "heading":
            blocks.append(f"<h1>{text}</h1>")
        elif block["type"] == "subheading":
            blocks.append(f"<h2>{text}</h2>")
        elif block["type"] == "paragraph":
            blocks.append(f"<p>{text}</p>")

    page_w, page_h = profile["page_width_in"], profile["page_height_in"]
    mt, mr, mb, ml = (
        profile["margin_top_in"],
        profile["margin_right_in"],
        profile["margin_bottom_in"],
        profile["margin_left_in"],
    )
    css = f"""
@page {{ size: {page_w}in {page_h}in; margin: {mt}in {mr}in {mb}in {ml}in; }}
html, body {{ margin: 0; padding: 0; }}
body {{ font-family: \"{profile['body_font']}\", sans-serif; font-size: {profile['body_size_pt']}pt; line-height: {profile['line_height']}; color: {profile['text_color']}; }}
.grant-page {{ box-sizing: border-box; background: #fff; margin: 0 auto; }}
.brand-row {{ font-size: {profile['brand_line_size_pt']}pt; font-weight: 700; display: flex; align-items: center; gap: 12px; margin-bottom: 18px; }}
.logo {{ max-height: .42in; max-width: 1.7in; object-fit: contain; }}
.doc-title {{ font-family: \"{profile['heading_font']}\", sans-serif; font-size: {profile['document_title_size_pt']}pt; margin: 0 0 16px; }}
h1 {{ font-family: \"{profile['heading_font']}\", sans-serif; font-size: {profile['heading_size_pt']}pt; margin: 14px 0 10px; border-bottom: 2px solid {profile['primary_color']}; padding-bottom: 6px; }}
h2 {{ font-family: \"{profile['heading_font']}\", sans-serif; font-size: {max(11.0, profile['heading_size_pt'] - 2)}pt; margin: 10px 0 6px; }}
p {{ margin: 0 0 {profile['paragraph_spacing_pt']}pt; white-space: pre-wrap; }}
@media screen {{
  body {{ background: #e7e7e7; padding: 16px; }}
  .grant-page {{ width: {page_w}in; min-height: {page_h}in; padding: {mt}in {mr}in {mb}in {ml}in; box-shadow: 0 3px 18px #0002; }}
}}
@media print {{
  .grant-page {{ width: auto; min-height: auto; padding: 0; box-shadow: none; }}
}}
"""
    brand = htmlmod.escape(profile.get("organization_name") or "")
    return (
        '<html><head><meta charset="utf-8"><style>'
        + css
        + '</style></head><body><article class="grant-page"><div class="brand-row">'
        + logo_html
        + f"<span>{brand}</span></div>"
        + "".join(blocks)
        + "</article></body></html>"
    )


def _diff_tokens(text: str) -> List[str]:
    return re.findall(r"\s+|[^\s]+", text or "")


def word_diff_html(baseline: str, proposed: str) -> str:
    old = _diff_tokens(baseline)
    new = _diff_tokens(proposed)
    matcher = difflib.SequenceMatcher(a=old, b=new, autojunk=False)
    parts: List[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            parts.append(htmlmod.escape("".join(new[j1:j2])))
        elif tag == "insert":
            parts.append('<mark class="agentic-add">' + htmlmod.escape("".join(new[j1:j2])) + "</mark>")
        elif tag == "delete":
            parts.append('<del class="agentic-remove">' + htmlmod.escape("".join(old[i1:i2])) + "</del>")
        else:
            parts.append('<del class="agentic-remove">' + htmlmod.escape("".join(old[i1:i2])) + "</del>")
            parts.append('<mark class="agentic-add">' + htmlmod.escape("".join(new[j1:j2])) + "</mark>")
    return "".join(parts)


def diff_preview_html(req: DiffPreviewRequest, profile: Dict[str, Any]) -> str:
    logo = logo_data_uri(profile.get("logo_path"))
    logo_html = f'<img class="logo" src="{logo}">' if logo else ""
    page_w, page_h = profile["page_width_in"], profile["page_height_in"]
    mt, mr, mb, ml = profile["margin_top_in"], profile["margin_right_in"], profile["margin_bottom_in"], profile["margin_left_in"]
    css = f"""
html,body{{margin:0;padding:0}}
body{{font-family:"{profile['body_font']}",sans-serif;font-size:{profile['body_size_pt']}pt;line-height:{profile['line_height']};color:{profile['text_color']};background:#e7e7e7;padding:16px}}
.grant-page{{box-sizing:border-box;background:#fff;margin:0 auto;width:{page_w}in;min-height:{page_h}in;padding:{mt}in {mr}in {mb}in {ml}in;box-shadow:0 3px 18px #0002}}
.brand-row{{font-size:{profile['brand_line_size_pt']}pt;font-weight:700;display:flex;align-items:center;gap:12px;margin-bottom:18px}}
.logo{{max-height:.42in;max-width:1.7in;object-fit:contain}}
h1{{font-family:"{profile['heading_font']}",sans-serif;font-size:{profile['heading_size_pt']}pt;margin:14px 0 10px;border-bottom:2px solid {profile['primary_color']};padding-bottom:6px}}
.diff-body{{white-space:pre-wrap;margin:0 0 {profile['paragraph_spacing_pt']}pt}}
.agentic-add{{background:#fff0a8;border-bottom:2px solid #c99400;padding:0 1px}}
.agentic-remove{{background:#ffe0e0;color:#8b3030;text-decoration:line-through;text-decoration-thickness:1.5px}}
.update-note{{font-family:Arial,sans-serif;font-size:9pt;background:#f4f7ff;border-left:4px solid #5267c9;padding:8px 10px;margin-bottom:14px}}
"""
    brand = htmlmod.escape(profile.get("organization_name") or "")
    title = htmlmod.escape(req.section.title)
    body = word_diff_html(req.baseline_body, req.section.body)
    reason = htmlmod.escape((req.update_reason or "Fresh public competitive intelligence").strip())
    summary = htmlmod.escape((req.update_summary or "Competitive positioning changed based on newly refreshed public evidence.").strip())
    diff_explanation = ("Highlighted text shows all changes relative to the pre-update version, including edits you made after the automatic proposal."
                        if req.includes_human_edits else
                        "Highlighted text shows the automatic proposal caused by the refreshed competitor intelligence.")
    notice = (f'<div class="update-note"><strong>⚡ {htmlmod.escape(COMPETITIVE_UPDATE_LABEL)}</strong> — {reason}. '
              f'{summary}<br><strong>{htmlmod.escape(diff_explanation)}</strong> Your previously approved version remains protected until you review and approve this version.</div>')
    return ('<html><head><meta charset="utf-8"><style>'+css+'</style></head><body><article class="grant-page">'
            +'<div class="brand-row">'+logo_html+f'<span>{brand}</span></div>'
            +notice
            +f'<h1>{title}</h1><div class="diff-body">{body}</div></article></body></html>')


def docx_from_ast(ast: Dict[str, Any], profile: Dict[str, Any], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(profile["page_width_in"])
    section.page_height = Inches(profile["page_height_in"])
    section.top_margin = Inches(profile["margin_top_in"])
    section.bottom_margin = Inches(profile["margin_bottom_in"])
    section.left_margin = Inches(profile["margin_left_in"])
    section.right_margin = Inches(profile["margin_right_in"])

    normal = doc.styles["Normal"]
    normal.font.name = profile["body_font"]
    normal.font.size = Pt(profile["body_size_pt"])
    for style_name, size in (
        ("Title", profile["document_title_size_pt"]),
        ("Heading 1", profile["heading_size_pt"]),
    ):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = profile["heading_font"]
            style.font.size = Pt(size)
            style.font.color.rgb = hex_rgb(profile["primary_color"])

    brand_paragraph = doc.add_paragraph()
    logo = profile.get("logo_path")
    if logo and Path(logo).exists():
        try:
            brand_paragraph.add_run().add_picture(str(logo), height=Inches(0.35))
        except Exception:
            pass
    brand_run = brand_paragraph.add_run(
        (" " if brand_paragraph.runs else "") + (profile.get("organization_name") or "")
    )
    brand_run.bold = True
    brand_run.font.size = Pt(profile["brand_line_size_pt"])

    for block in ast["blocks"]:
        if block["type"] == "document_title":
            doc.add_heading(block["text"], 0)
        elif block["type"] == "heading":
            doc.add_heading(block["text"], 1)
        elif block["type"] == "subheading":
            doc.add_heading(block["text"], 2)
        elif block["type"] == "paragraph":
            paragraph = doc.add_paragraph(block["text"])
            paragraph.paragraph_format.line_spacing = profile["line_height"]
            paragraph.paragraph_format.space_after = Pt(profile["paragraph_spacing_pt"])
    doc.save(path)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/design-profile")
def design_profile(req: DesignProfileRequest):
    return build_design_profile(req)


@app.post("/preview")
def preview(req: RenderRequest):
    profile = load_profile(req)
    ast = build_ast(req)
    return {"html": html_from_ast(ast, profile), "ast": ast, "design_profile": profile}


@app.post("/preview-diff")
def preview_diff(req: DiffPreviewRequest):
    render_req = RenderRequest(project_id=req.project_id,title=req.title,organization_name=req.organization_name,sections=[req.section],include_document_title=False,design_profile=req.design_profile)
    profile = load_profile(render_req)
    return {"html": diff_preview_html(req, profile), "design_profile": profile}



@app.post("/measure")
def measure(req: RenderRequest):
    profile = load_profile(req)
    ast = build_ast(req)
    pdf_bytes = HTML(string=html_from_ast(ast, profile), base_url=str(project_dir(req.project_id))).write_pdf()
    reader = PdfReader(io.BytesIO(pdf_bytes))
    by_section = {}
    total_words = 0
    for section in req.sections:
        words = len((section.body or "").split())
        total_words += words
        section_req = RenderRequest(
            project_id=req.project_id,
            title=section.title,
            sponsor=req.sponsor,
            organization_name=req.organization_name,
            sections=[section],
            include_document_title=False,
            design_profile=profile,
        )
        section_pdf = HTML(string=html_from_ast(build_ast(section_req), profile), base_url=str(project_dir(req.project_id))).write_pdf()
        key = section.section_key or re.sub(r"[^a-z0-9]+", "_", section.title.lower()).strip("_")
        by_section[key] = {"words": words, "pages": len(PdfReader(io.BytesIO(section_pdf)).pages)}
    return {"page_count": len(reader.pages), "word_count": total_words, "sections": by_section, "design_profile": profile}

@app.post("/package")
def package(req: PackageRequest):
    root = project_dir(req.project_id)
    output = OUTPUT_ROOT / req.project_id / "final"
    output.mkdir(parents=True, exist_ok=True)
    package_path = output / f"{safe(req.title)}_snapshot_{req.snapshot_id}_submission_package.zip"
    temporary_path = package_path.with_suffix(".zip.tmp")
    manifest = dict(req.manifest)
    manifest.update({"project_id": req.project_id, "snapshot_id": req.snapshot_id, "generated_paths": req.generated_paths})
    temporary_path.unlink(missing_ok=True)
    try:
        with zipfile.ZipFile(temporary_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            packaged_proposals = []
            for raw in req.generated_paths:
                path = Path(raw)
                try:
                    resolved = path.resolve()
                    if OUTPUT_ROOT.resolve() not in resolved.parents or not path.is_file():
                        continue
                except Exception:
                    continue
                digest = hashlib.sha256(path.read_bytes()).hexdigest()
                z.write(path, arcname=f"proposal/{path.name}")
                packaged_proposals.append({"filename": path.name, "sha256": digest})
            # Package only artifacts registered in the immutable snapshot manifest.
            # Files merely present in the shared submission directory are intentionally excluded.
            submission = (root / "submission").resolve()
            registered = manifest.get("submission_artifacts") or []
            packaged_artifacts = []
            for item in registered:
                if not isinstance(item, dict):
                    continue
                raw_path = item.get("path")
                if not raw_path:
                    continue
                try:
                    path = Path(raw_path).resolve()
                    if submission != path.parent and submission not in path.parents:
                        raise HTTPException(400, "registered artifact is outside the project submission directory")
                    if not path.is_file():
                        raise HTTPException(400, f"registered artifact is missing: {path.name}")
                except HTTPException:
                    raise
                except Exception as exc:
                    raise HTTPException(400, f"invalid registered artifact path: {exc}")
                digest = hashlib.sha256(path.read_bytes()).hexdigest()
                expected = str(item.get("sha256") or "").lower()
                if not expected or digest != expected:
                    raise HTTPException(409, f"registered artifact checksum changed: {path.name}")
                slot = safe(str(item.get("slot") or "attachments"))
                filename = safe(str(item.get("filename") or path.name))
                z.write(path, arcname=f"attachments/{slot}/{filename}")
                packaged_artifacts.append({**item, "packaged_sha256": digest})
            manifest["packaged_proposals"] = packaged_proposals
            manifest["packaged_artifacts"] = packaged_artifacts
            z.writestr("submission_manifest.json", json.dumps(manifest, indent=2))
        temporary_path.replace(package_path)
    except Exception:
        temporary_path.unlink(missing_ok=True)
        raise
    return {"path": str(package_path)}

@app.post("/render")
def render(req: RenderRequest):
    if req.format not in ("docx", "pdf"):
        raise HTTPException(400, "format must be docx or pdf")
    profile = load_profile(req)
    ast = build_ast(req)
    output = OUTPUT_ROOT / req.project_id / "final"
    output.mkdir(parents=True, exist_ok=True)
    suffix = f"_snapshot_{req.snapshot_id}" if req.snapshot_id is not None else "_Final"
    path = output / f"{safe(req.title)}{suffix}.{req.format}"
    if req.format == "docx":
        docx_from_ast(ast, profile, path)
    else:
        HTML(string=html_from_ast(ast, profile), base_url=str(project_dir(req.project_id))).write_pdf(path)
    (output / f"{safe(req.title)}{suffix}.document_ast.json").write_text(
        json.dumps(ast, indent=2)
    )
    return {
        "path": str(path),
        "ast_version": ast["version"],
        "snapshot_id": req.snapshot_id,
        "design_profile": profile,
    }
